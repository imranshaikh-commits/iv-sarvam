"""
Pass 2 — DOCX branding for Inspirit Vision (IV) proposals.

This module owns the visual identity applied to generated Word documents:
the IV logo, the navy/orange colour theme, running header/footer, section
dividers, heading styles and the client-logo placeholder on the title page.

Design-import rule (mirrors document_engine): this module is a LEAF. It is
imported BY document_engine and must NOT import it back, so there is no
circular dependency. It has no secrets and no network access, which keeps the
DOCX path renderable offline (keyless smoke test).

Brand identity:
  * primary accent  navy-purple  #231154  (headings, headers/footers, dividers)
  * secondary accent orange      #E85A24  (single accent: ticks, divider tail)
  * neutrals for body text.
Colour usage is intentionally restrained: navy structures the page, orange is
the one accent, everything else is neutral (design-foundations: restraint,
1 accent + neutrals, WCAG-AA contrast on white).
"""

from __future__ import annotations

import os
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --- palette ----------------------------------------------------------------
NAVY = RGBColor(0x23, 0x11, 0x54)         # primary accent
ORANGE = RGBColor(0xE8, 0x5A, 0x24)       # secondary accent (single accent)
NEUTRAL_DARK = RGBColor(0x33, 0x33, 0x3A)  # H2 / darker neutral (AA on white)
NEUTRAL_BODY = RGBColor(0x22, 0x22, 0x22)  # body text
NEUTRAL_MUTED = RGBColor(0x6B, 0x6B, 0x72)  # captions / placeholder label
DRAFT_RED = RGBColor(0xB0, 0x00, 0x00)     # safety banner (kept from Pass 1)

_NAVY_HEX = "231154"
_ORANGE_HEX = "E85A24"

BODY_FONT = "Calibri"      # sans-serif
HEADING_FONT = "Calibri"   # sans-serif

# --- asset locations (packaged next to this module) -------------------------
_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
LOGO_FULL = os.path.join(_ASSETS_DIR, "iv_logo.png")       # ~1600px wide
LOGO_HEADER = os.path.join(_ASSETS_DIR, "iv_logo_header.png")  # ~400px wide

# usable text width for US-Letter with the python-docx default 1" margins.
_CONTENT_WIDTH_IN = 6.5


# ---------------------------------------------------------------------------
# low-level OOXML helpers
# ---------------------------------------------------------------------------
def _force_font(style, font_name: str) -> None:
    """Pin ascii/hAnsi (and EA) font mapping on a style (python-docx quirk)."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font_name)


def _shd(fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    return shd


def _set_cell_background(cell, fill_hex: str) -> None:
    cell._tc.get_or_add_tcPr().append(_shd(fill_hex))


def _zero_cell_margins(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "start", "bottom", "end", "left", "right"):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _set_row_exact_height(row, points: float) -> None:
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(points * 20)))  # twips
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _paragraph_border(paragraph, edges: dict) -> None:
    """edges: {"left"|"top"|"bottom"|"right": {"sz","color","space","val"}}."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        if edge not in edges:
            continue
        spec = edges[edge]
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), spec.get("val", "single"))
        e.set(qn("w:sz"), str(spec.get("sz", 6)))
        e.set(qn("w:space"), str(spec.get("space", 4)))
        e.set(qn("w:color"), spec.get("color", _NAVY_HEX))
        pbdr.append(e)
    pPr.append(pbdr)


def _add_right_tab(paragraph, position_in: float = _CONTENT_WIDTH_IN) -> None:
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(position_in), WD_TAB_ALIGNMENT.RIGHT
    )


def _add_page_field(run) -> None:
    """Append a live Word PAGE field into an existing run."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _set_page_number_start(section, start: int = 0) -> None:
    """Restart page numbering so content counts from 1 after the title page.

    Title page = page ``start`` (0, hidden via different-first-page); the first
    content page becomes page 1.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))


# ---------------------------------------------------------------------------
# public: base styles
# ---------------------------------------------------------------------------
def configure_base_styles(document: Document) -> None:
    """Apply IV typography + colour to Normal/Heading styles.

    Heading 1/2 keep their built-in style *names* so the refreshable TOC field
    (\\o "1-3") still discovers them — only their appearance changes.
    """
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = NEUTRAL_BODY
    _force_font(normal, BODY_FONT)

    h1 = document.styles["Heading 1"]
    h1.font.name = HEADING_FONT
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    _force_font(h1, HEADING_FONT)

    h2 = document.styles["Heading 2"]
    h2.font.name = HEADING_FONT
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = NEUTRAL_DARK
    _force_font(h2, HEADING_FONT)


# ---------------------------------------------------------------------------
# public: dividers & headings
# ---------------------------------------------------------------------------
def add_divider(document: Document, navy_frac: float = 0.72,
                thickness_pt: float = 4.0) -> None:
    """A thin two-colour rule: a wide navy segment tailed by an orange accent."""
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cells = table.rows[0].cells
    cells[0].width = Inches(_CONTENT_WIDTH_IN * navy_frac)
    cells[1].width = Inches(_CONTENT_WIDTH_IN * (1 - navy_frac))
    _set_cell_background(cells[0], _NAVY_HEX)
    _set_cell_background(cells[1], _ORANGE_HEX)
    for c in cells:
        _zero_cell_margins(c)
        pf = c.paragraphs[0].paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = c.paragraphs[0].add_run(" ")
        run.font.size = Pt(1)
    _set_row_exact_height(table.rows[0], thickness_pt)


def add_section_heading(document: Document, text: str):
    """Start an H1 section: orange tick, then a navy left-border accent heading.

    The orange tick lives in its own (non-heading) paragraph so it never leaks
    into the TOC; the heading keeps the built-in Heading 1 style for the TOC.
    Returns the heading paragraph so callers can append inline flags.
    """
    tick = document.add_paragraph()
    tick.paragraph_format.space_before = Pt(16)
    tick.paragraph_format.space_after = Pt(0)
    tick.paragraph_format.left_indent = Pt(2)
    trun = tick.add_run("▬▬")  # ▬▬ short orange tick
    trun.font.color.rgb = ORANGE
    trun.font.size = Pt(9)
    trun.bold = True

    heading = document.add_heading(text, level=1)
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.left_indent = Pt(12)
    _paragraph_border(heading, {"left": {"sz": 22, "color": _NAVY_HEX, "space": 12}})
    return heading


# ---------------------------------------------------------------------------
# public: running header & footer
# ---------------------------------------------------------------------------
def apply_header_footer(document: Document, client_name: str) -> None:
    """Install the running header/footer and restart content page numbering.

    Header: small IV logo (left) + "Technical Proposal — {client}" (right),
            thin navy rule under.
    Footer: "Inspirit Vision — Confidential" (left) + "Page N" (right),
            thin navy rule above.
    The title page (first page) is left blank via different-first-page.
    """
    section = document.sections[0]
    section.different_first_page_header_footer = True
    _set_page_number_start(section, start=0)

    # --- header -------------------------------------------------------------
    header = section.header
    header.is_linked_to_previous = False
    hpara = header.paragraphs[0]
    hpara.text = ""
    hpara.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_right_tab(hpara)
    if os.path.exists(LOGO_HEADER):
        hpara.add_run().add_picture(LOGO_HEADER, width=Inches(1.15))
    else:  # pragma: no cover - asset always shipped
        brand = hpara.add_run("Inspirit Vision")
        brand.bold = True
        brand.font.color.rgb = NAVY
        brand.font.size = Pt(10)
    right = hpara.add_run(f"\tTechnical Proposal — {client_name}")
    right.font.size = Pt(9)
    right.font.color.rgb = NAVY
    _paragraph_border(hpara, {"bottom": {"sz": 6, "color": _NAVY_HEX, "space": 4}})

    # --- footer -------------------------------------------------------------
    footer = section.footer
    footer.is_linked_to_previous = False
    fpara = footer.paragraphs[0]
    fpara.text = ""
    fpara.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_right_tab(fpara)
    _paragraph_border(fpara, {"top": {"sz": 6, "color": _NAVY_HEX, "space": 4}})
    left = fpara.add_run("Inspirit Vision — Confidential")
    left.font.size = Pt(8)
    left.font.color.rgb = NEUTRAL_MUTED
    pnum = fpara.add_run("\tPage ")
    pnum.font.size = Pt(8)
    pnum.font.color.rgb = NEUTRAL_MUTED
    field_run = fpara.add_run()
    field_run.font.size = Pt(8)
    field_run.font.color.rgb = NEUTRAL_MUTED
    _add_page_field(field_run)


# ---------------------------------------------------------------------------
# public: title page
# ---------------------------------------------------------------------------
def _add_client_logo_block(document: Document,
                           client_logo_path: Optional[str]) -> None:
    """Bordered box on the title page: embeds a client logo or a placeholder."""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(2.4)
    _set_cell_border(cell, sz=8, color="B8B8C0")
    _set_row_exact_height(table.rows[0], 62)  # ~0.86in box
    cpara = cell.paragraphs[0]
    cpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if client_logo_path and os.path.exists(client_logo_path):
        cpara.add_run().add_picture(client_logo_path, width=Inches(2.0))
    else:
        label = cpara.add_run("Client Logo")
        label.font.color.rgb = NEUTRAL_MUTED
        label.font.size = Pt(11)
        label.italic = True


def _set_cell_border(cell, sz: int = 8, color: str = "B8B8C0") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def add_title_page(document: Document, metadata: dict,
                   client_logo_path: Optional[str] = None) -> None:
    """Render the branded title page (IV logo, title block, divider, client box).

    metadata keys: client_name, proposal_type, iam_vendor (optional),
    generated_at (optional), version (optional).
    """
    client_name = metadata.get("client_name") or "Client"
    proposal_type = metadata.get("proposal_type") or "implementation"
    iam_vendor = metadata.get("iam_vendor") or ""
    generated_at = metadata.get("generated_at") or ""
    version = metadata.get("version") or "Draft v1.0"

    document.add_paragraph()

    # IV logo centered near the top.
    if os.path.exists(LOGO_FULL):
        logo_p = document.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.add_run().add_picture(LOGO_FULL, width=Inches(3.3))
    document.add_paragraph()

    add_divider(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    trun = title.add_run(f"{proposal_type.upper()} PROPOSAL")
    trun.bold = True
    trun.font.size = Pt(30)
    trun.font.color.rgb = NAVY
    trun.font.name = HEADING_FONT

    prepared = document.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prun = prepared.add_run(f"Prepared for {client_name}")
    prun.font.size = Pt(18)
    prun.font.color.rgb = NAVY

    if iam_vendor:
        vend = document.add_paragraph()
        vend.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vrun = vend.add_run(f"IAM Platform: {iam_vendor}")
        vrun.font.size = Pt(13)
        vrun.font.color.rgb = NEUTRAL_DARK

    meta_line = document.add_paragraph()
    meta_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_line.paragraph_format.space_before = Pt(6)
    bits = [b for b in (generated_at, version) if b]
    mrun = meta_line.add_run("  •  ".join(bits))
    mrun.font.size = Pt(11)
    mrun.font.color.rgb = NEUTRAL_MUTED

    mark = document.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark.paragraph_format.space_before = Pt(10)
    mkrun = mark.add_run("DRAFT — FOR INTERNAL REVIEW ONLY  ·  CONFIDENTIAL")
    mkrun.bold = True
    mkrun.font.size = Pt(12)
    mkrun.font.color.rgb = DRAFT_RED

    document.add_paragraph()
    _add_client_logo_block(document, client_logo_path)
