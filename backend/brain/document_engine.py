"""
Sprint 5 — document-production engine.

Turns chat drafts into a downloadable, formatted Word (.docx) proposal.

Design note (circular-import avoidance):
  This module MUST NOT import app.py. app.py imports `generate_proposal` from
  here, so the dependency flows one way only. The reusable brain helpers
  (embed_query, retrieve_chunks, build_grounded_system, run_compliance_matrix,
  render_matrix_markdown) are PASSED IN as callables by the endpoint. That keeps
  the module importable with no secrets present, which is what lets the smoke
  test exercise it offline by passing stubs.

  The single OpenRouter network call is isolated in ``draft_with_openrouter`` so
  tests can monkeypatch exactly that one function.
"""

from __future__ import annotations

import asyncio
import io
import logging
import os
import re
from datetime import datetime, timezone
from typing import Awaitable, Callable, Optional

import httpx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import branding
import document_qa
from proposal_templates import (
    COMPLIANCE_SECTION_ID,
    SUBSECTION_FACETS,
    DepthTier,
    SectionSpec,
    get_depth_tier,
    get_template,
    topic_for,
)
import asset_selection

# Two per section: IV's proposals carry 37 images across 11 sections, and a
# third of those are per-deal architecture drawings we cannot reuse.
_ASSETS_PER_SECTION = int(os.environ.get("SHILPI_ASSETS_PER_SECTION", "2"))

log = logging.getLogger("shilpi-brain.doc-engine")

# --- config (env with safe defaults; NEVER required at import) --------------
# app.py hard-requires OPENROUTER_API_KEY at import; this module must not, so
# that it stays importable in a keyless environment (smoke test / CI).
OPENROUTER_BASE = os.environ.get("OPENROUTER_BASE", "https://openrouter.ai/api/v1")
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
# LLM models (kept in sync with app.py). document_engine must NOT import app
# (circular-import rule), so the constants are duplicated here — including the
# env names, so a swap moves BOTH the chat path and the drafting path. Setting
# only one would leave the proposal drafted by a different model than the chat,
# which is exactly the kind of split that makes a comparison run meaningless.
PRIMARY_LLM_MODEL = os.environ.get("SHILPI_PRIMARY_MODEL", "").strip() or "z-ai/glm-5.2"
FALLBACK_LLM_MODEL = os.environ.get("SHILPI_FALLBACK_MODEL", "").strip() or "qwen/qwen3-235b-a22b-2507"
TOP_K = int(os.environ.get("TOP_K", "8"))
DOC_CONCURRENCY = int(os.environ.get("DOC_CONCURRENCY", os.environ.get("COMPLIANCE_CONCURRENCY", "3")))

# Below this max-similarity the evidence is considered weak and the section is
# flagged for SME review.
WEAK_EVIDENCE_THRESHOLD = float(os.environ.get("DOC_WEAK_EVIDENCE", "0.55"))

SME_REVIEW_MARKER = "[SME REVIEW]"
# Marks a grounded, generic placeholder emitted when a subsection LLM call
# returns null/empty even after one retry. Not a fabricated fact — an explicit
# assumption that an SME must confirm/replace before client use.
ASSUMPTION_MARKER = "[ASSUMPTION]"

# Type aliases for the passed-in brain helpers.
EmbedFn = Callable[[httpx.AsyncClient, str], Awaitable[list[float]]]
RetrieveFn = Callable[..., Awaitable[list[dict]]]
BuildSystemFn = Callable[[list[dict]], str]

_SECTION_SYSTEM_TEMPLATE = """You are Shilpi, InspiritVision's internal proposal assistant (an IAM consulting firm).
You are drafting the "{title}" section of a {proposal_type} proposal for {client_name}{vendor_clause}.

SECTION PURPOSE: {purpose}

HARD RULES (non-negotiable):
1. Ground every material technical claim in the EVIDENCE below. Cite inline like [1], [3] referring to evidence numbers.
2. NEVER invent product versions, compliance/regulatory claims, pricing, SLAs, or client commitments.
   If the evidence does not support something you need, write a literal "{marker}" note explaining what is missing.
3. Draft ONLY this section as clean proposal prose (no markdown headings — the heading is added by the document builder). Be concise and specific; no filler.
4. Prefer specific technical content from the evidence (architectures, connectors, workflows, timelines, volumetrics) over generic methodology boilerplate.
5. This is a DRAFT for human review, never client-ready.

{evidence}

OUTPUT DISCIPLINE (mandatory):
- Write the section as it will appear in the client proposal. NEVER narrate your
  own research process. Do not write "the retrieved evidence", "the corpus",
  "Note on Evidence Applicability", or any commentary about what your sources do
  or do not contain. Sources are acknowledged ONLY through inline [N] markers.
- Where a fact is genuinely unavailable, insert the review marker once and move
  on. Do not explain at length what is missing or why.
- Never emit reasoning, deliberation or <think> blocks. Output only the section.
- Stop when the section is complete. Do not pad to fill space, and never repeat a
  word or phrase to extend length.
"""


def _vendor_clause(iam_vendor: Optional[str]) -> str:
    return f" using {iam_vendor}" if iam_vendor else ""


# Hard ceiling on any single draft call's token budget. Pass 3 depth tiers vary
# the budget DOWN for leaner tiers but must never raise a call above this — depth
# comes from more (fanned-out) calls, not from one runaway call.
# Lowered from 3500 after the Amlak run: the cap was raised to chase page count,
# but the generated proposal came out 38% LONGER than the human original while
# padding with degenerate synonym chains. Length was never the gap — fidelity
# was. 2500 leaves ample room for a substantive subsection.
MAX_DRAFT_TOKENS = 2500


# ---------------------------------------------------------------------------
# Discovery context routing
#
# The 22-area discovery interview captures sizing, timeline, integrations,
# commercials, NFRs and delivery model — and until this map existed NONE of it
# reached the drafting engine. ``generate_proposal`` accepted only ``rfp_text``,
# which app.py filled with ``business_objectives``, so 21 of 22 areas were
# collected, stored in Supabase and then discarded. The generated Amlak proposal
# complained that it lacked "Amlak-specific volumetrics/timeline/pricing" — it
# was telling the truth.
#
# Routing per section rather than passing the whole answers dict everywhere is
# deliberate: a single blob would bloat every call and invite every section to
# restate everything. Sizing belongs to architecture, milestones to commercials,
# pain points to the executive summary.
# Section id -> the discovery fields that section needs.
#
# THIS MAP MUST TRACK proposal_templates.py AND intake_template.py. It silently
# stopped doing both:
#
#   * Keys were the OLD section ids (client_context, solution_architecture,
#     technical_approach, implementation_methodology, integration_points). When
#     the template was rebuilt to IV's house structure every section was
#     renamed, and this map was not. Result: 11 of 12 implementation sections
#     received NO discovery answers at all -- only executive_summary matched.
#   * Field names were invented rather than taken from the intake schema
#     ("engagement_duration" for `duration`, "environments" for `envs`,
#     "sod_required" for `sod`, "hrms_authoritative_source" for
#     `integration_hrms`), so several that did match resolved to nothing.
#
# The visible symptom was 42 [SME REVIEW] markers in run 9, including on UAT
# sizing, connectors, hypercare and the support model -- all of which the
# consultant HAD supplied. The model was not hedging; it genuinely had not been
# told. A sizing table drafted without sizing inputs can only be generic, so
# this also explains part of the table thinness against IV's original.
#
# `test_document_engine.py` asserts every section id in every template appears
# here, and that every field name exists in the intake schema. Those two tests
# are the reason this cannot drift again.
_SECTION_DISCOVERY_FIELDS: dict[str, tuple[str, ...]] = {
    # --- implementation template ---
    "executive_summary": (
        "business_objectives", "pain_points", "differentiators", "decision_criteria",
        "duration", "app_count", "user_count", "audience",
    ),
    "company_profile": ("partner_positioning", "vendor_partner_positioning"),
    "similar_experience": (
        "case_studies_to_highlight", "case_studies_include", "case_studies_exclude",
        "similar_projects", "partner_positioning",
    ),
    "scope_understanding": (
        "business_objectives", "in_scope", "out_of_scope", "current_state",
        "existing_iam_platform", "pain_points", "app_count", "user_count",
        "identity_types", "apps_to_onboard",
    ),
    "solution_overview": (
        "sod", "access_review_cadence", "target_integrations", "audit",
        "monitoring", "identity_types", "differentiators", "regulations",
    ),
    "proposed_solution": (
        "deployment_model", "hardware_sizing_inputs", "cluster_topology",
        "ha_dr_requirements", "rto_rpo", "security_architecture_needs",
        "regions", "network_zones", "envs", "availability", "scalability",
        "performance", "target_integrations", "integration_hrms", "ad_exchange",
        "idp_sso", "apps_to_onboard", "app_count", "user_count", "directories",
        "current_hrms", "current_idp", "source_of_truth",
    ),
    "implementation_approach": (
        "delivery_phases", "delivery_milestones", "governance", "raci",
        "client_responsibilities", "dependencies", "assumptions",
    ),
    "project_timeline": (
        "duration", "timeline_milestones", "go_live_date", "delivery_phases",
        "delivery_milestones", "app_count",
    ),
    "assumptions_responsibilities": (
        "assumptions", "dependencies", "client_responsibilities", "out_of_scope",
        "rto_rpo", "versions",
    ),
    "knowledge_transfer": (
        "training", "kt", "hypercare", "support_model", "post_sla",
        "postgolive_reporting_cadence",
    ),
    "commercial": (
        "license_included", "pricing_model", "payment_milestones", "taxes",
        "travel", "support_terms", "validity_period", "currency",
    ),
    "compliance_matrix": ("rfp_text", "regulations", "certifications"),

    # --- migration template ---
    "current_state": (
        "current_state", "existing_iam_platform", "versions", "directories",
        "current_hrms", "current_idp", "source_of_truth", "pain_points",
        "is_migration", "tenants", "apps_to_onboard", "app_count",
    ),
    "target_state": (
        "deployment_model", "hardware_sizing_inputs", "cluster_topology",
        "ha_dr_requirements", "envs", "regions", "network_zones",
        "security_architecture_needs", "availability", "scalability",
    ),
    "migration_strategy": (
        "is_migration", "existing_iam_platform", "apps_to_onboard", "app_count",
        "identity_types", "out_of_scope", "delivery_phases", "data_residency",
    ),
    "rollback_risk": (
        "ha_dr_requirements", "rto_rpo", "availability", "dependencies",
        "assumptions", "envs",
    ),
    "decommissioning": (
        "existing_iam_platform", "data_retention", "out_of_scope",
        "license_included", "support_terms",
    ),

    # --- mss template ---
    "operating_model": (
        "support_model", "postgolive_reporting_cadence", "governance", "raci",
        "training", "kt", "hypercare",
    ),
    "service_model": (
        "support_model", "post_sla", "hypercare", "training", "kt",
        "delivery_phases", "envs",
    ),
    "sla_coverage": ("post_sla", "support_model", "availability", "monitoring", "audit"),
    "escalation_incident": ("post_sla", "support_model", "monitoring", "governance"),
    "assumptions_open_questions": (
        "assumptions", "dependencies", "client_responsibilities", "out_of_scope",
        "rto_rpo", "versions",
    ),
}

# Fields every section benefits from knowing.
_UNIVERSAL_DISCOVERY_FIELDS: tuple[str, ...] = (
    "client_name", "iam_vendor", "proposal_type", "industry", "country",
)

# Never route these into drafting prompts: control fields, or content that must
# not be paraphrased into client-facing prose.
_EXCLUDED_DISCOVERY_FIELDS = frozenset({
    "proposal_depth", "diagram_count", "required_diagram_types", "rfp_text",
    "_diagram_plan", "case_studies_exclude", "logo_url", "client_logo",
})


def _humanise_field(key: str) -> str:
    return key.replace("_", " ").strip().title()


def discovery_context_for(section_id: str, answers: Optional[dict],
                          limit: int = 3500) -> str:
    """Render the discovery answers relevant to one section.

    Returns an empty string when nothing relevant was captured, so the prompt is
    unchanged for callers that have no intake data (the REST path).
    """
    if not answers:
        return ""
    wanted = _SECTION_DISCOVERY_FIELDS.get(section_id, ())
    keys = list(_UNIVERSAL_DISCOVERY_FIELDS) + [k for k in wanted
                                                if k not in _UNIVERSAL_DISCOVERY_FIELDS]
    lines: list[str] = []
    for key in keys:
        if key in _EXCLUDED_DISCOVERY_FIELDS:
            continue
        val = answers.get(key)
        if val is None:
            continue
        text = " ".join(str(val).split())
        if not text or text.lower() in ("skip", "none", "n/a", "na", "-"):
            continue
        lines.append(f"- {_humanise_field(key)}: {text}")
    if not lines:
        return ""
    block = "\n".join(lines)
    if len(block) > limit:
        block = block[:limit].rsplit("\n", 1)[0]
    return block


def _draft_payload(model: str, system_prompt: str, user_prompt: str,
                   include_frequency_penalty: bool = True,
                   max_tokens: int = MAX_DRAFT_TOKENS) -> dict:
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "stream": False,
        "temperature": 0.4,
        # Cap runaway generation (same repetition-spiral risk as the compliance
        # classifier). LOW frequency_penalty preserves grounded citation/vendor terms.
        # Clamp to the hard ceiling so a bad depth config can never inflate a call.
        "max_tokens": min(int(max_tokens), MAX_DRAFT_TOKENS),
    }
    if include_frequency_penalty:
        # Raised from 0.2: the low value was chosen to preserve repeated vendor
        # and product terms, but it also let an exhausted model pad with synonym
        # chains. 0.45 still tolerates legitimate repetition of "SailPoint" or
        # "IdentityIQ" while discouraging runaway filler.
        payload["frequency_penalty"] = 0.45
    return payload


async def _post_draft(client: httpx.AsyncClient, payload: dict) -> str:
    resp = await client.post(
        f"{OPENROUTER_BASE}/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        },
        json=payload,
        timeout=180,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


async def draft_with_openrouter(
    client: httpx.AsyncClient,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int = MAX_DRAFT_TOKENS,
) -> str:
    """Isolated OpenRouter chat call with primary->fallback. Monkeypatched by the smoke test.

    Tries PRIMARY_LLM_MODEL, then FALLBACK_LLM_MODEL on HTTP/network/timeout
    error. Defensive: a 400 caused by an unsupported param (frequency_penalty)
    triggers a same-model retry WITHOUT frequency_penalty before falling back.

    ``max_tokens`` is the per-call budget from the active depth tier; it is
    clamped to MAX_DRAFT_TOKENS inside ``_draft_payload``.
    """
    last_exc: Exception | None = None
    for model in (PRIMARY_LLM_MODEL, FALLBACK_LLM_MODEL):
        try:
            content = await _post_draft(
                client, _draft_payload(model, system_prompt, user_prompt, max_tokens=max_tokens))
            log.info("OpenRouter draft model=%s", model)
            return content
        except httpx.HTTPStatusError as e:
            last_exc = e
            # A 400 may be an unsupported-param error (e.g. frequency_penalty).
            # Retry the SAME model once without it before falling back.
            if e.response is not None and e.response.status_code == 400:
                try:
                    content = await _post_draft(
                        client, _draft_payload(model, system_prompt, user_prompt,
                                               include_frequency_penalty=False,
                                               max_tokens=max_tokens))
                    log.info("OpenRouter draft model=%s (no frequency_penalty)", model)
                    return content
                except (httpx.HTTPStatusError, httpx.RequestError) as e2:
                    last_exc = e2
            if model == FALLBACK_LLM_MODEL:
                raise
            log.warning("draft failed on primary %s (%s); falling back to %s",
                        model, e, FALLBACK_LLM_MODEL)
        except httpx.RequestError as e:
            last_exc = e
            if model == FALLBACK_LLM_MODEL:
                raise
            log.warning("draft failed on primary %s (%s); falling back to %s",
                        model, e, FALLBACK_LLM_MODEL)
    raise last_exc  # pragma: no cover


def _fanout_queries(section_spec: SectionSpec, context: dict, fanout: int) -> list[str]:
    """Build up to ``fanout`` distinct retrieval queries for a section.

    Query 0 is the section's base query. Extra queries append a subsection facet
    keyword so retrieval surfaces evidence for different aspects of the section.
    """
    base = section_spec.render_query(context)
    if fanout <= 1:
        return [base]
    queries = [base]
    for _title, facet in SUBSECTION_FACETS[: max(0, fanout - 1)]:
        queries.append(f"{base} — {facet}")
    return queries[:fanout]


async def _retrieve_fanout(
    client: httpx.AsyncClient,
    section_spec: SectionSpec,
    context: dict,
    *,
    embed_fn: EmbedFn,
    retrieve_fn: RetrieveFn,
    top_k: int,
    fanout: int,
) -> list[dict]:
    """Run fanned-out retrieval, then merge + dedupe chunks by text.

    Deduped set is sorted by similarity (desc) and capped so wider fan-out gives
    richer evidence without an unbounded evidence block.
    """
    seen: set[str] = set()
    merged: list[dict] = []
    for query in _fanout_queries(section_spec, context, fanout):
        try:
            embedding = await embed_fn(client, query)
            # Pass the proposal type through so retrieval reserves slots for
            # the same kind of engagement. Without it a migration section is
            # drafted from greenfield implementation proposals: measured at 1 of
            # 8 chunks from a migration proposal, against a 35.5% base rate.
            chunks = await retrieve_fn(
                client, embedding, query, k=top_k,
                proposal_type=context.get("proposal_type") or None,
                # Retrieval runs ONCE per section, before subsections are
                # drafted, so there is no subsection heading here. The fan-out
                # QUERY text is used instead: it is derived from the section's
                # subsections and carries their content words ("sizing", "RACI",
                # "why SailPoint"), which is what topic_for matches on. Coarser
                # than a true per-subsection retrieval, and cheap.
                section_topic=topic_for(section_spec.id, query))
        except Exception as e:  # fail soft: one failed query must not sink the section
            log.error("draft_section retrieval failed for %s: %s", section_spec.id, e)
            chunks = []
        for c in chunks:
            key = (c.get("chunk_text") or c.get("heading") or "")[:160]
            if key and key in seen:
                continue
            if key:
                seen.add(key)
            merged.append(c)
    merged.sort(key=lambda c: float(c.get("similarity") or 0.0), reverse=True)
    # Cap evidence to keep prompts bounded: base top_k, plus headroom per extra query.
    cap = top_k * max(1, fanout)
    return merged[:cap]


# Reasoning models (GLM 5.2 among them) emit internal deliberation that must
# never reach a client document. A live Amlak proposal contained a raw </think>
# tag plus seven verbatim repetitions of the model's own musing about evidence
# quality. Nothing anywhere stripped these.
_THINK_BLOCK_RE = re.compile(r"<think\b[^>]*>.*?</think\s*>", re.S | re.I)
_STRAY_THINK_TAG_RE = re.compile(r"</?think\s*[^>]*>", re.I)

# Whole paragraphs where the model narrates its own retrieval instead of writing
# the proposal. These are answer content, not reasoning tokens, so the tag strip
# above does not catch them.
_META_PARAGRAPH_RE = re.compile(
    r"(?m)^[^\n]*\b("
    r"note on evidence applicability"
    r"|the retrieved evidence"
    r"|retrieved evidence (originates|comes|is drawn)"
    r"|based on the (retrieved )?(corpus|evidence base)"
    r"|the (corpus|evidence) (does not|doesn't) contain"
    r")\b[^\n]*$",
    re.I,
)


def strip_model_artifacts(text: str) -> str:
    """Remove reasoning blocks and retrieval meta-commentary from drafted text.

    Deliberately conservative: it removes whole lines that narrate the drafting
    process, and never touches ``[N]`` citation markers or SME-review markers,
    which are intentional signals for the human reviewer.
    """
    if not text:
        return text
    out = _THINK_BLOCK_RE.sub("", text)
    out = _STRAY_THINK_TAG_RE.sub("", out)
    out = _META_PARAGRAPH_RE.sub("", out)
    # Collapse the blank lines and orphaned rules those removals leave behind.
    out = re.sub(r"(?m)^\s*-{3,}\s*$\n(?=\s*$)", "", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def strip_degenerate_tail(text: str, run_threshold: int = 5) -> str:
    """Cut a draft at the point it degenerates into repetition.

    With max_tokens raised to 3500 and frequency_penalty held low, an exhausted
    model pads to the cap with synonym chains — a live proposal ended a section
    with "shattering shattering shattering ... vow-breakingly". Sentence-level
    repair is not possible; the only safe action is to truncate at the onset and
    keep the good prose before it.
    """
    if not text:
        return text
    words = text.split()
    run_start, run_len = 0, 1
    for i in range(1, len(words)):
        prev = words[i - 1].strip(".,;:!?").lower()
        cur = words[i].strip(".,;:!?").lower()
        if cur and cur == prev:
            run_len += 1
            if run_len >= run_threshold:
                cut = " ".join(words[:run_start]).rstrip()
                # Fall back to a sentence boundary so we do not end mid-clause.
                m = list(re.finditer(r"[.!?](\s|$)", cut))
                if m:
                    cut = cut[: m[-1].end()].rstrip()
                log.warning("degenerate repetition detected (%r x%d); truncating draft",
                            words[i], run_len)
                return cut
        else:
            run_len, run_start = 1, i
    return text


def _is_blank(text: Optional[str]) -> bool:
    """A drafted string is unusable if it is None or whitespace-only."""
    return text is None or not str(text).strip()


async def _draft_with_retry(
    client: httpx.AsyncClient,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int,
) -> Optional[str]:
    """Draft one (sub)section; retry ONCE if the model returns null/empty.

    Some models occasionally return a null (None) or whitespace-only message
    content on a successful HTTP response. That is not an exception, so it slips
    past ``draft_with_openrouter``'s error handling and used to crash the caller
    on ``.strip()`` — a soft-fail that silently dropped the subsection. Here we
    treat a blank body as a soft failure and retry once (which re-runs the
    existing primary->fallback path). Returns the stripped text, or None if it is
    still blank after the retry. Network/HTTP errors propagate to the caller.
    """
    content = await draft_with_openrouter(client, system_prompt, user_prompt, max_tokens=max_tokens)
    if _is_blank(content):
        log.warning("draft returned null/empty content; retrying once")
        content = await draft_with_openrouter(client, system_prompt, user_prompt, max_tokens=max_tokens)
    if _is_blank(content):
        return None

    # Single chokepoint for every drafted (sub)section.
    cleaned = strip_model_artifacts(content)
    cleaned, issues = document_qa.qa_text(cleaned)

    # A degenerate draft is padding, not content. Re-draft once — truncating
    # silently (the previous behaviour) hid the failure and lost real prose.
    if any(i.startswith("degenerate") for i in issues):
        log.warning("QA: degenerate draft (%s); re-drafting once",
                    "; ".join(i for i in issues if i.startswith("degenerate")))
        retry = await draft_with_openrouter(
            client, system_prompt,
            user_prompt + (
                "\n\nIMPORTANT: your previous attempt padded the section by repeating "
                "the same phrases. Write it again, shorter, with no repetition. Stop as "
                "soon as the substance is covered — length is not a goal."
            ),
            max_tokens=max_tokens)
        if not _is_blank(retry):
            retry_clean, retry_issues = document_qa.qa_text(
                strip_model_artifacts(retry))
            if not any(i.startswith("degenerate") for i in retry_issues):
                cleaned = retry_clean
            else:
                # Both attempts padded: keep the good prose, drop the padding.
                log.warning("QA: re-draft still degenerate; truncating at onset")
                cleaned = document_qa.truncate_at_degeneration(retry_clean)
        else:
            cleaned = document_qa.truncate_at_degeneration(cleaned)
    elif issues:
        log.info("QA: %s", "; ".join(issues))

    return cleaned if not _is_blank(cleaned) else None


def _assumption_placeholder(
    context: dict,
    section_title: str,
    sub_title: Optional[str] = None,
    facet: Optional[str] = None,
) -> str:
    """Grounded placeholder for a (sub)section the model could not draft.

    Derived ONLY from intake/context (client, proposal type, vendor, the
    section/subsection this call was meant to cover). Deliberately generic and
    marked ``[ASSUMPTION]`` + ``[SME REVIEW]`` — it invents no client facts,
    metrics, versions or commitments, only states that an SME must author it.
    """
    client_name = context.get("client_name") or "the client"
    ptype = context.get("proposal_type") or "implementation"
    focus = sub_title or section_title
    scope = f", which should cover {facet}," if facet else ""
    return (
        f"{ASSUMPTION_MARKER} The \"{focus}\" content{scope} could not be automatically "
        f"drafted from the available evidence for this {ptype} engagement for "
        f"{client_name}{_vendor_clause(context.get('iam_vendor'))}. This is a placeholder "
        f"assumption to be validated and replaced with grounded detail by a subject-matter "
        f"expert before any client use — no specifics have been inferred. {SME_REVIEW_MARKER}"
    )


# Prose subsections get a much tighter budget than table subsections.
#
# Measured: run 5 produced 16,051 prose words against the human proposal's 6,648
# across a near-identical subsection count (54 vs 53). That is ~300 words per
# subsection against IV's ~125. The cause is not the model choosing to ramble:
# a section was measured running dry at ~950 tokens and continuing to ~2,000, so
# it fills whatever budget it is given. The budget is therefore the lever.
#
# Table subsections keep the full budget - a 25-row RACI needs it.
# 420 tokens for a 220-word instruction left no headroom: dense technical prose
# tokenizes above 1.35x per word, so the model wrote to its word target, began
# an [SME REVIEW] marker and was guillotined mid-word. Run 8 shipped "[SME REV",
# "[SME RE" and "[SME REVIEW: specific SoD rule s".
#
# A hard cap should be a SAFETY NET against runaway generation, never the
# binding constraint on normal output. The word instruction does the shaping.
PROSE_SUBSECTION_TOKENS = int(os.environ.get("SHILPI_PROSE_SUBSECTION_TOKENS", "900"))
PROSE_SUBSECTION_WORDS = int(os.environ.get("SHILPI_PROSE_SUBSECTION_WORDS", "220"))
# IV's own proposals have a MEDIAN body paragraph of 29 words and a longest of
# 108. Run 8's median was 55 with twenty paragraphs over 100 and a longest of
# 180 -- because the instruction capped the SUBSECTION and said nothing about
# paragraphs, so the model wrote one 180-word block instead of four short ones.
PROSE_PARAGRAPH_WORDS = int(os.environ.get("SHILPI_PROSE_PARAGRAPH_WORDS", "60"))
_WANTS_TABLE_RE = re.compile(r"\bmarkdown\s+TABLE\b|\bas a (?:markdown )?TABLE\b", re.I)


async def draft_section(
    client: httpx.AsyncClient,
    section_spec: SectionSpec,
    context: dict,
    *,
    embed_fn: EmbedFn,
    retrieve_fn: RetrieveFn,
    build_grounded_system_fn: BuildSystemFn,
    top_k: int = TOP_K,
    fanout: int = 1,
    subsections: int = 1,
    max_tokens: int = MAX_DRAFT_TOKENS,
) -> dict:
    """Draft one proposal section, grounded in retrieved corpus evidence.

    Depth controls (Pass 3):
      fanout      : number of retrieval queries merged for this section's evidence.
      subsections : number of INDEPENDENT drafting calls. >1 splits the section
                    into focused facets (Overview / Detailed Design / ...), each
                    its own LLM call with the per-call ``max_tokens`` budget — so
                    depth grows via more calls, never a bigger single call.

    Returns: {"id","title","content","subsections","citations","max_similarity",
              "needs_sme_review"}.
    """
    chunks = await _retrieve_fanout(
        client, section_spec, context,
        embed_fn=embed_fn, retrieve_fn=retrieve_fn, top_k=top_k, fanout=fanout,
    )

    max_similarity = max((float(c.get("similarity") or 0.0) for c in chunks), default=0.0)
    needs_sme_review = (not chunks) or (max_similarity < WEAK_EVIDENCE_THRESHOLD)

    # Reuse the brain's evidence/system-prompt builder, then layer section-specific
    # drafting instructions on top so the model drafts THIS section.
    section_title = section_spec.render_title(context)
    evidence_block = build_grounded_system_fn(chunks)
    system_prompt = _SECTION_SYSTEM_TEMPLATE.format(
        title=section_title,
        proposal_type=context.get("proposal_type", "implementation"),
        client_name=context.get("client_name", "the client"),
        vendor_clause=_vendor_clause(context.get("iam_vendor")),
        purpose=section_spec.purpose,
        marker=SME_REVIEW_MARKER,
        evidence=evidence_block,
    )
    rfp_ctx = (context.get("rfp_text") or "")[:4000]
    # The discovery answers relevant to THIS section. Without this the drafting
    # engine saw only rfp_text and invented or omitted every captured specific.
    discovery_ctx = discovery_context_for(section_spec.id, context.get("discovery_answers"))
    client_facts = (f"CLIENT-SUPPLIED FACTS FOR THIS SECTION — these are "
                    f"authoritative and MUST be used verbatim where relevant. Do not "
                    f"replace them with generic statements, and do not claim they are "
                    f"unavailable:\n{discovery_ctx}\n\n") if discovery_ctx else ""

    # Plan the subsection facets.
    #
    # A section that defines its OWN subsections uses all of them, ignoring the
    # depth tier's subsection count: those headings are the section's structure,
    # not a knob. "Proposed Production Hardware Sizing" is not an optional
    # elaboration of "Proposed Solution", it is part of what that section IS.
    #
    # Only sections without their own subsections fall back to the generic
    # Overview / Detailed Design / Considerations triple, which is what every
    # section used to get and why the table of contents repeated three headings
    # seven times.
    own = section_spec.render_subsections(context)
    if own:
        facets = own
    else:
        n_sub = max(1, min(int(subsections), len(SUBSECTION_FACETS)))
        facets = SUBSECTION_FACETS[:n_sub] if n_sub > 1 else []

    async def _draft_once(user_prompt: str) -> Optional[str]:
        return await _draft_with_retry(client, system_prompt, user_prompt, max_tokens)

    subsection_results: list[dict] = []
    drafting_failed = False
    if not facets:
        user_prompt = (
            f"Draft the \"{section_title}\" section now, grounded in the EVIDENCE and citing inline as [N].\n\n"
            f"{client_facts}"
            f"RFP / requirement context:\n{rfp_ctx}"
        )
        try:
            content = await _draft_once(user_prompt)
        except Exception as e:
            log.error("draft_section drafting failed for %s: %s", section_spec.id, e)
            content = (
                f"{SME_REVIEW_MARKER}: drafting failed for this section ({e}). "
                "A subject-matter expert must author it manually."
            )
            drafting_failed = True
        # Null/empty draft (even after the retry inside _draft_once) must not
        # drop the section: emit a grounded placeholder instead.
        if _is_blank(content):
            log.warning("draft_section produced empty content for %s; using grounded placeholder",
                        section_spec.id)
            content = _assumption_placeholder(context, section_title)
            needs_sme_review = True
    else:
        # Independent drafting call per facet (the structured fan-out).
        parts: list[str] = []
        for sub_title, facet in facets:
            wants_table = _WANTS_TABLE_RE.search(facet) is not None
            budget = max_tokens if wants_table else min(max_tokens, PROSE_SUBSECTION_TOKENS)
            length_rule = (
                "Output the table and one short lead-in line, nothing else."
                if wants_table else
                f"Write NO MORE than {PROSE_SUBSECTION_WORDS} words in total, and "
                f"keep every paragraph under {PROSE_PARAGRAPH_WORDS} words - break "
                f"longer thoughts into separate short paragraphs. IV's proposals have "
                f"a median paragraph of 29 words; short paragraphs are the house style, "
                f"not a constraint. Density beats length. If you run out of grounded "
                f"material, stop - do not pad."
            )
            user_prompt = (
                f"Draft the \"{sub_title}\" subsection of the \"{section_title}\" section, "
                f"focusing specifically on {facet}. Ground every claim in the EVIDENCE and cite "
                f"inline as [N]. Do not repeat content that belongs in other subsections.\n"
                f"{length_rule}\n\n"
                f"{client_facts}"
                f"RFP / requirement context:\n{rfp_ctx}"
            )
            try:
                sub_content = await _draft_with_retry(
                    client, system_prompt, user_prompt, budget)
            except Exception as e:
                log.error("draft_section subsection %s failed for %s: %s", sub_title, section_spec.id, e)
                sub_content = (
                    f"{SME_REVIEW_MARKER}: drafting failed for this subsection ({e}). "
                    "A subject-matter expert must author it manually."
                )
                drafting_failed = True
            # Null/empty subsection (even after the retry inside _draft_once) is
            # the soft-fail that used to lose ~2/3 of full-depth content: emit a
            # grounded placeholder rather than an empty string.
            if _is_blank(sub_content):
                log.warning("draft_section subsection %s empty for %s; using grounded placeholder",
                            sub_title, section_spec.id)
                sub_content = _assumption_placeholder(context, section_title, sub_title, facet)
                needs_sme_review = True
            subsection_results.append({"title": sub_title, "content": sub_content})
            # An untitled subsection is continuous prose under the section
            # heading (IV's executive summary), so no "### " marker.
            parts.append(f"### {sub_title}\n\n{sub_content}"
                         if sub_title.strip() else sub_content)
        content = "\n\n".join(parts).strip()

    if drafting_failed:
        needs_sme_review = True

    content = (content or "").strip()
    if needs_sme_review and not content.startswith("[SME REVIEW"):
        content = (
            f"[SME REVIEW: weak evidence] Retrieval found "
            f"{'no' if not chunks else 'only low-similarity'} supporting evidence "
            f"(max similarity {max_similarity:.2f}). Verify and expand with an SME.\n\n"
            + content
        )

    return {
        "id": section_spec.id,
        "title": section_title,
        "content": content,
        "subsections": subsection_results,
        "citations": chunks,
        "max_similarity": max_similarity,
        "needs_sme_review": needs_sme_review,
    }


# ---------------------------------------------------------------------------
# DOCX assembly
# ---------------------------------------------------------------------------

_DRAFT_COLOR = RGBColor(0xB0, 0x00, 0x00)  # warning red (SME-review flag)
_CITATION_RE = re.compile(r"\[(\d+)\]")



def _add_static_toc(document: Document, sections: list[dict],
                    extra_titles: Optional[list[str]] = None) -> None:
    """Write a formatted contents table from the headings we already know.

    Renders as a real table (section | subsections) rather than a flat run of
    Normal paragraphs, which looked unformatted next to IV's own documents. A
    bare Word TOC field is not used because it shows "Right-click here and
    choose 'Update Field'..." until someone presses F9.
    """
    rows = []
    for sec in sections:
        title = (sec.get("title") or "").strip()
        if not title:
            continue
        subs = [ (sub.get("title") or "").strip()
                 for sub in (sec.get("subsections") or []) ]
        rows.append((title, [x for x in subs if x]))
    for title in extra_titles or []:
        rows.append((title, []))
    if not rows:
        return

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for idx, (title, subs) in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].width = Inches(0.45)
        cells[1].width = Inches(5.55)

        num = cells[0].paragraphs[0]
        num.paragraph_format.space_after = Pt(2)
        nrun = num.add_run(str(idx))
        nrun.bold = True
        nrun.font.color.rgb = branding.ORANGE

        body = cells[1].paragraphs[0]
        body.paragraph_format.space_after = Pt(2)
        body.add_run(title).bold = True
        for sub in subs:
            sp = cells[1].add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.18)
            sp.paragraph_format.space_after = Pt(0)
            r = sp.add_run(sub)
            r.font.size = Pt(9)
            r.font.color.rgb = branding.NEUTRAL_MUTED


def _add_toc_field(document: Document) -> None:
    """Insert a real, F9-refreshable Word Table of Contents field.

    Until the user updates the field in Word (right-click -> Update Field, or
    F9), Word shows the instructional placeholder text below.
    """
    para = document.add_paragraph()
    run = para.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Right-click here and choose 'Update Field' (or press F9) to build the "
        "table of contents from the section headings below."
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


# Inline emphasis the drafting model emits. ONLY asterisks are treated as
# markup: underscores are far too common in this domain's technical identifiers
# (data_retention, session_data, iam_vendor) and an underscore rule silently
# italicises across them.
_INLINE_MD_RE = re.compile(r"(\*\*.+?\*\*|\*(?!\s)[^*\n]+?(?<!\s)\*)", re.S)

# A bullet marker is "-", "•", or a SINGLE "*" followed by space. "**" is bold.
_BULLET_MARKER_RE = re.compile(r"^\s*(?:[-\u2022]|\*(?!\*))\s+")


def _add_inline_runs(para, text: str) -> None:
    """Add ``text`` to ``para`` converting **bold** / *italic* into real runs.

    Without this the model's markdown lands in the DOCX as literal asterisks —
    a generated proposal contained 147 stray ``**`` and 82 ``*`` sequences,
    which is immediately visible to anyone opening the document.
    """
    for token in _INLINE_MD_RE.split(text):
        if not token:
            continue
        if len(token) > 4 and token.startswith("**") and token.endswith("**"):
            para.add_run(token[2:-2]).bold = True
        elif len(token) > 2 and token.startswith("*") and token.endswith("*"):
            para.add_run(token[1:-1]).italic = True
        else:
            para.add_run(token)


def _is_gfm_separator(line: str) -> bool:
    """True for a GFM header/body separator such as ``|---|---|`` or ``---|---``.

    Deliberately stricter than a bare subset check: an EMPTY line is a subset of
    the allowed character set, which previously let a single stray ``|`` line be
    parsed as a one-row table.
    """
    s = line.strip()
    return bool(s) and "-" in s and set(s) <= set("|-: ")


def _parse_gfm_table(lines: list[str], i: int) -> Optional[tuple[list[str], list[list[str]], int]]:
    """If a GFM table starts at ``lines[i]``, return (headers, rows, next_index).

    Returns None when there is no table here, so callers fall through to their
    normal prose handling.
    """
    if i >= len(lines) or not lines[i].strip().startswith("|"):
        return None
    if i + 1 >= len(lines) or not _is_gfm_separator(lines[i + 1]):
        return None
    headers = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    j = i + 2
    rows: list[list[str]] = []
    while j < len(lines) and lines[j].strip().startswith("|"):
        rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
        j += 1
    return headers, rows, j


def _add_gfm_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Render parsed GFM table data as a native Word table.

    Cell text goes through ``_add_inline_runs`` so ``**Total**`` inside a BOQ or
    milestone row becomes a bold run rather than literal asterisks — the same
    failure that put 147 stray ``**`` into a generated proposal.
    """
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell_para = table.rows[0].cells[j].paragraphs[0]
        _add_inline_runs(cell_para, h)
        for run in cell_para.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for j in range(len(headers)):
            _add_inline_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")


def _add_body_paragraphs(document: Document, text: str) -> None:
    """Add body text, preserving [N] citation markers, splitting on blank lines.

    Markdown tables emitted inside a drafted section are rendered as native Word
    tables. Before this, only the compliance-matrix path could do that, so a
    sizing or BOQ table anywhere else in the document landed as a single
    paragraph of literal pipe characters.
    """
    if "|" in text:
        _add_body_with_tables(document, text)
        return
    _add_prose_paragraphs(document, text)


def _add_body_with_tables(document: Document, text: str) -> None:
    """Split body text into alternating prose runs and GFM tables, in order."""
    lines = text.splitlines()
    i = 0
    prose: list[str] = []

    def _flush() -> None:
        if prose:
            _add_prose_paragraphs(document, "\n".join(prose))
            prose.clear()

    while i < len(lines):
        parsed = _parse_gfm_table(lines, i)
        if parsed is None:
            prose.append(lines[i])
            i += 1
            continue
        headers, rows, i = parsed
        _flush()
        _add_gfm_table(document, headers, rows)
    _flush()


def _add_prose_paragraphs(document: Document, text: str) -> None:
    """Render body text into real Word paragraph styles.

    This used to decide per BLOCK: if a block began with "- " the whole block
    became bullets, otherwise the whole block became one paragraph. Drafted text
    is not shaped that way. A block like

        Assumptions & Open Questions:
        - The deployment model is not specified
        - Volumetrics are not provided

    begins with a label, so the label AND both dashes landed in a single
    paragraph as literal text. Measured on Amlak run 4: 42 paragraphs carried
    embedded newlines and 52 dash-lines rendered as body text, with List Bullet
    used zero times in the entire document.

    The decision is now per LINE, so each line gets the style it deserves.
    """
    for block in re.split(r"\n\s*\n", (text or "").strip()):
        block = block.strip()
        if not block:
            continue
        pending: list[str] = []

        def _flush() -> None:
            if not pending:
                return
            para = document.add_paragraph()
            _add_inline_runs(para, " ".join(pending))
            pending.clear()

        for raw in block.splitlines():
            line = raw.strip()
            if not line:
                _flush()
                continue

            heading = re.match(r"^(#{1,6})\s+(.*)$", line)
            if heading:
                _flush()
                _add_subheading(document, heading.group(2).strip(),
                                level=min(3, len(heading.group(1)) + 1))
                continue

            if _BULLET_MARKER_RE.match(line):
                _flush()
                # Strip ONLY the bullet marker. A blunt lstrip("-*• ") also ate
                # the opening ** of "- **Label:** text", leaving an unmatched
                # closing marker and losing the bold entirely.
                body = _BULLET_MARKER_RE.sub("", line).strip()
                if body:
                    _add_inline_runs(document.add_paragraph(style="List Bullet"), body)
                continue

            numbered = re.match(r"^\d+[.)]\s+(.*)$", line)
            if numbered and len(line.split()) <= _MAX_LIST_LINE_WORDS:
                _flush()
                _add_inline_runs(document.add_paragraph(style="List Number"),
                                 numbered.group(1).strip())
                continue

            if _is_pseudo_heading(line):
                _flush()
                _add_subheading(document, line.rstrip(":").strip(), level=3)
                continue

            pending.append(line)
        _flush()


# A numbered line longer than this is prose that happens to open with "1.", not
# a list item. Amlak run 3 contained a 783-word "list item".
_MAX_LIST_LINE_WORDS = 40

# A short line ending in a colon that introduces what follows. The model writes
# these as plain text ("Risks to Manage:", "Assumptions & Open Questions:") and
# they must become headings, not body prose. Bounded by length so a real
# sentence containing a colon is not promoted.
_PSEUDO_HEADING_RE = re.compile(r"^[A-Z][^.!?]{2,70}:$")


def _is_pseudo_heading(line: str) -> bool:
    return bool(_PSEUDO_HEADING_RE.match(line)) and len(line.split()) <= 10


def _add_subheading(document: Document, text: str, *, level: int = 3) -> None:
    """A bold sub-heading that is a real Word heading, so it reaches the TOC."""
    para = document.add_heading("", level=level)
    _add_inline_runs(para, text)
    for run in para.runs:
        run.bold = True


def _add_picture_fitted(document: Document, stream, *, max_w, max_h) -> None:
    """Insert an image scaled to fit inside (max_w, max_h), keeping aspect ratio."""
    try:
        from PIL import Image as _PILImage
        pos = stream.tell() if hasattr(stream, "tell") else None
        with _PILImage.open(stream) as im:
            w_px, h_px = im.size
        if pos is not None:
            stream.seek(pos)
        if w_px and h_px:
            scale = min(max_w / w_px, max_h / h_px)
            document.add_picture(stream, width=int(w_px * scale), height=int(h_px * scale))
            return
    except Exception as e:  # noqa: BLE001 — never lose a diagram over sizing
        log.warning("could not measure diagram for fitting (%s); using width only", e)
        if hasattr(stream, "seek"):
            stream.seek(0)
    document.add_picture(stream, width=max_w)


# Reusable assets are supporting material, not the point of the page: sized
# smaller than a generated architecture diagram so they illustrate rather than
# dominate. The text column is 5.9in wide with IV's margins.
# Diagrams are the point of the page; reusable assets illustrate. Different caps.
_DIAGRAM_MAX_W = Inches(6.0)
_DIAGRAM_MAX_H = Inches(7.5)

_IMAGE_MAX_W = Inches(4.5)
_IMAGE_MAX_H = Inches(3.2)


def _embeddable_diagrams(diagrams: Optional[list[dict]]) -> list[dict]:
    """Diagrams that will actually be embedded: approved AND carrying an image.

    Single source of truth. The contents list previously re-implemented a looser
    version of this check and advertised a diagram section the assembler then
    declined to create.
    """
    out: list[dict] = []
    for d in diagrams or []:
        if (d.get("status") or "").strip() != "approved":
            continue
        title = d.get("title") or "Architecture Diagram"
        image = d.get("image_bytes")
        path = d.get("image_path")
        if image:
            out.append({"title": title, "stream": io.BytesIO(image)})
        elif path and os.path.exists(path):
            out.append({"title": title, "stream": path})
    return out


def _add_approved_diagrams(document: Document, diagrams: Optional[list[dict]]) -> None:
    """Embed ONLY approved, rendered architecture diagrams as images.

    A diagram is embedded iff status == 'approved' and it carries usable image
    bytes (``image_bytes``) or a readable ``image_path``. Everything else
    (draft/needs_review/rejected, or no render) is skipped — this is the DOCX-side
    enforcement of the approval gate. If nothing qualifies, no section is added,
    so proposals without approved diagrams are unchanged.
    """
    embeddable = _embeddable_diagrams(diagrams)
    if not embeddable:
        return

    document.add_page_break()
    branding.add_section_heading(document, "Solution Architecture Diagrams")
    intro = document.add_paragraph()
    irun = intro.add_run(
        "The following architecture diagrams have been reviewed and approved for inclusion. "
        "They are draft artefacts for human review and not client-ready commitments."
    )
    irun.italic = True
    irun.font.size = Pt(10)
    for item in embeddable:
        document.add_heading(item["title"], level=2)
        try:
            # Width alone lets a tall D2 render scale to any height — a live
            # proposal had architecture diagrams spanning 2+ pages. Constrain by
            # BOTH dimensions, preserving aspect ratio, so a diagram always fits
            # on one page.
            _add_picture_fitted(document, item["stream"],
                                max_w=Inches(6.0), max_h=Inches(7.5))
        except Exception as e:  # noqa: BLE001 — a bad image must not break the doc
            log.error("failed to embed diagram '%s' (skipping): %s", item["title"], e)


async def _attach_assets(client, sections: list[dict], context: dict,
                         asset_fns: dict) -> None:
    """Put approved, section-appropriate images on each drafted section.

    `asset_fns` supplies {"library": async () -> list[dict],
                          "download": async (storage_path) -> stream|None}.

    Every image that reaches a document has been approved by a human, matched
    to the section by its vision description, and checked against the vendor
    being proposed. `architecture` assets are excluded upstream because they
    depict a single client's estate.
    """
    library = await asset_fns["library"](client)
    if not library:
        return

    vendor = context.get("iam_vendor")
    used: set[str] = set()
    placed = 0
    for sec in sections:
        chosen = asset_selection.select_assets(
            library, sec.get("id") or "", vendor, limit=_ASSETS_PER_SECTION)
        attached = []
        for a in chosen:
            # One image appears once per document, however many sections it
            # suits -- the same picture twice reads as a mistake.
            if a["storage_path"] in used:
                continue
            stream = await asset_fns["download"](client, a["storage_path"])
            if stream is None:
                continue
            used.add(a["storage_path"])
            attached.append({"id": a["id"], "stream": stream})
        if attached:
            sec["assets"] = attached
            placed += len(attached)
    log.info("assets: %d image(s) placed across %d section(s)",
             placed, sum(1 for s in sections if s.get("assets")))



# Which drafted subsection each diagram belongs beside.
#
# IV embeds every diagram INSIDE the subsection that explains it: the deployment
# architecture diagram sits under "Proposed Deployment Architecture", right after
# the prose describing it. Shilpi collected all diagrams into a trailing
# "Solution Architecture Diagrams" section, so in run 9 a reader had to hold the
# text in their head and go looking forty pages later.
#
# Matched on the diagram's TYPE and title against the subsection heading, since
# the diagram plan and the template are written independently.
_DIAGRAM_PLACEMENT: tuple[tuple[str, str], ...] = (
    (r"joiner|integration|hrms|lifecycle|jml", r"HRMS Integration|Joiner"),
    (r"deployment", r"Proposed Deployment Architecture"),
    (r"security|network", r"Proposed Production Architecture"),
    (r"solution|reference|architecture", r"Proposed Future IAM State"),
    (r"migration|cutover", r"Migration Pattern|Proposed Target Architecture"),
)

_DIAGRAM_PLACEMENT_RE = tuple(
    (re.compile(d, re.I), re.compile(h, re.I)) for d, h in _DIAGRAM_PLACEMENT)


def _placement_for(diagram: dict) -> Optional[re.Pattern]:
    """The subsection-heading pattern this diagram should sit under, if any."""
    key = f"{diagram.get('diagram_type') or ''} {diagram.get('title') or ''}"
    for dpat, hpat in _DIAGRAM_PLACEMENT_RE:
        if dpat.search(key):
            return hpat
    return None


def _claim_diagram(embeddable: list[dict], heading: str,
                   used: set) -> Optional[dict]:
    """The diagram belonging under this subsection heading, once."""
    for item in embeddable:
        if id(item) in used:
            continue
        pattern = _placement_for(item)
        if pattern and pattern.search(heading or ""):
            used.add(id(item))
            return item
    return None


def _embed_diagram(document: Document, item: dict) -> None:
    try:
        _add_picture_fitted(document, item["stream"],
                            max_w=_DIAGRAM_MAX_W, max_h=_DIAGRAM_MAX_H)
    except Exception as e:  # noqa: BLE001 - a bad render must not sink the section
        log.warning("could not embed diagram %s: %s", item.get("title"), e)


def _render_section_assets(document: Document, sec: dict) -> None:
    """Embed the approved images chosen for this section.

    Only assets a human has approved reach this point, and only `corporate` and
    `product` kinds -- `architecture` assets depict a specific client's estate
    and are excluded in asset_selection.py.

    No captions: they were generated from the vision description and leaked
    another client's project name into run 8.
    """
    for asset in sec.get("assets") or []:
        stream = asset.get("stream")
        if not stream:
            continue
        try:
            _add_picture_fitted(document, stream,
                                max_w=_IMAGE_MAX_W, max_h=_IMAGE_MAX_H)
        except Exception as e:  # noqa: BLE001 - a bad image must not sink a section
            log.warning("could not embed asset %s: %s", asset.get("id"), e)


def assemble_docx(
    metadata: dict,
    sections: list[dict],
    compliance_markdown: Optional[str] = None,
    client_logo_path: Optional[str] = None,
    include_appendices: bool = False,
    diagrams: Optional[list[dict]] = None,
) -> bytes:
    """Build a professional, IV-branded Word document and return its bytes.

    metadata keys: client_name, proposal_type, iam_vendor (optional),
    generated_at (optional), version (optional).

    client_logo_path: optional path to a client logo image embedded in the
    title-page box. When None (default) a bordered "Client Logo" placeholder is
    drawn instead. Logos are never sourced online in this pass.

    include_appendices: when True (full depth) appends the appendix pack —
    RACI, timeline, sizing, integration inventory and risks — as real DOCX
    tables. Where intake/retrieval data is absent, conservative
    assumption-marked placeholder rows are used (never fabricated specifics).

    diagrams (Pass 4): optional list of architecture-diagram dicts. ONLY diagrams
    whose status == 'approved' AND that carry a rendered image (``image_bytes``
    or a readable ``image_path``) are embedded — draft/rejected/needs_review
    diagrams are silently skipped. When no diagram qualifies, the document is
    byte-for-byte the same as before (Pass 1-3 unchanged).
    """
    metadata = {
        **metadata,
        "generated_at": metadata.get("generated_at")
        or datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
    }

    document = Document()
    branding.configure_base_styles(document)
    branding.apply_header_footer(document, metadata.get("client_name") or "Client",
                                 client_logo_path=client_logo_path)

    # --- Title page (IV branding) ------------------------------------------
    branding.add_title_page(document, metadata, client_logo_path=client_logo_path)
    document.add_page_break()

    # --- Table of contents (real, refreshable Word TOC field) --------------
    # NOT a section heading: a contents page that lists itself as section 1 is
    # a Shilpi artefact, not something IV's proposals do. Styled to match, but
    # outside the heading hierarchy so it stays out of the contents it renders.
    _toc_title = document.add_paragraph()
    _toc_run = _toc_title.add_run("Table of Contents")
    _toc_run.bold = True
    _toc_run.font.size = Pt(16)
    _toc_run.font.color.rgb = branding.NAVY
    _toc_extra = [t for t in (
        "Solution Architecture Diagrams" if _embeddable_diagrams(diagrams) else None,
        "Compliance Matrix" if any(
            s.get("id") == COMPLIANCE_SECTION_ID for s in sections) else None,
    ) if t]
    _add_static_toc(document, sections, _toc_extra)
    document.add_page_break()

    # --- Sections -----------------------------------------------------------
    # Diagrams are placed INSIDE the subsection that explains them, the way IV
    # does it. Whatever no subsection claims still gets the trailing gallery, so
    # an unmatched diagram is never silently dropped.
    _inline_diagrams = _embeddable_diagrams(diagrams)
    _placed_diagrams: set = set()

    aggregated_assumptions: list[str] = []
    for sec in sections:
        heading = branding.add_section_heading(document, sec.get("title", "Untitled"))
        if sec.get("needs_sme_review"):
            flag = heading.add_run("   [SME REVIEW REQUIRED]")
            flag.font.size = Pt(10)
            flag.font.color.rgb = _DRAFT_COLOR
        # Reusable images sit directly under the section heading they were
        # chosen for. They used to render AFTER the whole section body, so an
        # image appeared beneath whatever the last subsection heading happened
        # to be -- run 9 put a CIAM governance pyramid under "Assumptions &
        # Open Questions" and a change-control graphic under "Training and
        # Post-Production Support". The selection was right; the position made
        # it read as wrong.
        _render_section_assets(document, sec)

        subs = sec.get("subsections") or []
        if subs:
            # Multi-subsection (full depth): render each facet under an H2 heading.
            for sub in subs:
                # An EMPTY title means "continuous prose under the section
                # heading", which is how IV writes its executive summary. Adding
                # an empty H2 would put a blank line in the table of contents.
                title = (sub.get("title") or "").strip()
                if title:
                    document.add_heading(title, level=2)
                _add_body_paragraphs(document, sub.get("content", ""))
                # The diagram that explains THIS subsection, immediately after
                # the prose describing it -- the way IV places them.
                claimed = _claim_diagram(_inline_diagrams, title, _placed_diagrams)
                if claimed:
                    _embed_diagram(document, claimed)
        else:
            _add_body_paragraphs(document, sec.get("content", ""))

        # Assets are rendered directly under the SECTION heading (above), not
        # here. See _render_section_assets.

        # Opportunistically collect assumption-ish lines for the aggregate section.
        if "assumption" not in (sec.get("id") or "").lower():
            for line in (sec.get("content") or "").splitlines():
                s = line.strip().lstrip("-*• ").strip()
                if s and re.match(r"(?i)^(assumption|assume|open question|dependency)", s):
                    aggregated_assumptions.append(s)

    # --- Assumptions & Open Questions (ensure one always exists) ------------
    has_dedicated_assumptions = any("assumption" in (s.get("id") or "").lower() for s in sections)
    if not has_dedicated_assumptions:
        branding.add_section_heading(document, "Assumptions & Open Questions")
        if aggregated_assumptions:
            for a in aggregated_assumptions:
                document.add_paragraph(a, style="List Bullet")
        else:
            document.add_paragraph(
                "No explicit assumptions were captured during drafting. "
                "An SME should confirm scope, dependencies, and open questions before client use.",
                style="List Bullet",
            )

    # --- Architecture Diagrams (Pass 4 — approved only) --------------------
    _add_approved_diagrams(document, diagrams)

    # --- Compliance Matrix (optional) --------------------------------------
    if compliance_markdown:
        document.add_page_break()
        # H2: a requirement-coverage matrix is a supporting artefact, not one
        # of the proposal's top-level sections.
        document.add_heading("Compliance Matrix", level=2)
        _add_markdown_ish(document, compliance_markdown)

    # --- Citation Appendix: REMOVED -----------------------------------------
    # This listed every retrieved corpus chunk by CLIENT NAME with a similarity
    # score. In the Amlak run 3 document it named Al Qadsiah Club 51 times, plus
    # Ministry of Energy and National Water Company: three other IV clients, in
    # a document addressed to a fourth. 78 paragraphs of it.
    #
    # The directive was already "citations are removed completely" — the inline
    # [N] markers were stripped but the appendix they pointed at was left
    # behind, which is the same built-but-never-finished shape as the rest.
    # Retrieval provenance belongs in logs, not in a client deliverable.

    # --- Appendices (full depth only) --------------------------------------
    # The appendix pack predates the template rebuild, when the body was seven
    # generic sections with no RACI, timeline, sizing or commercial content.
    # The body now carries all four as real tables, so run 5 printed each twice
    # with DIFFERENT numbers: a 25-row RACI built from discovery answers in the
    # body, and a 9-row generic placeholder one in Appendix A. Two versions of
    # the same table in one proposal is worse than neither, because a reader
    # cannot tell which is authoritative.
    if include_appendices:
        _add_appendices(document, metadata,
                        skip=_superseded_appendices({s.get("id") for s in sections}))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_markdown_ish(document: Document, text: str) -> None:
    """Render the compliance matrix markdown into DOCX in a robust, simple way.

    We don't build a full markdown parser: headings (#) become paragraphs, GFM
    tables become native Word tables, everything else becomes plain paragraphs.
    """
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # GFM table block — same parser/renderer the section body path uses, so
        # a fix to one is a fix to both.
        parsed = _parse_gfm_table(lines, i)
        if parsed is not None:
            headers, rows, i = parsed
            _add_gfm_table(document, headers, rows)
            continue
        if stripped.startswith("#"):
            para = document.add_paragraph()
            run = para.add_run(stripped.lstrip("# ").strip())
            run.bold = True
            i += 1
            continue
        if stripped.startswith(">"):
            note = document.add_paragraph()
            nr = note.add_run(stripped.lstrip("> ").strip())
            nr.italic = True
            i += 1
            continue
        document.add_paragraph(stripped)
        i += 1


# ---------------------------------------------------------------------------
# Appendices (Pass 3 — full depth). Deterministic, no LLM calls, no fabrication.
# Missing intake/retrieval data yields conservative assumption-marked rows.
# ---------------------------------------------------------------------------

_APPENDIX_ASSUMPTION = "[ASSUMPTION]"


def _appendix_note(document: Document, text: str) -> None:
    note = document.add_paragraph()
    run = note.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def _appendix_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Render a native Word table with a bold header row."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        run = table.rows[0].cells[j].paragraphs[0].add_run(h)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for j in range(len(headers)):
            cells[j].text = row[j] if j < len(row) else ""


# Appendix letter -> the body section that now carries the same content.
_APPENDIX_SUPERSEDED_BY = {
    "A": "implementation_approach",   # RACI matrix
    "B": "project_timeline",          # indicative timeline
    "C": "proposed_solution",         # sizing & volumetrics
    "F": "commercial",                # commercial structure
}


def _superseded_appendices(body_section_ids: set) -> frozenset:
    """Appendices the body template already covers, so they must not print twice.

    D (integration inventory) and E (risk register) have no body counterpart and
    are always kept.
    """
    return frozenset(letter for letter, sec_id in _APPENDIX_SUPERSEDED_BY.items()
                     if sec_id in body_section_ids)


def _add_appendices(document: Document, metadata: dict,
                    skip: frozenset = frozenset()) -> None:
    """Append the full-depth appendix pack as real DOCX sections/tables.

    Every row is either grounded in supplied metadata or an explicit
    ``[ASSUMPTION]`` placeholder — SMEs must confirm before client use. No
    specific figures, dates or SLAs are fabricated.
    """
    vendor = (metadata.get("iam_vendor") or "the selected IAM platform").strip() or "the selected IAM platform"
    ptype = (metadata.get("proposal_type") or "implementation").strip().lower()

    if len(skip) >= 6:
        return

    # Rather than reindenting six large literal blocks behind guards, the emit
    # helpers below no-op while the current appendix is in `skip`. Each block
    # sets `_cur` first. Fewer moving parts than restructuring, and the skip
    # decision stays in one place.
    _cur = [""]

    def _heading(text: str) -> None:
        # Level 2, not a section heading. "Appendices" is the H1; A through F sit
        # UNDER it. Emitting each as an H1 gave run 9 eighteen top-level
        # sections against IV's eleven, and a table of contents listing
        # "Appendices", "Appendix D", "Appendix E" and "Appendix F" as peers.
        if _cur[0] not in skip:
            document.add_heading(text, level=2)

    def _table(_doc, cols, rows) -> None:
        # Signature mirrors _appendix_table so the call sites are unchanged.
        if _cur[0] not in skip:
            _appendix_table(_doc, cols, rows)

    def _note(text: str) -> None:
        if _cur[0] not in skip:
            _appendix_note(document, text)

    document.add_page_break()
    branding.add_section_heading(document, "Appendices")
    _appendix_note(
        document,
        "The following appendices are structured planning artefacts. Rows marked "
        f"{_APPENDIX_ASSUMPTION} are conservative placeholders to be confirmed with the "
        "client and an SME during discovery — they are not commitments.",
    )

    # A. RACI matrix
    _cur[0] = "A"
    _heading("Appendix A — RACI Matrix")
    _table(
        document,
        ["Activity / Workstream", "InspiritVision", "Client", "Vendor"],
        [
            ["Solution design & architecture", "R/A", "C", "C"],
            ["Environment provisioning", "C", "R/A", "C"],
            ["Configuration & build", "R/A", "C", "I"],
            ["Integration & connector setup", "R/A", "C", "C"],
            ["Testing & UAT", "R", "A", "I"],
            ["Go-live & cutover", "R/A", "C", "I"],
            ["Knowledge transfer & handover", "R/A", "C", "I"],
            [f"{_APPENDIX_ASSUMPTION} Additional workstreams", "TBC", "TBC", "TBC"],
        ],
    )
    _note("R = Responsible, A = Accountable, C = Consulted, I = Informed.")

    # B. Timeline / phasing
    _cur[0] = "B"
    _heading("Appendix B — Indicative Timeline")
    _dur = " ".join(str((metadata.get("discovery_answers") or {}).get("engagement_duration") or "").split())
    if _dur and _dur.lower() not in ("skip", "none", "n/a", "-"):
        _note(f"Client-supplied engagement duration: {_dur}.")
    _table(
        document,
        ["Phase", "Key Activities", "Indicative Duration"],
        [
            ["Discovery & Design", "Requirements, current-state review, target design", f"{_APPENDIX_ASSUMPTION} TBC"],
            ["Build & Configure", "Platform config, connectors, workflows", f"{_APPENDIX_ASSUMPTION} TBC"],
            ["Test & Validate", "SIT, UAT, remediation", f"{_APPENDIX_ASSUMPTION} TBC"],
            ["Deploy & Stabilise", "Cutover, hypercare, handover", f"{_APPENDIX_ASSUMPTION} TBC"],
        ],
    )
    _note("Durations are confirmed once scope and volumetrics are baselined in discovery.")

    # C. Sizing
    _cur[0] = "C"
    _heading("Appendix C — Sizing & Volumetrics")
    # Use what discovery actually captured. These rows used to be hardcoded
    # "TBC / Client to confirm" even when the consultant had supplied exact
    # figures at intake — the same discard bug as the drafting path.
    _ans = metadata.get("discovery_answers") or {}

    def _row(label: str, *keys: str) -> list[str]:
        for k in keys:
            v = " ".join(str(_ans.get(k) or "").split())
            if v and v.lower() not in ("skip", "none", "n/a", "na", "-"):
                return [label, v, "Client-supplied at discovery"]
        return [label, f"{_APPENDIX_ASSUMPTION} TBC", "Client to confirm"]

    _rows = [
        _row("Identities / users", "user_count"),
        _row("Target applications", "app_count", "applications_to_onboard"),
        _row("Environments", "environments"),
        _row("Deployment model", "deployment_model"),
        _row("Hardware sizing", "hardware_sizing_inputs"),
        _row("Cluster topology", "cluster_topology"),
        _row("HA / DR", "ha_dr_requirements"),
        _row("Availability target", "availability"),
        _row("Peak transaction volume", "performance"),
    ]
    _appendix_table(document, ["Dimension", "Value", "Source"], _rows)

    # D. Integration inventory
    _cur[0] = "D"
    _heading("Appendix D — Integration Inventory")
    _table(
        document,
        ["System / Application", "Integration Type", f"{vendor} Connector", "Notes"],
        [
            ["Directory / HR source", "Authoritative source", f"{_APPENDIX_ASSUMPTION} TBC", "System of record for identities"],
            ["Core business applications", "Provisioning target", f"{_APPENDIX_ASSUMPTION} TBC", "Confirm inventory in discovery"],
            ["Downstream / custom apps", f"{_APPENDIX_ASSUMPTION} TBC", f"{_APPENDIX_ASSUMPTION} TBC", "May require custom connector"],
        ],
    )

    # E. Risks
    _cur[0] = "E"
    _heading("Appendix E — Risk Register")
    _table(
        document,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["Scope / requirements change", "Medium", "High", "Change control, phased delivery"],
            ["Application onboarding delays", "Medium", "Medium", "Early inventory, prioritised backlog"],
            ["Data quality in source systems", "Medium", "High", "Data-quality assessment in discovery"],
            ["Resource / SME availability", "Medium", "Medium", "Agreed RACI and governance cadence"],
            [f"{_APPENDIX_ASSUMPTION} {ptype}-specific risks", "TBC", "TBC", "To be assessed with client"],
        ],
    )

    # F. Commercial structure. The human proposals always carry one, and the
    _cur[0] = "F"
    # milestone basis IS captured at discovery. Figures stay with the human.
    _c = metadata.get("discovery_answers") or {}
    _crows = []
    for label, key in (("Licence included", "license_included"),
                       ("Pricing model", "pricing_model"),
                       ("Payment milestones", "payment_milestones"),
                       ("Taxes", "taxes"),
                       ("Travel", "travel"),
                       ("Support terms", "support_terms")):
        val = " ".join(str(_c.get(key) or "").split())
        if val and val.lower() not in ("skip", "none", "n/a", "na", "-"):
            _crows.append([label, val])
    if _crows:
        branding.add_section_heading(document, "Appendix F — Commercial Structure")
        _appendix_note(
            document,
            "Structure as captured at discovery. Pricing figures are deliberately "
            "omitted and must be completed by the commercial owner before issue.")
        _appendix_table(document, ["Item", "Basis"], _crows)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

async def generate_proposal(
    client: httpx.AsyncClient,
    rfp_text: str,
    client_name: str,
    proposal_type: str,
    iam_vendor: Optional[str] = None,
    *,
    embed_fn: EmbedFn,
    retrieve_fn: RetrieveFn,
    build_grounded_system_fn: BuildSystemFn,
    run_compliance_matrix_fn: Optional[Callable[..., Awaitable[object]]] = None,
    render_matrix_markdown_fn: Optional[Callable[[object], str]] = None,
    sections: Optional[list[str]] = None,
    include_compliance_matrix: bool = False,
    top_k: int = TOP_K,
    proposal_depth: Optional[str] = None,
    diagrams: Optional[list[dict]] = None,
    discovery_answers: Optional[dict] = None,
    client_logo_path: Optional[str] = None,
    asset_fns: Optional[dict] = None,
) -> dict:
    """Orchestrate: pick template, draft sections concurrently, assemble DOCX.

    The brain helpers are injected so this module never imports app.py.

    ``proposal_depth`` (brief|standard|full) controls long-form depth via
    STRUCTURED fan-out — retrieval queries + independent drafting calls per
    section, plus (full) an appendix pack. Unknown/absent values fall back to the
    safe ``standard`` tier, preserving existing behaviour.
    Returns {"docx_bytes", "sections_meta", "filename", ...}.
    """
    template = get_template(proposal_type)  # raises ValueError on bad type
    tier: DepthTier = get_depth_tier(proposal_depth)

    # Choose which sections to draft. Compliance is opt-in and handled separately.
    if sections:
        wanted = {s.strip().lower() for s in sections}
        chosen = [s for s in template if s.id in wanted or s.title.lower() in wanted]
        if not chosen:
            chosen = [s for s in template if not s.optional]
    else:
        chosen = [s for s in template if not s.optional]

    draft_specs = [s for s in chosen if s.id != COMPLIANCE_SECTION_ID]

    context = {
        "client_name": client_name,
        "iam_vendor": iam_vendor or "",
        "proposal_type": proposal_type,
        "rfp_text": rfp_text or "",
        "discovery_answers": discovery_answers or {},
    }

    sem = asyncio.Semaphore(DOC_CONCURRENCY)

    async def _draft(spec: SectionSpec) -> dict:
        async with sem:
            return await draft_section(
                client, spec, context,
                embed_fn=embed_fn,
                retrieve_fn=retrieve_fn,
                build_grounded_system_fn=build_grounded_system_fn,
                top_k=top_k,
                fanout=tier.retrieval_fanout,
                subsections=tier.subsections_per_section,
                max_tokens=tier.per_call_max_tokens,
            )

    drafted = await asyncio.gather(*[_draft(s) for s in draft_specs])

    # Optional compliance matrix.
    compliance_markdown: Optional[str] = None
    want_compliance = include_compliance_matrix or (
        sections is not None and COMPLIANCE_SECTION_ID in {s.strip().lower() for s in sections}
    )
    if want_compliance and run_compliance_matrix_fn and render_matrix_markdown_fn:
        try:
            matrix = await run_compliance_matrix_fn(client, rfp_text, None, top_k)
            compliance_markdown = render_matrix_markdown_fn(matrix)
        except Exception as e:
            log.error("Compliance matrix generation failed: %s", e)
            compliance_markdown = (
                f"{SME_REVIEW_MARKER}: compliance matrix generation failed ({e}). "
                "Run the /v1/compliance-matrix endpoint separately."
            )

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    metadata = {
        "client_name": client_name,
        "proposal_type": proposal_type,
        "iam_vendor": iam_vendor,
        "generated_at": generated_at,
        # Appendices read captured figures from here instead of printing "TBC"
        # over data the consultant already supplied.
        "discovery_answers": discovery_answers or {},
    }
    # Attach reusable images to each section before assembly. Selection is
    # pure (asset_selection.py); fetching bytes is the caller's job, passed in
    # as `asset_fns` so document_engine keeps no Supabase dependency -- the same
    # arrangement as retrieve_fn and embed_fn.
    if asset_fns:
        try:
            await _attach_assets(client, list(drafted), context, asset_fns)
        except Exception as e:  # noqa: BLE001
            # A proposal without images is a worse proposal, not a failed one.
            log.warning("asset attachment failed, drafting without images: %s", e)

    docx_bytes = assemble_docx(
        metadata, list(drafted), compliance_markdown,
        client_logo_path=client_logo_path,
        include_appendices=tier.include_appendices,
        diagrams=diagrams,
    )

    safe_client = re.sub(r"[^A-Za-z0-9]+", "_", client_name).strip("_") or "Client"
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"Shilpi_Proposal_{safe_client}_{stamp}.docx"

    sections_meta = [
        {
            "id": d["id"],
            "title": d["title"],
            "max_similarity": d["max_similarity"],
            "needs_sme_review": d["needs_sme_review"],
            "citation_count": len(d.get("citations") or []),
        }
        for d in drafted
    ]
    draft_markdown = "\n\n".join(
        f"## {d['title']}\n\n{d.get('content', '')}".rstrip() for d in drafted
    )

    return {
        "docx_bytes": docx_bytes,
        "sections_meta": sections_meta,
        "draft_markdown": draft_markdown,
        "filename": filename,
        "included_compliance_matrix": compliance_markdown is not None,
        "proposal_depth": tier.name,
        "included_appendices": tier.include_appendices,
    }
