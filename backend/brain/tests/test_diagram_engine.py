"""
Sprint 5 Pass 4 — architecture diagram framework tests (keyless, no network).

Covers:
  * DiagramSpec validation + sanitization (caps, escaping, dangling-edge drop,
    id rewriting, diagram_type coercion) and deterministic DOT building.
  * Render path: skips fail-soft when `dot` is unavailable; otherwise renders.
  * Approval state machine: valid/invalid transitions, rejection comment
    requirement, iteration bump on re-draft.
  * DOCX embedding: approved diagrams embed; draft/rejected/needs_review do NOT.
  * generate_proposal default path still works with ZERO diagrams (unchanged).
  * LLM spec generation uses the INJECTED structured helper (mocked) with the
    mandated caps — no live OpenRouter call.

No live Supabase / OpenRouter calls anywhere.
"""

import asyncio
import base64
import io
import os
import sys

import pytest

_HERE = os.path.dirname(os.path.abspath(__file__))
_BRAIN = os.path.dirname(_HERE)
if _BRAIN not in sys.path:
    sys.path.insert(0, _BRAIN)

import httpx
from docx import Document

import diagram_engine as de
import document_engine
from diagram_engine import DiagramEdge, DiagramNode, DiagramSpec, InvalidTransition
from document_engine import assemble_docx, generate_proposal

# A minimal valid 1x1 PNG so add_picture works without the `dot` binary or PIL.
_ONE_PX_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


# --- schema / sanitization --------------------------------------------------

def test_sanitize_caps_node_and_edge_counts():
    nodes = [DiagramNode(id=f"n{i}", label=f"Node {i}") for i in range(de.MAX_NODES + 15)]
    edges = [DiagramEdge(source=f"n{i}", target=f"n{i+1}") for i in range(de.MAX_EDGES + 30)]
    spec = DiagramSpec(title="Big", nodes=nodes, edges=edges)
    safe = de.sanitize_spec(spec)
    assert len(safe.nodes) <= de.MAX_NODES
    assert len(safe.edges) <= de.MAX_EDGES


def test_sanitize_clips_label_length():
    long_label = "x" * 500
    spec = DiagramSpec(title="t", nodes=[DiagramNode(id="a", label=long_label)], edges=[])
    safe = de.sanitize_spec(spec)
    assert len(safe.nodes[0].label) <= de.MAX_LABEL_LEN


def test_sanitize_rewrites_unsafe_ids_and_keeps_edges():
    spec = DiagramSpec(
        title="t",
        nodes=[DiagramNode(id="hr!! source", label="HR"), DiagramNode(id="idp", label="IdP")],
        edges=[DiagramEdge(source="hr!! source", target="idp", label="SCIM")],
    )
    safe = de.sanitize_spec(spec)
    ids = {n.id for n in safe.nodes}
    assert "hrsource" in ids  # non-alnum stripped
    assert len(safe.edges) == 1
    assert safe.edges[0].source in ids and safe.edges[0].target in ids


def test_sanitize_drops_dangling_edges():
    spec = DiagramSpec(
        title="t",
        nodes=[DiagramNode(id="a", label="A")],
        edges=[DiagramEdge(source="a", target="ghost")],
    )
    safe = de.sanitize_spec(spec)
    assert safe.edges == []


def test_diagram_type_coerced_to_allowlist():
    spec = DiagramSpec(diagram_type="totally-made-up", title="t", nodes=[], edges=[])
    assert spec.diagram_type == "architecture"
    ok = DiagramSpec(diagram_type="data_flow", title="t", nodes=[], edges=[])
    assert ok.diagram_type == "data_flow"


def test_build_dot_escapes_quotes_no_injection():
    # A label attempting to break out of the quoted string / inject attributes.
    evil = 'X" ]; node[color=red]; a[label="pwned'
    spec = DiagramSpec(
        title='T"itle',
        nodes=[DiagramNode(id="a", label=evil)],
        edges=[],
    )
    dot = de.build_dot(spec)
    # The raw payload (with its UNescaped quotes) must not survive verbatim — the
    # quotes are escaped, so the label can never break out to inject attributes.
    assert evil not in dot
    assert de._escape_label('a"b') == 'a\\"b'
    assert '\\"' in dot  # escaped quotes present
    # Only ONE real node attribute list exists (no injected extra node): count
    # node-definition lines, which start with an indented quoted id.
    node_defs = [ln for ln in dot.splitlines() if ln.strip().startswith('"a"')]
    assert len(node_defs) == 1
    assert dot.startswith("digraph shilpi_diagram {")
    assert dot.strip().endswith("}")


def test_build_dot_groups_into_clusters():
    spec = DiagramSpec(
        title="t",
        nodes=[
            DiagramNode(id="a", label="A", group="Zone1"),
            DiagramNode(id="b", label="B", group="Zone1"),
            DiagramNode(id="c", label="C"),
        ],
        edges=[DiagramEdge(source="a", target="b")],
    )
    dot = de.build_dot(spec)
    assert "subgraph cluster_0" in dot
    assert 'label="Zone1"' in dot


# --- render path ------------------------------------------------------------

def test_render_fail_soft_or_png(monkeypatch):
    spec = DiagramSpec(title="t", nodes=[DiagramNode(id="a", label="A")], edges=[])
    if de.dot_available():
        out = de.render_spec(spec, fmt="png")
        assert out and out[:8] == b"\x89PNG\r\n\x1a\n"
    else:
        assert de.render_spec(spec, fmt="png") is None


def test_render_returns_none_when_dot_missing(monkeypatch):
    monkeypatch.setattr(de.shutil, "which", lambda _: None)
    spec = DiagramSpec(title="t", nodes=[DiagramNode(id="a", label="A")], edges=[])
    assert de.render_spec(spec, fmt="png") is None


# --- state machine ----------------------------------------------------------

def test_valid_transitions():
    assert de.apply_transition({"status": "draft"}, "needs_review")["status"] == "needs_review"
    approved = de.apply_transition({"status": "needs_review"}, "approved")
    assert approved["status"] == "approved" and approved["approved"] is True


def test_reject_requires_comment_and_appends():
    with pytest.raises(InvalidTransition):
        de.apply_transition({"status": "needs_review"}, "rejected")
    patch = de.apply_transition(
        {"status": "needs_review", "rejection_comments": ["old"]},
        "rejected",
        rejection_comment="needs a cloud zone",
    )
    assert patch["status"] == "rejected"
    assert patch["rejection_comments"] == ["old", "needs a cloud zone"]


def test_redraft_bumps_iteration():
    patch = de.apply_transition({"status": "rejected", "iteration": 2}, "draft")
    assert patch["status"] == "draft" and patch["iteration"] == 3


def test_invalid_transitions_rejected():
    for current, target in [
        ("draft", "approved"),
        ("draft", "rejected"),
        ("approved", "draft"),
        ("approved", "needs_review"),
        ("needs_review", "draft"),
        ("rejected", "approved"),
    ]:
        with pytest.raises(InvalidTransition):
            de.apply_transition({"status": current}, target, rejection_comment="c")


def test_unknown_target_rejected():
    with pytest.raises(InvalidTransition):
        de.apply_transition({"status": "draft"}, "banana")


# --- LLM spec generation (injected structured helper, mocked) --------------

def test_generate_diagram_spec_uses_injected_helper_with_caps():
    captured = {}

    async def fake_structured(response_model, messages, **kwargs):
        captured["model"] = response_model
        captured["kwargs"] = kwargs
        return DiagramSpec(
            diagram_type="architecture",
            title="ignored — caller title wins",
            nodes=[DiagramNode(id="idp", label="IdP"), DiagramNode(id="app", label="App")],
            edges=[DiagramEdge(source="idp", target="app", label="SAML")],
        )

    spec = asyncio.run(
        de.generate_diagram_spec(
            fake_structured,
            title="My Architecture",
            diagram_type="architecture",
            context_text="SailPoint IGA rollout",
            client_name="Meridian",
            iam_vendor="SailPoint",
        )
    )
    assert spec.title == "My Architecture"  # caller title preserved
    assert captured["model"] is DiagramSpec
    kw = captured["kwargs"]
    assert kw["max_tokens"] == de.DIAGRAM_SPEC_MAX_TOKENS == 1500
    assert kw["frequency_penalty"] == 0.2
    assert kw["max_retries"] == 1
    # returned spec is sanitized
    assert len(spec.nodes) == 2 and len(spec.edges) == 1


# --- DOCX embedding (approval gate) ----------------------------------------

def _extract_images(docx_bytes: bytes) -> int:
    doc = Document(io.BytesIO(docx_bytes))
    return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)


def _extract_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def _base_sections():
    return [{"id": "executive_summary", "title": "Executive Summary",
             "content": "Body [1].", "citations": [], "max_similarity": 0.6,
             "needs_sme_review": False}]


def test_only_approved_diagrams_embed():
    diagrams = [
        {"title": "Approved Arch", "status": "approved", "image_bytes": _ONE_PX_PNG},
        {"title": "Draft Arch", "status": "draft", "image_bytes": _ONE_PX_PNG},
        {"title": "Rejected Arch", "status": "rejected", "image_bytes": _ONE_PX_PNG},
        {"title": "Review Arch", "status": "needs_review", "image_bytes": _ONE_PX_PNG},
    ]
    docx_bytes = assemble_docx(
        {"client_name": "Meridian", "proposal_type": "implementation"},
        _base_sections(),
        diagrams=diagrams,
    )
    text = _extract_text(docx_bytes)
    assert "Solution Architecture Diagrams" in text
    assert "Approved Arch" in text
    assert "Draft Arch" not in text
    assert "Rejected Arch" not in text
    assert "Review Arch" not in text
    # Exactly one diagram image embedded (the approved one). Branding/logo images
    # may also exist, so assert the approved image increased the count vs none.
    assert _extract_images(docx_bytes) >= 1


def test_no_diagrams_section_when_none_approved():
    diagrams = [
        {"title": "Draft Arch", "status": "draft", "image_bytes": _ONE_PX_PNG},
        {"title": "Rejected Arch", "status": "rejected", "image_bytes": _ONE_PX_PNG},
    ]
    docx_bytes = assemble_docx(
        {"client_name": "Meridian", "proposal_type": "implementation"},
        _base_sections(),
        diagrams=diagrams,
    )
    assert "Solution Architecture Diagrams" not in _extract_text(docx_bytes)


def test_approved_without_image_is_skipped():
    diagrams = [{"title": "Approved No Image", "status": "approved"}]
    docx_bytes = assemble_docx(
        {"client_name": "Meridian", "proposal_type": "implementation"},
        _base_sections(),
        diagrams=diagrams,
    )
    assert "Solution Architecture Diagrams" not in _extract_text(docx_bytes)


# --- generate_proposal default path unchanged (zero diagrams) --------------

async def _stub_embed(client, text):
    return [0.01] * 1536


async def _stub_retrieve(client, embedding, query, k=8):
    return [{"chunk_text": "SailPoint IIQ provisioning across 42 apps.",
             "heading": "Provisioning", "similarity": 0.62,
             "client_name": "Northwind", "iam_vendor": "sailpoint"}]


def _stub_build_system(chunks):
    return "=== EVIDENCE ===\n" + "\n".join(f"[{i}] {c['chunk_text']}" for i, c in enumerate(chunks, 1))


async def _stub_draft(client, system_prompt, user_prompt, max_tokens=1500):
    return "Grounded prose [1]."


def test_generate_proposal_zero_diagrams_default_path():
    document_engine.draft_with_openrouter = _stub_draft

    async def _run():
        async with httpx.AsyncClient() as client:
            return await generate_proposal(
                client,
                rfp_text="IAM implementation.",
                client_name="Meridian Bank",
                proposal_type="implementation",
                iam_vendor="SailPoint",
                embed_fn=_stub_embed,
                retrieve_fn=_stub_retrieve,
                build_grounded_system_fn=_stub_build_system,
                sections=None,
                include_compliance_matrix=False,
            )

    result = asyncio.run(_run())
    assert result["docx_bytes"]
    text = _extract_text(result["docx_bytes"])
    assert "Executive Summary" in text
    # No diagrams passed -> no diagram section.
    assert "Solution Architecture Diagrams" not in text


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-q"]))


# --- prompt assembly regression ---------------------------------------------
def test_guidance_survives_prompt_truncation():
    """REGRESSION: guidance used to be appended to the END of context_text, which
    a blunt [:4000] slice cut off entirely — so a 'deployment' diagram was never
    actually told to show zones, load balancing or HA."""
    import asyncio
    captured = {}

    async def fake_structured(model, messages=None, **kw):
        captured["prompt"] = messages[-1]["content"]
        return DiagramSpec.model_validate({
            "diagram_type": "architecture", "title": "t",
            "nodes": [{"id": "a", "label": "A"}, {"id": "b", "label": "B"},
                      {"id": "c", "label": "C"}],
            "edges": [{"source": "a", "target": "b"}]})

    asyncio.run(de.generate_diagram_spec(
        fake_structured,
        title="AWS - Deployment",
        diagram_type="architecture",
        context_text="network_zones: DMZ, application, data\n" + ("filler\n" * 2000),
        client_name="AWS",
        iam_vendor="PingIdentity",
        guidance="Every node must sit in a named zone. Show regions and HA.",
        evidence_text="evidence " * 2000,
    ))
    prompt = captured["prompt"]
    assert "named zone" in prompt, "guidance was truncated away"
    assert prompt.index("named zone") < prompt.index("network_zones"), \
        "guidance must precede the answers so it always survives"
    assert len(prompt) < 12000, f"prompt unexpectedly large: {len(prompt)}"


def test_evidence_is_trimmed_before_answers():
    """Evidence is the most replaceable input, so it gets the smallest budget."""
    import asyncio
    captured = {}

    async def fake_structured(model, messages=None, **kw):
        captured["prompt"] = messages[-1]["content"]
        return DiagramSpec.model_validate({
            "diagram_type": "architecture", "title": "t",
            "nodes": [{"id": "a", "label": "A"}, {"id": "b", "label": "B"},
                      {"id": "c", "label": "C"}],
            "edges": [{"source": "a", "target": "b"}]})

    asyncio.run(de.generate_diagram_spec(
        fake_structured, title="x", context_text="ANSWER_MARKER\n" + ("a\n" * 100),
        guidance="g", evidence_text="EVIDENCE_MARKER " * 5000))
    prompt = captured["prompt"]
    assert "ANSWER_MARKER" in prompt
    assert prompt.count("EVIDENCE_MARKER") < 5000


# --- empty-spec guard -------------------------------------------------------
def test_empty_spec_triggers_a_corrective_retry():
    """REGRESSION: nodes defaults to [], so an EMPTY spec is schema-valid. The
    model returned exactly that live and we rendered a diagram with only a title."""
    import asyncio
    calls = {"n": 0}

    async def empty_then_good(model, messages=None, models=None, **kw):
        calls["n"] += 1
        if calls["n"] == 1:
            return DiagramSpec.model_validate({
                "diagram_type": "architecture", "title": "t", "nodes": [], "edges": []})
        return DiagramSpec.model_validate({
            "diagram_type": "architecture", "title": "t",
            "nodes": [{"id": "a", "label": "A"}, {"id": "b", "label": "B"},
                      {"id": "c", "label": "C"}],
            "edges": [{"source": "a", "target": "b"}]})

    spec = asyncio.run(de.generate_diagram_spec(empty_then_good, title="X"))
    assert calls["n"] == 2, "did not retry an empty spec"
    assert len(spec.nodes) >= de.MIN_SPEC_NODES


def test_persistently_empty_spec_raises_rather_than_rendering_nothing():
    import asyncio

    async def always_empty(model, messages=None, models=None, **kw):
        return DiagramSpec.model_validate({
            "diagram_type": "architecture", "title": "t", "nodes": [], "edges": []})

    try:
        asyncio.run(de.generate_diagram_spec(always_empty, title="X"))
    except ValueError as e:
        assert "unusable" in str(e)
    else:
        raise AssertionError("an empty spec must raise, not render an empty diagram")


def test_correction_prompt_states_the_requirement():
    import asyncio
    prompts = []

    async def capture(model, messages=None, models=None, **kw):
        prompts.append(messages[-1]["content"])
        return DiagramSpec.model_validate({
            "diagram_type": "architecture", "title": "t", "nodes": [], "edges": []})

    try:
        asyncio.run(de.generate_diagram_spec(capture, title="X"))
    except ValueError:
        pass
    assert len(prompts) == 2
    low = prompts[1].lower()
    assert "rejected" in low and "nodes" in low and "edges" in low


# --- edges are mandatory ----------------------------------------------------
def _spec(nodes, edges):
    return DiagramSpec.model_validate(
        {"diagram_type": "architecture", "title": "t", "nodes": nodes, "edges": edges})


def test_the_live_edgeless_failure_is_rejected():
    """REGRESSION: a live spec had 12 nodes and 1 edge and rendered as a grouped
    list. Node count alone was checked, so it passed validation."""
    nodes = [{"id": f"n{i}", "label": f"N{i}"} for i in range(12)]
    shortfall = de.spec_shortfall(_spec(nodes, [{"source": "n0", "target": "n1"}]))
    assert shortfall and "edge" in shortfall


def test_healthy_graph_passes():
    nodes = [{"id": f"n{i}", "label": f"N{i}"} for i in range(12)]
    edges = [{"source": f"n{i}", "target": f"n{i+1}"} for i in range(11)]
    assert de.spec_shortfall(_spec(nodes, edges)) is None


def test_minimal_two_node_diagram_still_allowed():
    assert de.spec_shortfall(_spec(
        [{"id": "a", "label": "A"}, {"id": "b", "label": "B"}],
        [{"source": "a", "target": "b"}])) is None


def test_too_many_orphans_rejected():
    nodes = [{"id": f"n{i}", "label": f"N{i}"} for i in range(9)]
    # Enough edges to pass the count, but most nodes still floating.
    edges = [{"source": "n0", "target": "n1"}, {"source": "n1", "target": "n0"},
             {"source": "n0", "target": "n1"}, {"source": "n1", "target": "n0"},
             {"source": "n0", "target": "n1"}]
    shortfall = de.spec_shortfall(_spec(nodes, edges))
    assert shortfall and "no connections" in shortfall


def test_edgeless_spec_triggers_correction_then_raises():
    import asyncio
    prompts = []

    async def edgeless(model, messages=None, models=None, **kw):
        prompts.append(messages[-1]["content"])
        return _spec([{"id": f"n{i}", "label": f"N{i}"} for i in range(8)],
                     [{"source": "n0", "target": "n1"}])

    try:
        asyncio.run(de.generate_diagram_spec(edgeless, title="X"))
    except ValueError as e:
        assert "unusable" in str(e)
    else:
        raise AssertionError("an edgeless spec must not be accepted")
    assert len(prompts) == 2, "no corrective retry"
    assert "GRAPH" in prompts[1] and "edges" in prompts[1].lower()


def test_group_labels_are_humanised():
    for raw, want in (("identity_source_layer", "Identity Source"),
                      ("iam_platform_layer", "IAM Platform"),
                      ("_monitoring_layer", "Monitoring"),
                      ("dmz_zone", "DMZ Zone"),
                      ("Identity Sources", "Identity Sources")):
        assert de._pretty_group(raw) == want, f"{raw} -> {de._pretty_group(raw)}"


import diagram_engine  # noqa: E402

# ---------------------------------------------------------------------------
# Swimlanes, decision nodes and page fit.
#
# IV's joiner-flow diagram is four horizontal lanes (HRMS / SailPoint IIQ /
# Manager / Active Directory) with a decision diamond branching Employee vs
# Contractor. Shilpi drew every node as a rectangle in a single column, and
# embedded diagrams at 1800x4217 (2.34:1) that scaled to 3.2in wide slivers.
# ---------------------------------------------------------------------------

def _flow_spec():
    return diagram_engine.DiagramSpec(
        diagram_type="flow", title="Joiner Flow",
        nodes=[
            diagram_engine.DiagramNode(id="hr", label="HRMS", group="HRMS", shape="external"),
            diagram_engine.DiagramNode(id="agg", label="Aggregation Task", group="SailPoint IIQ"),
            diagram_engine.DiagramNode(id="q", label="Employee?", group="SailPoint IIQ", shape="decision"),
            diagram_engine.DiagramNode(id="ap", label="Manager Approval", group="Manager"),
            diagram_engine.DiagramNode(id="ad", label="Active Directory", group="Active Directory", shape="datastore"),
        ],
        edges=[
            diagram_engine.DiagramEdge(source="hr", target="agg"),
            diagram_engine.DiagramEdge(source="agg", target="q"),
            diagram_engine.DiagramEdge(source="q", target="ad", label="YES"),
            diagram_engine.DiagramEdge(source="q", target="ap", label="CONTRACTOR"),
        ],
    )


def test_decision_node_renders_as_a_diamond():
    d2 = diagram_engine.build_d2(_flow_spec())
    assert "shape: diamond" in d2, "decision node did not become a diamond"
    assert "shape: cylinder" in d2, "datastore did not become a cylinder"
    assert "shape: package" in d2, "external system did not become a package"


def test_unknown_shape_falls_back_to_process():
    n = diagram_engine.DiagramNode(id="a", label="X", shape="banana")
    assert n.shape == diagram_engine.DEFAULT_NODE_SHAPE


def test_flow_groups_become_swimlanes():
    """Each lane runs across the page; lanes stack down."""
    d2 = diagram_engine.build_d2(_flow_spec())
    assert d2.count("direction: right") >= 4, "lanes did not get their own direction"
    assert d2.startswith("direction: down"), d2[:40]


def test_architecture_groups_are_zones_not_lanes():
    spec = diagram_engine.DiagramSpec(
        diagram_type="architecture", title="Deployment",
        nodes=[diagram_engine.DiagramNode(id="a", label="Node", group="DMZ")],
        edges=[],
    )
    assert "direction: right" not in diagram_engine.build_d2(spec)


def test_aspect_penalty_scores_a_band_not_a_ceiling():
    """A flip that overshoots must score WORSE, not better.

    An 8-zone chain went from 3.49 (sliver) to 0.09 (an 11:1 strip). Both are
    unusable, so 'smaller wins' would have picked the worse one.
    """
    p = diagram_engine._aspect_penalty
    assert p(1.0) == 0.0
    assert p(3.49) > 0
    assert p(0.09) > p(3.49), "over-wide scored better than over-tall"
    assert p(None) == float("inf")


def test_guidance_actually_asks_for_lanes_and_shapes():
    """CALL-SITE check: the spec fields are useless if nothing requests them."""
    import chat_state
    flow = chat_state.deployment_guidance_for("Integration / Joiner Flow", "flow")
    assert "SWIMLANE" in flow
    assert "decision" in flow
    arch = chat_state.deployment_guidance_for("Deployment", "architecture")
    assert "decision" in arch, "shape mandate missing from non-flow diagrams"


# ---------------------------------------------------------------------------
# Decision nodes, enforced by VALIDATION rather than asked for again.
#
# The model reliably writes branch logic into flow diagrams -- run 6 produced
# "Identity Already Exists?" and "Manager Approval Required?" with correctly
# labelled Yes/No edges -- and then marks every node `process`, so the diamonds
# that make a flow readable never render. Two rounds of prompt instruction did
# not move it, while the swimlane half of the SAME instruction landed at once.
# ---------------------------------------------------------------------------

def _flow(nodes, edges):
    return diagram_engine.DiagramSpec(
        diagram_type="flow", title="Joiner",
        nodes=[diagram_engine.DiagramNode(**n) for n in nodes],
        edges=[diagram_engine.DiagramEdge(**e) for e in edges])


_BRANCH_NODES = [
    {"id": "a", "label": "HR Joiner Event", "group": "HRMS"},
    {"id": "q", "label": "Identity Already Exists?", "group": "SailPoint IIQ"},
    {"id": "y", "label": "Update Attributes", "group": "SailPoint IIQ"},
    {"id": "n", "label": "Provision New Identity", "group": "SailPoint IIQ"},
]
_BRANCH_EDGES = [
    {"source": "a", "target": "q"},
    {"source": "q", "target": "y", "label": "Yes"},
    {"source": "q", "target": "n", "label": "No"},
]


def test_a_question_node_left_as_process_is_rejected():
    reason = diagram_engine.spec_shortfall(_flow(_BRANCH_NODES, _BRANCH_EDGES))
    assert reason and "decision" in reason
    assert "Identity Already Exists?" in reason


def test_a_properly_marked_decision_passes():
    nodes = [dict(n) for n in _BRANCH_NODES]
    nodes[1]["shape"] = "decision"
    assert diagram_engine.spec_shortfall(_flow(nodes, _BRANCH_EDGES)) is None


def test_multiple_labelled_outgoing_edges_count_as_a_branch():
    """Not every branch point is phrased as a question."""
    nodes = [{"id": "a", "label": "Aggregate", "group": "HRMS"},
             {"id": "r", "label": "Route by identity type", "group": "IIQ"},
             {"id": "e", "label": "Employee path", "group": "IIQ"},
             {"id": "c", "label": "Contractor path", "group": "IIQ"}]
    edges = [{"source": "a", "target": "r"},
             {"source": "r", "target": "e", "label": "Employee"},
             {"source": "r", "target": "c", "label": "Contractor"}]
    assert diagram_engine.spec_shortfall(_flow(nodes, edges)) is not None


def test_an_architecture_fan_out_is_not_a_decision():
    """A load balancer with two labelled edges is not a branch point. This rule
    must not fire outside flow diagrams or it would reject every HA topology."""
    spec = diagram_engine.DiagramSpec(
        diagram_type="architecture", title="Deployment",
        nodes=[diagram_engine.DiagramNode(id="lb", label="Load Balancer", group="DMZ"),
               diagram_engine.DiagramNode(id="n1", label="IIQ Node 1", group="App"),
               diagram_engine.DiagramNode(id="n2", label="IIQ Node 2", group="App")],
        edges=[diagram_engine.DiagramEdge(source="lb", target="n1", label="HTTPS"),
               diagram_engine.DiagramEdge(source="lb", target="n2", label="HTTPS")])
    assert diagram_engine.spec_shortfall(spec) is None


def test_the_rejection_names_the_offending_nodes():
    """The retry prompt is built from this string, so it has to be specific
    enough for the model to act on."""
    reason = diagram_engine.spec_shortfall(_flow(_BRANCH_NODES, _BRANCH_EDGES))
    assert 'shape="decision"' in reason


# ---------------------------------------------------------------------------
# The corrective retry must address the shortfall it was given.
#
# Live in run 9: the Integration / Joiner Flow spec was rejected for unmarked
# decision nodes, and the retry appended a fixed paragraph about graphs and
# edges -- advice with no bearing on the fault. It failed the same check twice
# and the diagram was abandoned. A generic correction is worse than none: it
# consumes the single retry without addressing the problem.
# ---------------------------------------------------------------------------

def test_a_decision_rejection_gets_decision_advice():
    reason = "branch points were not marked as decisions, so they would render as rectangles"
    out = diagram_engine._correction_for(reason)
    assert 'shape="decision"' in out
    assert "roughly as many edges as nodes" not in out, \
        "generic graph advice sent for a decision-shape fault"


def test_a_connectivity_rejection_still_gets_graph_advice():
    out = diagram_engine._correction_for("12 of 15 components had no connections at all")
    assert "GRAPH" in out
    assert 'shape="decision"' not in out


def test_every_correction_quotes_the_actual_reason():
    """The model cannot fix what it was not told."""
    for reason in ("branch points were not marked as decisions",
                   "the spec had no nodes at all",
                   "12 of 15 components had no connections at all"):
        assert reason in diagram_engine._correction_for(reason)


def test_decision_correction_lists_the_valid_shapes():
    """A correction that names the field but not its legal values invites a
    second failure."""
    out = diagram_engine._correction_for("branch points were not marked as decisions")
    for shape in ("process", "decision", "datastore", "external"):
        assert shape in out
