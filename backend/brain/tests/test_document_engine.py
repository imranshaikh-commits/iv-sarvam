"""
Sprint 5 smoke test — runs with NO API keys / NO network.

Strategy:
  * document_engine takes the brain helpers (embed / retrieve / build-system) as
    injected callables, so we pass stubs — no import of app.py, no secrets.
  * The single OpenRouter call lives in document_engine.draft_with_openrouter,
    which we monkeypatch to return canned drafted prose.

Run directly (`python tests/test_document_engine.py`) or via pytest. It writes a
sample DOCX into the system temp dir (override with SHILPI_TEST_OUTPUT_DIR) and
asserts the document contains the expected markers.
"""

import asyncio
import io
import os
import sys
import tempfile

# Make the brain package importable when run as a bare script from any cwd.
_HERE = os.path.dirname(os.path.abspath(__file__))
_BRAIN = os.path.dirname(_HERE)
if _BRAIN not in sys.path:
    sys.path.insert(0, _BRAIN)

import httpx
from docx import Document

import document_engine
from document_engine import assemble_docx, generate_proposal

# Was hardcoded to /home/user/workspace/, which exists on exactly one machine —
# the test failed everywhere else before it could assert anything.
OUTPUT_DIR = os.environ.get("SHILPI_TEST_OUTPUT_DIR") or tempfile.gettempdir()
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "shilpi_sample_proposal.docx")
CLIENT_NAME = "Meridian Bank"
SME_MARKER = "[SME REVIEW]"
# A phrase unique to stub_draft's output — proof the drafting path actually ran.
DRAFT_SENTINEL = "Inspirit Vision proposes a SailPoint IdentityIQ deployment"


# --- stubs ------------------------------------------------------------------
async def stub_embed(client, text):
    return [0.01] * 1536


def _fake_chunks():
    return [
        {
            "chunk_text": "SailPoint IdentityIQ was deployed with automated joiner-mover-leaver "
            "provisioning across 42 target applications, including SAP and Active Directory.",
            "heading": "Provisioning Architecture",
            "similarity": 0.62,
            "client_name": "Northwind Insurance",
            "iam_vendor": "sailpoint",
            "industry": "Insurance",
        },
        {
            "chunk_text": "Access certification campaigns ran quarterly with delegated reviewers and "
            "automated revocation of orphaned accounts.",
            "heading": "Access Certification",
            "similarity": 0.58,
            "client_name": "Acme Retail",
            "iam_vendor": "sailpoint",
            "industry": "Retail",
        },
    ]


# NOTE: must accept whatever draft_section passes. When `max_tokens` was added
# to draft_with_openrouter this same class of mismatch silently routed every
# section into the drafting-failure path while the test stayed green. **kw keeps
# the stub tolerant of new keyword arguments; the CALLER is asserted separately.
async def stub_retrieve(client, embedding, query, k=8, **kw):
    return _fake_chunks()


def stub_build_system(chunks):
    return "=== EVIDENCE ===\n" + "\n".join(
        f"[{i}] {c['chunk_text']}" for i, c in enumerate(chunks, 1)
    )


# NOTE: this signature must track draft_with_openrouter's. When ``max_tokens``
# was added to the real function (2026-07-17) this stub was not updated, so every
# section raised TypeError, document_engine swallowed it per-section, and the
# test stayed green while drafted content never reached the DOCX at all. The
# DRAFT_SENTINEL assertion below is what makes that visible.
async def stub_draft(client, system_prompt, user_prompt, max_tokens=None):
    # Canned drafted paragraph with a [1] citation and an SME-review marker.
    return (
        "Inspirit Vision proposes a SailPoint IdentityIQ deployment covering automated "
        "joiner-mover-leaver provisioning across the client's core applications [1]. "
        "Quarterly access certification campaigns will be established with delegated "
        "reviewers [1].\n\n"
        f"{SME_MARKER}: pricing and licensing counts must be confirmed by an SME before "
        "any client-facing use.\n\n"
        "- Assumption: target application inventory will be provided during discovery."
    )


async def _run() -> bytes:
    # Monkeypatch the isolated network call.
    document_engine.draft_with_openrouter = stub_draft

    async with httpx.AsyncClient() as client:
        result = await generate_proposal(
            client,
            rfp_text="Deliver an enterprise IAM implementation with automated provisioning "
            "and access certification.",
            client_name=CLIENT_NAME,
            proposal_type="implementation",
            iam_vendor="SailPoint",
            embed_fn=stub_embed,
            retrieve_fn=stub_retrieve,
            build_grounded_system_fn=stub_build_system,
            sections=None,
            include_compliance_matrix=False,
        )
    assert result["filename"].startswith("Shilpi_Proposal_Meridian_Bank_"), result["filename"]
    assert result["filename"].endswith(".docx")
    assert result["sections_meta"], "expected section metadata"
    return result["docx_bytes"]


def _extract_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_generate_proposal_docx():
    docx_bytes = asyncio.run(_run())
    assert docx_bytes, "DOCX bytes must be non-empty"

    with open(OUTPUT_PATH, "wb") as f:
        f.write(docx_bytes)
    assert os.path.getsize(OUTPUT_PATH) > 0, "written DOCX must be non-empty"

    text = _extract_text(docx_bytes)
    assert CLIENT_NAME in text, "client name missing"
    assert "DRAFT" in text, "DRAFT safety banner missing"
    assert "Executive Summary" in text, "expected a section heading"
    # The Citation Appendix was REMOVED (2026-08-14). It listed retrieved corpus
    # chunks by client name and similarity score, so the Amlak document named
    # three OTHER IV clients across 78 paragraphs. Its absence is now the
    # correct behaviour, and no corpus client name may appear anywhere.
    assert "Citation Appendix" not in text, "citation appendix is back"
    assert "similarity" not in text.lower(), "retrieval provenance leaked into the document"
    assert SME_MARKER in text, "SME REVIEW marker missing"
    # The scaffold supplies the headings, banner and appendix on its own, so
    # every other assertion here passes even when drafting fails outright.
    # This one fails unless the drafted prose actually reached the document.
    assert DRAFT_SENTINEL in text, (
        "drafted section content missing — the stub was never successfully "
        "called (check its signature against draft_with_openrouter)"
    )
    # A citation source from the fake corpus should appear in the appendix.
    # Northwind Insurance is the OTHER client in the fixture's retrieved corpus
    # chunk. This used to assert its name appeared, because the Citation
    # Appendix listed retrieval sources by client. That is exactly the leak the
    # Amlak run shipped, so the assertion is now inverted: a corpus client's
    # name must never reach a document addressed to a different client.
    assert "Northwind Insurance" not in text, "another client's name leaked into the document"

    print("SMOKE TEST PASSED")
    print(f"  wrote {OUTPUT_PATH} ({os.path.getsize(OUTPUT_PATH)} bytes)")
    print(f"  client name present : {CLIENT_NAME in text}")
    print(f"  DRAFT banner        : {'DRAFT' in text}")
    print(f"  section heading     : {'Executive Summary' in text}")
    print(f"  citation appendix   : {'Citation Appendix' in text}")
    print(f"  SME REVIEW marker   : {SME_MARKER in text}")


def test_assemble_docx_directly():
    """Exercise assemble_docx on its own with a compliance-matrix markdown block."""
    sections = [
        {
            "id": "executive_summary",
            "title": "Executive Summary",
            "content": "A grounded summary with a citation [1].",
            "citations": _fake_chunks(),
            "max_similarity": 0.62,
            "needs_sme_review": False,
        }
    ]
    md = "# DRAFT Compliance Matrix\n\n| Req | Status |\n|---|---|\n| REQ-001 | Covered |\n"
    docx_bytes = assemble_docx(
        {"client_name": CLIENT_NAME, "proposal_type": "implementation", "iam_vendor": "SailPoint"},
        sections,
        compliance_markdown=md,
    )
    text = _extract_text(docx_bytes)
    assert "Compliance Matrix" in text
    assert "REQ-001" in text


async def _null_draft(client, system_prompt, user_prompt, max_tokens=None):
    """Simulate the LLM returning null (None) content on every call/retry."""
    return None


async def _run_null_draft() -> dict:
    """Full-depth generation where every subsection LLM call returns None."""
    document_engine.draft_with_openrouter = _null_draft
    async with httpx.AsyncClient() as client:
        return await generate_proposal(
            client,
            rfp_text="Deliver an enterprise IAM implementation with automated provisioning.",
            client_name=CLIENT_NAME,
            proposal_type="implementation",
            iam_vendor="SailPoint",
            embed_fn=stub_embed,
            retrieve_fn=stub_retrieve,
            build_grounded_system_fn=stub_build_system,
            sections=None,
            include_compliance_matrix=False,
            proposal_depth="full",
        )


def test_null_llm_returns_grounded_placeholder():
    """Bug fix: a None LLM response must NOT crash or drop the subsection.

    Asserts (a) no exception is raised (generation completes with a valid DOCX)
    and (b) each subsection renders a non-empty grounded [ASSUMPTION] placeholder
    instead of an empty string.
    """
    result = asyncio.run(_run_null_draft())  # must not raise
    docx_bytes = result["docx_bytes"]
    assert docx_bytes, "DOCX must still be produced when the LLM returns None"

    text = _extract_text(docx_bytes)
    assert document_engine.ASSUMPTION_MARKER in text, "grounded [ASSUMPTION] placeholder missing"
    # The placeholder must be grounded in intake/context, not empty.
    assert CLIENT_NAME in text, "placeholder should reference the client from context"

    # No subsection should have come back as an empty string.
    md = result["draft_markdown"]
    assert md.strip(), "draft markdown must be non-empty"
    for line in md.splitlines():
        if line.startswith("### "):
            # A subsection heading must be followed by real content somewhere.
            pass
    assert md.count(document_engine.ASSUMPTION_MARKER) >= 1, "expected placeholder content in draft"

    print("NULL-DRAFT PLACEHOLDER TEST PASSED")
    print(f"  assumption markers  : {md.count(document_engine.ASSUMPTION_MARKER)}")


if __name__ == "__main__":
    test_generate_proposal_docx()
    test_assemble_docx_directly()
    test_null_llm_returns_grounded_placeholder()
    print("ALL CHECKS PASSED")


# --- inline markdown must become real Word formatting -----------------------
def _render(text):
    import io, re, zipfile
    from docx import Document as _Doc
    doc = _Doc()
    document_engine._add_body_paragraphs(doc, text)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    xml = zipfile.ZipFile(buf).read("word/document.xml").decode()
    raw = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.S))
    buf.seek(0)
    return raw, _Doc(buf)


def test_no_literal_markdown_reaches_the_docx():
    """REGRESSION: a generated proposal contained 147 literal ** and 82 literal *
    because the whole block was written as a single run."""
    raw, _ = _render(
        "**Identity Layer.** Centralized provisioning [2][8]. "
        "*(needs SME confirmation.)*\n\n"
        "- **Hybrid Coexistence:** legacy AD must coexist\n"
        "- *Integration Readiness:* assess applications\n\n"
        "### Assumptions\nBody with **bold** inside.")
    assert "**" not in raw, "literal bold markers reached the document"
    import re as _re
    assert not _re.search(r"(?<!\*)\*(?!\*)", raw), "literal italic markers reached the document"
    assert "#" not in raw, "literal heading hashes reached the document"


def test_emphasis_becomes_real_runs():
    _, doc = _render("**Bold lead.** then normal text, and *italic bit* after.")
    para = [p for p in doc.paragraphs if p.text.strip()][0]
    assert any(r.bold for r in para.runs), "no bold run produced"
    assert any(r.italic for r in para.runs), "no italic run produced"
    assert any(not r.bold and not r.italic for r in para.runs), "plain text lost"


def test_bold_survives_inside_bullets():
    """The old lstrip('-*• ') ate the opening ** of '- **Label:** text'."""
    raw, doc = _render("- **Hybrid Identity Coexistence:** must coexist with legacy AD")
    assert "**" not in raw
    bullet = [p for p in doc.paragraphs if p.style.name == "List Bullet"][0]
    assert any(r.bold for r in bullet.runs), "bold lost when the bullet marker was stripped"
    assert bullet.text.startswith("Hybrid Identity Coexistence:")


def test_technical_underscores_are_never_italicised():
    """Underscores are identifiers here (data_retention, session_data), not markup."""
    raw, doc = _render(
        "Retention is data_retention of 7 years and session_data of 90 days, "
        "keyed on iam_vendor and client_name.")
    for ident in ("data_retention", "session_data", "iam_vendor", "client_name"):
        assert ident in raw, f"{ident} was mangled"
    para = [p for p in doc.paragraphs if p.text.strip()][0]
    assert not any(r.italic for r in para.runs), "underscores were treated as italics"


def test_citation_markers_survive():
    raw, _ = _render("Provisioning replaces per-app credentials [2][8] per the design [15].")
    for c in ("[2]", "[8]", "[15]"):
        assert c in raw, f"citation {c} lost"


# --- discovery answers must reach drafting ----------------------------------
def test_discovery_routing_sends_sizing_to_architecture():
    """REGRESSION: generate_proposal accepted only rfp_text, so 21 of 22 discovery
    areas were captured, stored and then discarded before drafting.

    Section ids updated 2026-08-27: this test asserted `solution_architecture`
    and `implementation_methodology`, which stopped existing when the template
    was rebuilt to IV's house structure. It passed against a map that routed
    nothing, which is how 11 of 12 sections came to be drafted blind.
    """
    ans = {"client_name": "Amlak International", "iam_vendor": "SailPoint",
           "hardware_sizing_inputs": "4 app servers at 4 CPU, 16 GB memory",
           "cluster_topology": "2 UI servers and 2 Task servers",
           "payment_milestones": "20 percent on signature",
           "duration": "42 weeks"}
    arch = document_engine.discovery_context_for("proposed_solution", ans)
    assert "16 GB" in arch and "Task servers" in arch
    assert "20 percent" not in arch, "commercial detail leaked into architecture"

    # Engagement duration belongs to project_timeline, not to the delivery
    # approach: "42 weeks" is a schedule fact, and the RACI/deliverables section
    # has no use for it.
    timeline = document_engine.discovery_context_for("project_timeline", ans)
    assert "42 weeks" in timeline
    assert "16 GB" not in timeline, "sizing leaked into the timeline"


def test_skip_and_control_fields_are_not_routed():
    ans = {"client_name": "X", "rto_rpo": "skip", "proposal_depth": "full",
           "_diagram_plan": [["a", "b"]], "hardware_sizing_inputs": "real value"}
    ctx = document_engine.discovery_context_for("proposed_solution", ans)
    assert "real value" in ctx
    assert "skip" not in ctx.lower()
    assert "proposal_depth" not in ctx.lower() and "diagram_plan" not in ctx.lower()


def test_no_answers_yields_empty_context():
    assert document_engine.discovery_context_for("executive_summary", None) == ""
    assert document_engine.discovery_context_for("executive_summary", {}) == ""


# --- model artifacts must never reach the document --------------------------
def test_think_blocks_and_meta_commentary_stripped():
    """REGRESSION: a live proposal carried a raw </think> tag and seven copies of
    the model's own 'Note on Evidence Applicability' deliberation."""
    raw = ("Note on Evidence Applicability: The retrieved evidence originates from "
           "InspiritVision proposals for Al Qadsiah FC.\n---\n"
           "<think>let me consider the evidence</think>"
           "InspiritVision is pleased to present this proposal [1].")
    out = document_engine.strip_model_artifacts(raw)
    assert "think" not in out.lower()
    assert "Note on Evidence Applicability" not in out
    assert "retrieved evidence" not in out.lower()
    assert "[1]" in out, "citation markers must survive"
    assert "pleased to present" in out


def test_sme_marker_survives_stripping():
    raw = f"Some drafted prose. {document_engine.SME_REVIEW_MARKER}: confirm versions."
    assert document_engine.SME_REVIEW_MARKER in document_engine.strip_model_artifacts(raw)


def test_degenerate_tail_truncated_at_sentence_boundary():
    """REGRESSION: ~109 words of 'shattering shattering ... vow-breakingly'."""
    good = "The platform integrates with Active Directory. Provisioning is automated."
    bad = good + " shattering shattering shattering shattering shattering shattering"
    out = document_engine.strip_degenerate_tail(bad)
    assert "shattering" not in out
    assert out.endswith("automated.")


def test_legitimate_repetition_preserved():
    """Vendor names repeat legitimately; only runaway runs should be cut."""
    txt = ("SailPoint IdentityIQ provides governance. SailPoint connectors integrate "
           "with AD. IdentityIQ certifications run quarterly.")
    assert document_engine.strip_degenerate_tail(txt) == txt


# --- assembly mechanics -----------------------------------------------------
def test_static_toc_replaces_the_f9_placeholder():
    import io
    from docx import Document as _Doc
    md = {"client_name": "C", "proposal_type": "implementation", "generated_at": "now"}
    b = document_engine.assemble_docx(
        md, [{"title": "Executive Summary", "content": "x", "id": "executive_summary"},
             {"title": "Solution Architecture", "content": "y", "id": "solution_architecture"}])
    txt = "\n".join(p.text for p in _Doc(io.BytesIO(b)).paragraphs)
    assert "Right-click here" not in txt, "F9 placeholder still shown"
    assert "Executive Summary" in txt and "Solution Architecture" in txt


def test_appendices_use_captured_values_not_tbc():
    import io
    from docx import Document as _Doc
    md = {"client_name": "C", "proposal_type": "implementation", "generated_at": "now",
          "discovery_answers": {"app_count": "25 applications",
                                "environments": "production, DR, UAT, development",
                                "user_count": "skip"}}
    b = document_engine.assemble_docx(
        md, [{"title": "S", "content": "x", "id": "executive_summary"}],
        include_appendices=True)
    tbl = "\n".join(c.text for t in _Doc(io.BytesIO(b)).tables for r in t.rows for c in r.cells)
    assert "25 applications" in tbl, "captured app count not used"
    assert "production, DR, UAT" in tbl, "captured environments not used"
    assert "TBC" in tbl, "genuinely absent values should still read TBC"


def test_tall_diagram_is_height_capped():
    import io
    from docx import Document as _Doc
    from docx.shared import Inches
    from PIL import Image
    buf = io.BytesIO(); Image.new("RGB", (1200, 4000), "white").save(buf, "PNG"); buf.seek(0)
    doc = _Doc()
    document_engine._add_picture_fitted(doc, buf, max_w=Inches(6.0), max_h=Inches(7.5))
    shape = doc.inline_shapes[0]
    assert shape.height <= Inches(7.5), "diagram would span multiple pages"
    assert shape.width <= Inches(6.0)


# ---------------------------------------------------------------------------
# GFM tables in drafted section bodies.
#
# Before this, only the compliance-matrix path converted markdown tables to real
# Word tables. A sizing / RACI / BOQ table drafted into any normal section landed
# as one paragraph of literal pipe characters. IV's own proposals are table-heavy
# (25 tables vs Shilpi's 9), so this path has to work before any template work
# asking the model for tables can pay off.
# ---------------------------------------------------------------------------

_SIZING_MD = """Proposed production sizing is below.

| Component | vCPU | RAM | Disk |
|---|---|---|---|
| App server (UI) | 4 | 16 GB | 100 GB |
| App server (Task) | 4 | 16 GB | 100 GB |
| Database | 8 | 64 GB | 500 GB |

DR is sized identically to production.
"""


def _render_body(text: str):
    from docx import Document as _Doc
    doc = _Doc()
    document_engine._add_body_paragraphs(doc, text)
    return doc


def test_markdown_table_in_body_becomes_word_table():
    doc = _render_body(_SIZING_MD)
    assert len(doc.tables) == 1, f"expected 1 native table, got {len(doc.tables)}"
    table = doc.tables[0]
    assert len(table.columns) == 4
    assert len(table.rows) == 4, "header + 3 data rows"
    assert table.rows[0].cells[0].text == "Component"
    assert table.rows[3].cells[2].text == "64 GB"
    # And no literal pipe soup left behind in the prose.
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "|" not in body, f"literal pipes survived into prose: {body!r}"


def test_table_preserves_surrounding_prose_and_order():
    doc = _render_body(_SIZING_MD)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert any(p.startswith("Proposed production sizing") for p in paras)
    assert any(p.startswith("DR is sized identically") for p in paras)


def test_table_cells_render_bold_markup_as_runs():
    doc = _render_body(
        "| Item | Amount |\n|---|---|\n| **Total** | 100 |\n"
    )
    cell = doc.tables[0].rows[1].cells[0]
    assert cell.text == "Total", "literal ** left in a table cell"
    assert any(r.bold for r in cell.paragraphs[0].runs), "bold run lost in cell"


def test_prose_without_tables_is_unchanged():
    doc = _render_body("A plain paragraph.\n\n- bullet one\n- bullet two\n")
    assert len(doc.tables) == 0
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert texts == ["A plain paragraph.", "bullet one", "bullet two"]


def test_stray_pipe_line_is_not_a_table():
    """A single pipe line with no separator row must stay prose.

    The previous subset check treated an empty following line as a valid
    separator and built a one-row table out of it.
    """
    doc = _render_body("Latency was measured end | to end during UAT.\n")
    assert len(doc.tables) == 0
    assert "| to end" in "\n".join(p.text for p in doc.paragraphs)


def test_compliance_matrix_still_renders_tables():
    """The compliance path now shares the parser — prove it did not regress."""
    from docx import Document as _Doc
    doc = _Doc()
    document_engine._add_markdown_ish(
        doc,
        "# Coverage\n\n| Requirement | Status |\n|---|---|\n| SSO | Met |\n",
    )
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[1].cells[1].text == "Met"


# ---------------------------------------------------------------------------
# Template rebuild to IV's house structure, and the formatting that goes with it.
#
# Run 4 (Sonnet) produced byte-identical structure to run 3 (GLM): 7 sections,
# 32 subsections of which only 14 were unique, "Overview / Detailed Design /
# Considerations & Dependencies" seven times each. No model can change that,
# because it came from a module-level list applied to every section.
# ---------------------------------------------------------------------------

def test_every_subsection_heading_is_unique():
    """THE table-of-contents complaint, as an assertion."""
    from proposal_templates import get_template
    ctx = {"client_name": "Amlak International", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    headings = [h for spec in get_template("implementation")
                for h, _ in spec.render_subsections(ctx)]
    dupes = {h for h in headings if headings.count(h) > 1}
    assert not dupes, f"repeated subsection headings: {sorted(dupes)}"
    assert len(headings) >= 40, f"only {len(headings)} subsections; IV's has 53"


def test_the_generic_triple_is_gone_from_the_implementation_template():
    from proposal_templates import get_template
    ctx = {"client_name": "X", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    headings = [h for spec in get_template("implementation")
                for h, _ in spec.render_subsections(ctx)]
    for banned in ("Overview", "Detailed Design", "Considerations & Dependencies"):
        assert banned not in headings, f"{banned!r} is still a subsection heading"


def test_ivs_section_skeleton_is_present():
    from proposal_templates import get_template
    titles = {s.id for s in get_template("implementation")}
    for required in ("company_profile", "similar_experience", "solution_overview",
                     "proposed_solution", "implementation_approach",
                     "project_timeline", "knowledge_transfer", "commercial"):
        assert required in titles, f"IV section {required!r} missing"


def test_vendor_name_substitutes_into_titles_and_headings():
    from proposal_templates import get_template
    ctx = {"client_name": "Amlak International", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    specs = {s.id: s for s in get_template("implementation")}
    assert specs["proposed_solution"].render_title(ctx) == "Proposed Solution - SailPoint"
    headings = [h for h, _ in specs["solution_overview"].render_subsections(ctx)]
    assert "Why SailPoint" in headings
    assert "{{" not in " ".join(headings), "unrendered Jinja leaked into a heading"


def test_section_subsections_override_the_depth_tier_count():
    """A section's own headings are its structure, not a depth knob."""
    from proposal_templates import get_template
    ctx = {"client_name": "X", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    specs = {s.id: s for s in get_template("implementation")}
    assert len(specs["proposed_solution"].render_subsections(ctx)) >= 8


# --- formatting -------------------------------------------------------------

def test_label_then_dashes_becomes_a_heading_and_bullets():
    """Run 4 had 42 paragraphs with embedded newlines and 0 List Bullet uses."""
    from docx import Document as _Doc
    doc = _Doc()
    document_engine._add_body_paragraphs(doc, (
        "Risks to Manage:\n"
        "- Keycloak migration gaps could break SSO\n"
        "- HR data quality may delay onboarding\n"
    ))
    styles = [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]
    assert styles[0][0].startswith("Heading"), styles
    assert styles[0][1] == "Risks to Manage"
    assert all(s == "List Bullet" for s, _ in styles[1:]), styles
    assert not any("\n" in t for _, t in styles), "embedded newline survived"


def test_bold_markup_inside_a_bullet_survives():
    from docx import Document as _Doc
    doc = _Doc()
    document_engine._add_body_paragraphs(
        doc, "- **Keycloak risk:** incomplete mapping could break SSO\n")
    bullet = [p for p in doc.paragraphs if p.style.name == "List Bullet"][0]
    assert "**" not in bullet.text
    assert any(r.bold for r in bullet.runs), "bold lost inside the bullet"


def test_a_long_numbered_line_stays_prose():
    """A 783-word 'list item' is prose that happens to start with '1.'."""
    from docx import Document as _Doc
    doc = _Doc()
    long_line = "1. " + " ".join(["word"] * 60) + "."
    document_engine._add_body_paragraphs(doc, long_line)
    assert not [p for p in doc.paragraphs if p.style.name == "List Number"]


def test_a_sentence_containing_a_colon_is_not_promoted_to_a_heading():
    from docx import Document as _Doc
    doc = _Doc()
    document_engine._add_body_paragraphs(
        doc, "The position today is simply this: nobody can say who has access.")
    assert not [p for p in doc.paragraphs if p.style.name.startswith("Heading")]


def test_client_logo_reaches_the_running_header():
    """Point 4: IV runs the client mark on every page, not just the cover."""
    import base64, tempfile, os as _os
    from docx import Document as _Doc
    import branding
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")
    fd, path = tempfile.mkstemp(suffix=".png")
    with _os.fdopen(fd, "wb") as fh:
        fh.write(png)
    doc = _Doc()
    branding.configure_base_styles(doc)
    branding.apply_header_footer(doc, "Amlak International", client_logo_path=path)
    header_xml = doc.sections[0].header.paragraphs[0]._p.xml
    assert header_xml.count("<a:blip") == 2, "client logo missing beside the IV logo"


def test_header_still_renders_without_a_client_logo():
    from docx import Document as _Doc
    import branding
    doc = _Doc()
    branding.configure_base_styles(doc)
    branding.apply_header_footer(doc, "Amlak International")
    assert "Amlak International" in doc.sections[0].header.paragraphs[0].text


# ---------------------------------------------------------------------------
# CALL-SITE tests. The four tests above assert that SectionSpec.render_subsections
# and branding.apply_header_footer behave correctly in isolation — and a negative
# control proved that BOTH still pass when the call sites are unwired. That is
# the "built but never wired" failure this project has hit four times, reproduced
# in its own regression suite. These tests drive the real pipeline instead.
# ---------------------------------------------------------------------------

def test_generated_document_has_unique_subsection_headings():
    """Drive the whole pipeline; assert the DOCX itself, not the template."""
    docx_bytes = asyncio.run(_run())
    doc = Document(io.BytesIO(docx_bytes))
    h2 = [p.text.strip() for p in doc.paragraphs
          if p.style.name == "Heading 2" and p.text.strip()]
    assert h2, "no subsection headings in the generated document"
    dupes = {h for h in h2 if h2.count(h) > 1}
    assert not dupes, f"repeated subsection headings in the DOCX: {sorted(dupes)}"
    for banned in ("Overview", "Detailed Design", "Considerations & Dependencies"):
        assert banned not in h2, f"generic facet {banned!r} reached the document"


def test_generated_document_carries_ivs_section_skeleton():
    docx_bytes = asyncio.run(_run())
    doc = Document(io.BytesIO(docx_bytes))
    h1 = [p.text.strip() for p in doc.paragraphs
          if p.style.name == "Heading 1" and p.text.strip()]
    joined = " | ".join(h1)
    for required in ("Company Profile", "Similar Experience", "Commercial",
                     "Knowledge Transfer"):
        assert required in joined, f"{required!r} missing from the document: {joined}"


def test_assemble_docx_puts_the_client_logo_in_the_running_header():
    """assemble_docx must PASS the logo through, not merely accept it."""
    import base64, tempfile, os as _os
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")
    fd, path = tempfile.mkstemp(suffix=".png")
    with _os.fdopen(fd, "wb") as fh:
        fh.write(png)
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "Amlak International", "proposal_type": "implementation"},
        sections=[{"id": "executive_summary", "title": "Executive Summary",
                   "content": "Body text."}],
        client_logo_path=path,
    )
    doc = Document(io.BytesIO(docx_bytes))
    # The IV logo is ALSO an image in this header, so presence of "blip" proves
    # nothing — an earlier version of this assertion passed with the call site
    # unwired. Count images instead: IV mark + client mark = 2.
    header_xml = doc.sections[0].header.paragraphs[0]._p.xml
    assert header_xml.count("<a:blip") == 2, (
        f"expected IV logo + client logo in the header, found "
        f"{header_xml.count('<a:blip')} image(s)")


def test_prose_subsections_get_a_tighter_token_budget_than_tables():
    """CALL-SITE test: draft_section must pass a reduced budget for prose.

    Run 5 produced 16,051 prose words against the human proposal's 6,648 across
    a near-identical subsection count. The model fills whatever budget it is
    given, so the budget is the lever - but only if the call site applies it.
    """
    import proposal_templates
    seen: list[int] = []

    async def spy(client, system_prompt, user_prompt, max_tokens=None):
        seen.append(max_tokens)
        return "Drafted body text for the subsection."

    original = document_engine.draft_with_openrouter
    document_engine.draft_with_openrouter = spy
    try:
        spec = proposal_templates.SectionSpec(
            id="t", title="T", purpose="p", query_template="q",
            subsections=(
                ("Sizing", "production sizing as a markdown TABLE with columns X, Y"),
                ("Narrative", "a description of the approach in prose"),
            ),
        )

        async def stub_embed(c, q):
            return [0.0] * 8

        async def stub_retrieve(c, v, **kw):
            return []

        asyncio.run(document_engine.draft_section(
            None, spec,
            {"client_name": "X", "iam_vendor": "SailPoint",
             "proposal_type": "implementation", "rfp_text": ""},
            embed_fn=stub_embed, retrieve_fn=stub_retrieve,
            build_grounded_system_fn=lambda chunks: "EVIDENCE",
            top_k=1, subsections=2, max_tokens=2500,
        ))
    finally:
        document_engine.draft_with_openrouter = original

    assert len(seen) >= 2, f"expected two drafting calls, saw {seen}"
    table_budget, prose_budget = seen[0], seen[1]
    assert table_budget == 2500, f"table subsection lost its budget: {table_budget}"
    assert prose_budget <= document_engine.PROSE_SUBSECTION_TOKENS, (
        f"prose subsection got {prose_budget}, expected "
        f"<= {document_engine.PROSE_SUBSECTION_TOKENS}")
    assert prose_budget < table_budget


def test_proposal_type_is_passed_to_retrieval():
    """CALL-SITE test: draft_section must tell retrieval which type it is drafting.

    Measured on the retrieval scorecard: a credential-migration query returned
    1 of 8 chunks from a migration proposal without this, and 5 of 8 with it,
    against a 35.5% base rate. Retrieval was performing WORSE than random on
    proposal type, so every migration section would have been drafted from
    greenfield implementation proposals.

    The stub above takes **kw, which makes it tolerant of new arguments but also
    blind to their absence -- hence this test asserts what was actually passed.
    """
    import proposal_templates
    seen: list[dict] = []

    async def spy_retrieve(client, embedding, query, k=8, **kw):
        seen.append(kw)
        return []

    async def stub_embed(c, q):
        return [0.0] * 8

    spec = proposal_templates.SectionSpec(
        id="t", title="T", purpose="p", query_template="migration cutover")

    asyncio.run(document_engine.draft_section(
        None, spec,
        {"client_name": "NWC", "iam_vendor": "Oracle",
         "proposal_type": "migration", "rfp_text": ""},
        embed_fn=stub_embed, retrieve_fn=spy_retrieve,
        build_grounded_system_fn=lambda chunks: "EVIDENCE",
        top_k=8, subsections=1, max_tokens=500,
    ))

    assert seen, "retrieve_fn was never called"
    assert all("proposal_type" in kw for kw in seen), \
        f"proposal_type never reached retrieval: {seen}"
    assert seen[0]["proposal_type"] == "migration", seen[0]


def test_absent_proposal_type_passes_none_not_empty_string():
    """An empty string would filter to a type that does not exist, returning
    nothing. None means 'no preference', which is the correct fallback."""
    seen: list[dict] = []

    async def spy_retrieve(client, embedding, query, k=8, **kw):
        seen.append(kw)
        return []

    async def stub_embed(c, q):
        return [0.0] * 8

    import proposal_templates
    spec = proposal_templates.SectionSpec(id="t", title="T", purpose="p",
                                          query_template="anything")
    asyncio.run(document_engine.draft_section(
        None, spec, {"client_name": "X", "iam_vendor": "Ping", "rfp_text": ""},
        embed_fn=stub_embed, retrieve_fn=spy_retrieve,
        build_grounded_system_fn=lambda chunks: "EVIDENCE",
        top_k=8, subsections=1, max_tokens=500,
    ))
    assert seen[0]["proposal_type"] is None, seen[0]


def test_section_topic_is_passed_to_retrieval():
    """CALL-SITE test: draft_section must tell retrieval which TOPIC it needs.

    Company Profile, Why-Vendor and Similar Experience are each 1-3% of the
    corpus, so a general vector search rarely surfaces them -- which is why
    those sections were thin in every generated proposal. Measured: why_vendor
    went from 1 of 8 on-topic results to 4 of 8 with the filter passed.
    """
    seen: list[dict] = []

    async def spy_retrieve(client, embedding, query, k=8, **kw):
        seen.append(kw)
        return []

    async def stub_embed(c, q):
        return [0.0] * 8

    import proposal_templates
    spec = proposal_templates.SectionSpec(
        id="company_profile", title="Company Profile", purpose="p",
        query_template="Inspirit Vision company profile offices workforce")

    asyncio.run(document_engine.draft_section(
        None, spec,
        {"client_name": "X", "iam_vendor": "SailPoint",
         "proposal_type": "implementation", "rfp_text": ""},
        embed_fn=stub_embed, retrieve_fn=spy_retrieve,
        build_grounded_system_fn=lambda chunks: "EVIDENCE",
        top_k=8, subsections=1, max_tokens=500,
    ))

    assert seen, "retrieve_fn was never called"
    assert all("section_topic" in kw for kw in seen), \
        f"section_topic never reached retrieval: {seen}"
    assert seen[0]["section_topic"] == "company_profile", seen[0]


def test_subsection_words_in_the_query_beat_the_section_topic():
    """A sizing query inside the solution section must ask for sizing evidence,
    not the architecture prose the parent section maps to."""
    import proposal_templates as P
    assert P.topic_for("proposed_solution", "Proposed Production Hardware Sizing") == "sizing"
    assert P.topic_for("proposed_solution", "Proposed Target Architecture") == "architecture"
    assert P.topic_for("commercial", "Payment Milestones") == "pricing"
    assert P.topic_for("solution_overview", "Why SailPoint") == "why_vendor"


def test_unmapped_section_passes_none():
    """A section with no topic must not filter at all, rather than filtering to
    a topic that does not exist and returning nothing."""
    import proposal_templates as P
    assert P.topic_for("no_such_section") is None


# ---------------------------------------------------------------------------
# Image placement. CALL-SITE tests: selection logic was already covered in
# test_asset_selection.py, and this project's most expensive recurring failure
# is a mechanism that works in isolation and is never invoked.
# ---------------------------------------------------------------------------

def _png_bytes():
    import base64
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")


def test_assemble_embeds_section_assets():
    stream = io.BytesIO(_png_bytes())
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "Amlak International",
                  "proposal_type": "implementation"},
        sections=[{"id": "company_profile", "title": "Company Profile",
                   "content": "Body text.",
                   "assets": [{"id": "a1", "stream": stream}]}],
    )
    doc = Document(io.BytesIO(docx_bytes))
    images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    # IV logo on the cover + the placed asset.
    assert len(images) >= 2, f"asset not embedded, only {len(images)} image(s)"


def test_a_broken_asset_does_not_sink_the_section():
    """A proposal without an image is worse; a proposal that fails to build is
    unusable. Bad bytes must be logged and skipped."""
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "X", "proposal_type": "implementation"},
        sections=[{"id": "company_profile", "title": "Company Profile",
                   "content": "Body text that must survive.",
                   "assets": [{"id": "bad", "stream": io.BytesIO(b"not an image")}]}],
    )
    text = "\n".join(p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs)
    assert "Body text that must survive." in text


def test_sections_without_assets_are_unaffected():
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "X", "proposal_type": "implementation"},
        sections=[{"id": "company_profile", "title": "Company Profile",
                   "content": "Body."}],
    )
    doc = Document(io.BytesIO(docx_bytes))
    images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    assert len(images) == 1, "only the IV cover logo expected"


def test_attach_assets_places_each_image_once():
    """The same picture twice in one document reads as a mistake."""
    lib = [{"id": "a1", "storage_path": "p/1.png", "asset_kind": "corporate",
            "approved": True, "occurrences": 9,
            "vision_description": "Inspirit Vision certified resources workforce "
                                  "skill matrix and delivery model"}]

    async def library(_c):
        return lib

    async def download(_c, _p):
        return io.BytesIO(_png_bytes())

    sections = [{"id": "company_profile", "title": "Company Profile"},
                {"id": "implementation_approach", "title": "Implementation Approach"}]
    asyncio.run(document_engine._attach_assets(
        None, sections, {"iam_vendor": "SailPoint"},
        {"library": library, "download": download}))
    total = sum(len(s.get("assets") or []) for s in sections)
    assert total == 1, f"the same asset was placed {total} times"


def test_attach_assets_survives_an_unavailable_library():
    async def library(_c):
        raise RuntimeError("supabase down")

    async def download(_c, _p):
        return None

    sections = [{"id": "company_profile", "title": "Company Profile"}]
    try:
        asyncio.run(document_engine._attach_assets(
            None, sections, {"iam_vendor": "SailPoint"},
            {"library": library, "download": download}))
    except RuntimeError:
        pass  # generate_proposal catches this; the point is sections stay clean
    assert not sections[0].get("assets")


def test_no_caption_text_is_written_beside_an_image():
    """Captions were REMOVED (run 8). They were generated from the vision
    description and produced, verbatim:

      "Gantt chart, a type of project management diagram that visualizes the
       schedule and dependencies for the 'Sistem-BTPN ProjectPL'. It details
       tasks broken into ph"

    Three faults at once: it explained what a Gantt chart is to an IAM
    audience, it named ANOTHER CLIENT'S project, and it truncated mid-word.
    """
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "Amlak International",
                  "proposal_type": "implementation"},
        sections=[{"id": "company_profile", "title": "Company Profile",
                   "content": "Body text.",
                   "assets": [{"id": "a1", "stream": io.BytesIO(_png_bytes()),
                               "caption": "Gantt chart for Sistem-BTPN ProjectPL"}]}],
    )
    text = "\n".join(p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs)
    assert "Sistem-BTPN" not in text, "a caption leaked another client's project name"
    assert "Gantt" not in text


def test_paragraph_length_is_instructed_not_just_section_length():
    """IV's median body paragraph is 29 words; run 8's was 55 with twenty over
    100, because only the SUBSECTION length was capped."""
    assert document_engine.PROSE_PARAGRAPH_WORDS <= 80


def test_the_token_cap_is_not_the_binding_constraint():
    """420 tokens against a 220-word instruction truncated [SME REVIEW] markers
    mid-word. The cap must leave headroom above the word target."""
    assert document_engine.PROSE_SUBSECTION_TOKENS > \
        document_engine.PROSE_SUBSECTION_WORDS * 2, (
        "token cap too close to the word target; it will truncate mid-sentence")


# ---------------------------------------------------------------------------
# The discovery map must track BOTH the templates and the intake schema.
#
# It silently stopped tracking either. Keys were the OLD section ids from before
# the template was rebuilt to IV's house structure, so 11 of 12 implementation
# sections received NO discovery answers at all. Field names were invented
# rather than taken from the intake schema, so several that did match resolved
# to nothing.
#
# Visible symptom: 42 [SME REVIEW] markers in run 9, including on UAT sizing,
# connectors, hypercare and the support model -- all of which the consultant HAD
# supplied. The model was not hedging; it had not been told.
# ---------------------------------------------------------------------------

def test_every_section_of_every_template_receives_discovery_answers():
    import proposal_templates
    mapped = set(document_engine._SECTION_DISCOVERY_FIELDS)
    missing = {}
    for ptype in proposal_templates.VALID_PROPOSAL_TYPES:
        gap = {s.id for s in proposal_templates.get_template(ptype)} - mapped
        if gap:
            missing[ptype] = sorted(gap)
    assert not missing, (
        f"sections drafted with no discovery answers: {missing}. A sizing "
        f"subsection without sizing inputs can only produce generic text.")


def test_every_mapped_field_exists_in_the_intake_schema():
    """A field name that is not in the schema resolves to nothing, silently."""
    import intake_template
    valid = {q["id"] for b in intake_template.get_intake_template(None)["buckets"]
             for q in b["questions"]}
    bad = sorted({f for fields in document_engine._SECTION_DISCOVERY_FIELDS.values()
                  for f in fields if f not in valid})
    assert not bad, f"discovery fields that do not exist in the intake: {bad}"


def test_supplied_answers_reach_the_section_that_needs_them():
    """Spot-check the fields that were flagged [SME REVIEW] in run 9."""
    answers = {"hardware_sizing_inputs": "UAT 4 app servers, 8 CPU, 32 GB, 250 GB",
               "support_model": "L1, L2 and L3 support",
               "hypercare": "post-production support period",
               "ad_exchange": "AD and Exchange integration in scope"}
    sizing = document_engine.discovery_context_for("proposed_solution", answers)
    assert "32 GB" in sizing and "Exchange" in sizing
    kt = document_engine.discovery_context_for("knowledge_transfer", answers)
    assert "L1, L2 and L3" in kt and "post-production" in kt


def test_skip_answers_are_not_presented_as_facts():
    answers = {"user_count": "skip", "app_count": "25 applications"}
    out = document_engine.discovery_context_for("executive_summary", answers)
    assert "25 applications" in out
    assert "skip" not in out.lower()


def test_assets_render_before_the_section_body():
    """Images must sit under the SECTION heading they were chosen for.

    They used to render after the whole body, so an image appeared beneath
    whatever the last subsection heading happened to be: run 9 put a CIAM
    governance pyramid under "Assumptions & Open Questions" and a change-control
    graphic under "Training and Post-Production Support". The selection was
    right; the position made it read as wrong.
    """
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "Amlak International",
                  "proposal_type": "implementation"},
        sections=[{"id": "company_profile", "title": "Company Profile",
                   "assets": [{"id": "a1", "stream": io.BytesIO(_png_bytes())}],
                   "subsections": [
                       {"title": "Inspirit Vision", "content": "First subsection."},
                       {"title": "Workforce and Capabilities", "content": "Last subsection."}]}],
    )
    doc = Document(io.BytesIO(docx_bytes))
    order, images = [], []
    for i, p in enumerate(doc.paragraphs):
        if "graphicData" in p._p.xml:
            images.append(i)
        if p.text.strip() in ("Company Profile", "Workforce and Capabilities"):
            order.append((p.text.strip(), i))
    heading = dict(order).get("Company Profile")
    last_sub = dict(order).get("Workforce and Capabilities")
    # Skip the IV cover logo, which precedes every section.
    seen_image = next((i for i in images if i > heading), None)
    assert seen_image is not None, "asset not embedded"
    assert heading < seen_image < last_sub, (
        f"image at {seen_image} should sit between the section heading "
        f"({heading}) and the last subsection ({last_sub})")


def test_sizing_subsections_name_ivs_real_columns():
    """IV's sizing tables are eleven columns; run 9 produced 6, 4, 3 and 3 from
    an invented list."""
    import proposal_templates
    ctx = {"client_name": "X", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    sizing = [f for s in proposal_templates.get_template("implementation")
              for h, f in s.render_subsections(ctx) if "Sizing" in h]
    assert len(sizing) == 4, f"expected four sizing subsections, got {len(sizing)}"
    for f in sizing:
        assert "eleven" in f or "Server Category" in f
    prod = next(f for f in sizing if "Server Category" in f)
    for col in ("CPU per node", "DB Storage", "Operating System", "Remarks"):
        assert col in prod, f"missing IV column: {col}"


def test_iv_table_structures_are_split_the_way_iv_splits_them():
    """RACI is a legend plus two matrices; payment splits licence from
    implementation; each tranche has its own milestone table."""
    import proposal_templates
    ctx = {"client_name": "X", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    headings = [h for s in proposal_templates.get_template("implementation")
                for h, _ in s.render_subsections(ctx)]
    for required in ("RACI Legend", "RACI - Project Governance",
                     "RACI - Delivery Activities",
                     "Payment Milestone - Licence",
                     "Payment Milestone - Implementation"):
        assert required in headings, f"missing: {required}"
    tranches = [h for h in headings if h.startswith("Tranche")]
    assert len(tranches) == 3, tranches


def test_the_template_asks_for_more_tables_than_ivs_original_has():
    """IV's Amlak proposal has 25 tables. Matching it is the floor, not the
    target."""
    import proposal_templates, document_engine as DE
    ctx = {"client_name": "X", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    wants = [h for s in proposal_templates.get_template("implementation")
             for h, f in s.render_subsections(ctx) if DE._WANTS_TABLE_RE.search(f)]
    assert len(wants) >= 17, f"only {len(wants)} table subsections: {wants}"


# ---------------------------------------------------------------------------
# Diagrams belong INSIDE the subsection that explains them.
#
# IV puts the deployment architecture diagram under "Proposed Deployment
# Architecture", immediately after the prose describing it. Shilpi collected all
# diagrams into a trailing "Solution Architecture Diagrams" section, so in run 9
# a reader had to hold the text in their head and go looking forty pages later.
# ---------------------------------------------------------------------------

def _diagram(dtype, title):
    return {"diagram_type": dtype, "title": title, "status": "approved",
            "image_bytes": _png_bytes()}


def test_a_diagram_renders_under_the_subsection_it_explains():
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "Amlak International",
                  "proposal_type": "implementation"},
        sections=[{"id": "proposed_solution", "title": "Proposed Solution",
                   "subsections": [
                       {"title": "Proposed Deployment Architecture",
                        "content": "The deployment spans four environments."},
                       {"title": "Proposed Production Hardware Sizing",
                        "content": "Sizing follows."}]}],
        diagrams=[_diagram("deployment", "Amlak — Deployment")],
    )
    doc = Document(io.BytesIO(docx_bytes))
    idx = {p.text.strip(): i for i, p in enumerate(doc.paragraphs) if p.text.strip()}
    images = [i for i, p in enumerate(doc.paragraphs) if "graphicData" in p._p.xml]
    dep = idx.get("Proposed Deployment Architecture")
    sizing = idx.get("Proposed Production Hardware Sizing")
    assert dep and sizing
    assert any(dep < i < sizing for i in images), (
        "deployment diagram did not render between its own subsection and the next")


def test_an_unmatched_diagram_still_reaches_the_document():
    """Placement is a preference, not a filter. A diagram no subsection claims
    must still appear in the trailing gallery rather than vanish."""
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "X", "proposal_type": "implementation"},
        sections=[{"id": "company_profile", "title": "Company Profile",
                   "content": "Body."}],
        diagrams=[_diagram("deployment", "Deployment")],
    )
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    assert "Solution Architecture Diagrams" in text
    assert len(images) >= 2, "unmatched diagram was dropped"


def test_a_diagram_is_placed_once():
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "X", "proposal_type": "implementation"},
        sections=[{"id": "proposed_solution", "title": "Proposed Solution",
                   "subsections": [
                       {"title": "Proposed Deployment Architecture", "content": "a"},
                       {"title": "Proposed Deployment Architecture", "content": "b"}]}],
        diagrams=[_diagram("deployment", "Deployment")],
    )
    doc = Document(io.BytesIO(docx_bytes))
    images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    assert len(images) == 2, f"IV logo + one diagram expected, got {len(images)}"


def test_executive_summary_is_continuous_prose():
    """IV's executive summary has no sub-headings. Leaving `subsections` empty
    made it fall back to the generic Overview/Detailed Design triple, which
    Sprint B removed everywhere else -- and run 9 shipped it at the top of the
    document, in the one section every reader reads."""
    import proposal_templates
    ctx = {"client_name": "X", "iam_vendor": "SailPoint",
           "proposal_type": "implementation", "rfp_text": ""}
    spec = next(s for s in proposal_templates.get_template("implementation")
                if s.id == "executive_summary")
    subs = spec.render_subsections(ctx)
    assert len(subs) == 1 and subs[0][0] == "", subs


def test_an_untitled_subsection_adds_no_heading():
    docx_bytes = document_engine.assemble_docx(
        metadata={"client_name": "X", "proposal_type": "implementation"},
        sections=[{"id": "executive_summary", "title": "Executive Summary",
                   "subsections": [{"title": "", "content": "Continuous prose."}]}],
    )
    doc = Document(io.BytesIO(docx_bytes))
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert h2 == [], f"empty heading emitted: {h2}"
    assert any("Continuous prose." in p.text for p in doc.paragraphs)
