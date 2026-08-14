"""QA gate tests — written against the ACTUAL defects in the 2026-08-07 Amlak run."""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import document_qa as qa  # noqa: E402


def test_phrase_level_degeneration_detected():
    """REGRESSION: the previous check looked for the same WORD repeating. The
    model's real failure was the same PHRASE repeating, so it sailed through."""
    bad = ("The platform enforces controls. " + "stated constraints " * 16)
    degenerate, reason = qa.is_degenerate(bad)
    assert degenerate, "phrase-level padding not detected"
    assert "repeats" in reason


def test_real_amlak_degenerate_paragraph():
    """Verbatim shape from the live run: long sentence then a phrase loop."""
    real = ("Amlak International will gain automated user lifecycle management. "
            + "encrypted storage secrets TLS HTTPS web app connections role-based UI "
              "task RBAC ABAC model auditable logs SoD policy checks " * 6)
    degenerate, _ = qa.is_degenerate(real)
    assert degenerate


def test_legitimate_technical_prose_not_flagged():
    """Vendor names and IAM terms repeat legitimately — no false positives."""
    good = (
        "SailPoint IdentityIQ provides access certification for managers, application "
        "owners and entitlement owners. The platform integrates with Active Directory "
        "for identity lookup and with the HRMS as the authoritative source for joiner, "
        "mover and leaver events. Provisioning covers on-premises and cloud "
        "applications, and de-provisioning is automated on termination. SailPoint also "
        "supports segregation-of-duties policy checks with detective and preventive "
        "controls, plus password management including self-service reset."
    )
    degenerate, reason = qa.is_degenerate(good)
    assert not degenerate, f"false positive on clean prose: {reason}"


def test_short_paragraphs_never_flagged():
    assert not qa.is_degenerate("Short text repeated repeated repeated.")[0]


def test_citations_fully_stripped():
    """User directive: citation markers gone completely from client output."""
    src = "unified architecture [6]. The scope covers deployment [1][7]. HRMS feeds it [3]."
    out = qa.strip_citations(src)
    assert "[" not in out and "]" not in out
    assert "architecture." in out, "punctuation not tidied after stripping"
    assert "deployment." in out


def test_meta_commentary_new_variants_removed():
    """REGRESSION: banning 'the retrieved evidence' just made the model rephrase
    to 'Evidence does not confirm...' — so match the SUBJECT, not the wording."""
    for sentence in (
        "Evidence does not confirm specific SailPoint product edition version number.",
        "Evidence cited originates from prior engagements with different clients.",
        "The retrieved evidence originates from InspiritVision proposals.",
        "The evidence base lacks Amlak-specific volumetrics.",
    ):
        assert qa.strip_meta_commentary(sentence).strip() == "", sentence


def test_real_proposal_content_survives_meta_stripping():
    keep = ("SailPoint IdentityIQ delivers certification campaigns for managers, "
            "application owners and entitlement owners.")
    assert qa.strip_meta_commentary(keep).strip() == keep


def test_meta_sentences_are_extractable_for_review():
    src = ("The platform is deployed on-premises. Evidence does not confirm the "
           "HRMS product used as authoritative source.")
    found = qa.extract_meta_commentary(src)
    assert len(found) == 1 and "HRMS product" in found[0]


def test_qa_text_reports_what_it_changed():
    src = "Architecture is centralised [2]. Evidence does not confirm the edition."
    out, issues = qa.qa_text(src)
    assert "[2]" not in out
    assert any("citation" in i for i in issues)
    assert any("meta-commentary" in i for i in issues)


def test_truncation_keeps_good_prose_and_ends_cleanly():
    good = "The platform integrates with Active Directory. Provisioning is automated."
    out = qa.truncate_at_degeneration(good + " " + "stated constraints " * 16)
    assert "stated constraints stated constraints" not in out
    assert out.rstrip().endswith((".", "!", "?"))


def test_qa_report_names_degenerate_sections():
    sections = [
        {"title": "Good Section", "content": "Clean prose about identity governance."},
        {"title": "Bad Section",
         "content": "Intro sentence here. " + "stated constraints " * 16},
    ]
    rep = qa.qa_report(sections)
    assert "Bad Section" in rep["degenerate_sections"]
    assert "Good Section" not in rep["degenerate_sections"]


# ---------------------------------------------------------------------------
# Run-3 failure modes. Each of these passed the gate as it stood, which is why
# they reached a client-facing document.
# ---------------------------------------------------------------------------

# Verbatim shape of the collapse seen in Amlak run 3: a single-pass thesaurus
# walk. It never repeats a phrase, so the n-gram check scores it a clean zero.
_THESAURUS_WALK = (
    "The governance model is defined below. category classification grouping "
    "categorization taxonomy hierarchy structure organization arrangement ordering "
    "sequence layout format presentation display appearance look feel style design "
    "pattern template model framework architecture blueprint plan scheme system "
    "method approach technique procedure process workflow methodology strategy "
    "tactic roadmap path route direction trajectory course journey voyage expedition "
    "trek hike climb ascent rise growth development evolution progression advancement "
    "improvement enhancement refinement optimization maximization utilization "
    "exploitation leverage harnessing employing using applying implementing executing "
    "deploying launching initiating starting beginning commencing originating creating "
    "generating producing building constructing assembling forming shaping molding "
    "crafting designing planning conceiving imagining envisioning visualizing"
)


def test_thesaurus_walk_is_caught_although_nothing_repeats():
    """THE run-3 miss. Repetition score is zero; it is still not prose."""
    count, _ = qa.repetition_score(_THESAURUS_WALK)
    assert count < qa.DEGENERATE_REPEATS, "precondition: nothing repeats here"
    degenerate, reason = qa.is_degenerate(_THESAURUS_WALK)
    assert degenerate, "thesaurus walk slipped through the gate again"
    assert "punctuation" in reason or "function-word" in reason


def test_phrase_repetition_is_still_caught():
    """Regression guard: the run-2 failure mode must not stop being detected."""
    text = "Opening sentence. " + "stated constraints and requirements today " * 5
    degenerate, reason = qa.is_degenerate(text)
    assert degenerate and "repeats" in reason


def test_normal_technical_prose_is_not_flagged():
    """False positives cause pointless re-drafts of good content."""
    text = (
        "SailPoint IdentityIQ will be deployed across four environments. The "
        "production tier comprises two UI servers and two task servers running "
        "on Tomcat, each provisioned with 4 CPU and 16 GB of memory. A single "
        "database server provides 8 CPU, 64 GB of memory and 500 GB of "
        "RAID-protected storage. The disaster recovery environment is sized "
        "identically to production, while UAT is reduced to 32 GB of memory and "
        "250 GB of storage. A development environment is also provisioned."
    )
    assert qa.is_degenerate(text) == (False, "")


def test_a_long_markdown_table_is_not_flagged_as_degenerate():
    """Table rows have almost no function words by construction."""
    rows = "\n".join(f"| App {i} | LDAP | Direct | Phase 1 |" for i in range(30))
    text = "The application onboarding plan is below.\n\n" + \
           "| Application | Protocol | Connector | Tranche |\n|---|---|---|---|\n" + rows
    assert qa.is_degenerate(text)[0] is False


def test_truncation_keeps_the_good_head_of_a_collapsed_block():
    kept = qa.truncate_at_degeneration(_THESAURUS_WALK)
    assert kept.startswith("The governance model is defined below.")
    assert "taxonomy hierarchy structure organization" not in kept
    assert qa.is_degenerate(kept)[0] is False, "truncated text is still degenerate"


def test_parenthetical_review_markers_are_removed():
    """45 of these survived run 3 because the old regex required a dash."""
    text = ("The deployment model is on-premises. (needs SME confirmation) "
            "Four environments are in scope (needs SME confirmation on sizing).")
    out = qa.strip_review_markers(text)
    assert "SME" not in out, out
    assert "The deployment model is on-premises." in out
    assert "Four environments are in scope" in out


def test_dash_led_review_markers_are_still_removed():
    text = "DR is included - needs SME confirmation on the replication target."
    assert "SME" not in qa.strip_review_markers(text)


def test_tbc_and_unverified_asides_are_removed():
    for aside in ["(TBC)", "[to be confirmed]", "(unverified)", "(TBD)"]:
        out = qa.strip_review_markers(f"Sizing is agreed {aside} for production.")
        assert "confirm" not in out.lower() and "TB" not in out, aside


def test_em_dashes_are_replaced():
    """IV uses 2 per 9,900 words; run 3 used 203 per 15,551."""
    out = qa.strip_em_dashes("The platform is modern — and it is fast.")
    assert "—" not in out and "–" not in out
    assert "modern, and it is fast" in out


def test_em_dash_as_a_label_separator_keeps_heading_shape():
    out = qa.strip_em_dashes("Tranche 2 — Life Cycle Management")
    assert out == "Tranche 2 - Life Cycle Management", out


def test_qa_text_applies_every_cleaner_and_reports_them():
    text = ("Deployment is on-premises [1]. (needs SME confirmation) "
            "The architecture is proven — and repeatable.")
    out, issues = qa.qa_text(text)
    assert "[1]" not in out
    assert "SME" not in out
    assert "—" not in out
    joined = " ".join(issues)
    assert "citation" in joined and "review aside" in joined and "em-dash" in joined


def test_real_run3_paragraphs_are_caught_and_the_human_control_is_clean():
    """The only honest test: the actual document that failed, plus its control.

    Skipped when the fixtures are absent (they are client data and are not in
    the repo).
    """
    import os
    from docx import Document
    gen = "/mnt/user-data/uploads/Shilpi_Proposal_Amlak_International_20260811_111337.docx"
    human = "/mnt/user-data/uploads/InspiritVision_Amlak_Technical_Proposal_SP_V1_0.docx"
    if not (os.path.exists(gen) and os.path.exists(human)):
        return
    flagged = [t for t in (p.text for p in Document(gen).paragraphs)
               if t.strip() and qa.is_degenerate(t)[0]]
    assert len(flagged) >= 3, f"expected the known bad paragraphs, got {len(flagged)}"
    control = [t for t in (p.text for p in Document(human).paragraphs)
               if t.strip() and qa.is_degenerate(t)[0]]
    assert control == [], f"FALSE POSITIVES on the human proposal: {len(control)}"


if __name__ == "__main__":
    passed = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
            passed += 1
    print(f"ALL {passed} DOCUMENT QA TESTS PASSED")
