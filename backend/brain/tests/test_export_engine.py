"""
Sprint 5 Pass 6 — export pipeline tests. Keyless, no network, no live soffice.

Covers:
  (a) lite DOCX compression reduces size + preserves a valid DOCX
  (b) fail-soft: no images / already small / cannot reach target
  (c) PDF export via a MOCKED soffice subprocess + graceful missing-binary handling
  (d) Supabase upload + signed-URL (MOCKED), incl. missing-bucket / failure cases
  (e) legacy /v1/generate-proposal response unchanged by default (DOCX binary)
  (f) opt-in export returns JSON with DOCX/PDF metadata + signed URLs

soffice / OpenRouter / Supabase are all mocked — nothing real is invoked.
"""

import asyncio
import io
import os
import subprocess
import sys

# Make the brain package importable when run from any cwd.
_HERE = os.path.dirname(os.path.abspath(__file__))
_BRAIN = os.path.dirname(_HERE)
if _BRAIN not in sys.path:
    sys.path.insert(0, _BRAIN)

# Dummy (non-secret) env so app.py + supabase_client import without real creds.
os.environ.setdefault("OPENROUTER_API_KEY", "test-key")
os.environ.setdefault("SUPABASE_URL", "https://example.supabase.co")
os.environ.setdefault("SUPABASE_KEY", "test-key")

from docx import Document
from PIL import Image

import export_engine


# --- helpers ----------------------------------------------------------------

def _docx_bytes(document: Document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _image_stream(size, fmt="JPEG") -> io.BytesIO:
    """A high-entropy image so JPEG can't trivially compress it (forces the
    downscale/quality loop to do real work)."""
    w, h = size
    img = Image.frombytes("RGB", (w, h), os.urandom(w * h * 3))
    bio = io.BytesIO()
    if fmt == "JPEG":
        img.save(bio, format="JPEG", quality=95)
    else:
        img.save(bio, format="PNG")
    bio.seek(0)
    return bio


def _docx_with_image(size=(1200, 1200), fmt="JPEG") -> bytes:
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Proposal", level=1)
    doc.add_paragraph("Body text.")
    doc.add_picture(_image_stream(size, fmt), width=Inches(6.0))
    return _docx_bytes(doc)


def _plain_docx() -> bytes:
    doc = Document()
    doc.add_heading("Proposal", level=1)
    doc.add_paragraph("No images here.")
    return _docx_bytes(doc)


def _opens_as_docx(data: bytes) -> bool:
    try:
        Document(io.BytesIO(data))
        return True
    except Exception:
        return False


# --- (a) compression reduces size + preserves valid DOCX --------------------

def test_compress_reduces_size_and_preserves_docx():
    original = _docx_with_image(size=(1400, 1400), fmt="JPEG")
    # Force the loop with a small target well under the image-heavy original.
    out, meta = export_engine.compress_docx_lite(original, max_bytes=60_000)
    assert meta["compressed"] is True
    assert meta["final_size"] < meta["original_size"]
    assert len(out) < len(original)
    assert meta["images_processed"] == 1
    assert _opens_as_docx(out), "compressed DOCX must still open as a valid Word doc"


# --- (b) fail-soft branches -------------------------------------------------

def test_compress_no_images_is_noop():
    original = _plain_docx()
    out, meta = export_engine.compress_docx_lite(original, max_bytes=1)
    assert out == original  # unchanged bytes
    assert meta["images_processed"] == 0
    assert "no embedded images" in meta["note"]


def test_compress_already_small_is_noop():
    original = _docx_with_image(size=(200, 200), fmt="JPEG")
    out, meta = export_engine.compress_docx_lite(original, max_bytes=export_engine.LITE_MAX_BYTES)
    assert out == original
    assert "already under target" in meta["note"]


def test_compress_cannot_reach_target_warns_but_survives():
    original = _docx_with_image(size=(1400, 1400), fmt="JPEG")
    out, meta = export_engine.compress_docx_lite(original, max_bytes=1)  # impossible
    assert meta["under_target"] is False
    assert "warning" in meta
    assert _opens_as_docx(out), "best-effort result must remain a valid DOCX"


def test_compress_bad_zip_fails_soft():
    out, meta = export_engine.compress_docx_lite(b"not a zip at all")
    assert out == b"not a zip at all"
    assert "warning" in meta


# --- (c) PDF export (mocked soffice) ----------------------------------------

def test_export_pdf_mocked_success(monkeypatch):
    monkeypatch.setattr(export_engine, "_soffice_binary", lambda: "/usr/bin/soffice")

    def fake_run(cmd, **kwargs):
        outdir = cmd[cmd.index("--outdir") + 1]
        with open(os.path.join(outdir, "proposal.pdf"), "wb") as f:
            f.write(b"%PDF-1.4 mocked pdf\n%%EOF")
        return subprocess.CompletedProcess(cmd, 0, b"", b"")

    monkeypatch.setattr(export_engine.subprocess, "run", fake_run)

    pdf, err = export_engine.export_pdf(_plain_docx())
    assert err is None
    assert pdf is not None and pdf.startswith(b"%PDF")


def test_export_pdf_missing_binary_fails_soft(monkeypatch):
    monkeypatch.setattr(export_engine, "_soffice_binary", lambda: None)
    pdf, err = export_engine.export_pdf(_plain_docx())
    assert pdf is None
    assert "not found" in err


def test_export_pdf_conversion_error_fails_soft(monkeypatch):
    monkeypatch.setattr(export_engine, "_soffice_binary", lambda: "/usr/bin/soffice")

    def boom(cmd, **kwargs):
        raise subprocess.CalledProcessError(1, cmd, stderr=b"conversion boom")

    monkeypatch.setattr(export_engine.subprocess, "run", boom)
    pdf, err = export_engine.export_pdf(_plain_docx())
    assert pdf is None
    assert "PDF conversion failed" in err


def test_export_pdf_no_output_file_fails_soft(monkeypatch):
    monkeypatch.setattr(export_engine, "_soffice_binary", lambda: "/usr/bin/soffice")
    # Succeeds but writes nothing -> no output file.
    monkeypatch.setattr(
        export_engine.subprocess, "run",
        lambda cmd, **kw: subprocess.CompletedProcess(cmd, 0, b"", b""),
    )
    pdf, err = export_engine.export_pdf(_plain_docx())
    assert pdf is None
    assert "no output" in err


# --- (d) Supabase storage upload + signed URL (mocked) ----------------------

import supabase_client  # noqa: E402


class _FakeResp:
    def __init__(self, *, json_data=None, raise_exc=None):
        self._json = json_data or {}
        self._raise = raise_exc

    def raise_for_status(self):
        if self._raise:
            raise self._raise

    def json(self):
        return self._json


class _FakeClient:
    """Minimal async httpx.AsyncClient stand-in returning a queued response."""

    def __init__(self, resp=None, exc=None):
        self._resp = resp
        self._exc = exc
        self.calls = []

    async def post(self, url, **kwargs):
        self.calls.append((url, kwargs))
        if self._exc:
            raise self._exc
        return self._resp


def test_upload_generated_draft_success():
    client = _FakeClient(resp=_FakeResp())
    path = asyncio.run(
        supabase_client.upload_generated_draft(client, "prop/file.docx", b"data", "application/x")
    )
    assert path == "prop/file.docx"
    assert client.calls and "generated-drafts" in client.calls[0][0]


def test_upload_generated_draft_missing_bucket_fails_soft():
    client = _FakeClient(exc=RuntimeError("bucket not found"))
    path = asyncio.run(
        supabase_client.upload_generated_draft(client, "prop/file.docx", b"data", "application/x")
    )
    assert path is None


def test_create_signed_url_success():
    resp = _FakeResp(json_data={"signedURL": "/object/sign/generated-drafts/prop/file.docx?token=abc"})
    client = _FakeClient(resp=resp)
    url = asyncio.run(supabase_client.create_signed_url(client, "prop/file.docx"))
    assert url == (
        f"{supabase_client.SUPABASE_URL}/storage/v1"
        "/object/sign/generated-drafts/prop/file.docx?token=abc"
    )


def test_create_signed_url_failure_fails_soft():
    client = _FakeClient(exc=RuntimeError("no bucket"))
    url = asyncio.run(supabase_client.create_signed_url(client, "prop/file.docx"))
    assert url is None


def test_insert_generated_proposal_uses_valid_status():
    """Regression: status must be 'drafting' (a value allowed by the
    generated_proposals_status_check CHECK constraint), never 'draft' (which 400s
    and left generated_proposals empty)."""
    resp = _FakeResp(json_data=[{"id": "prop-999"}])
    client = _FakeClient(resp=resp)
    pid = asyncio.run(
        supabase_client.insert_generated_proposal(
            client,
            org_id="org-1",
            client_name="Test Client",
            proposal_type="implementation",
        )
    )
    assert pid == "prop-999"
    assert client.calls, "expected an insert POST"
    payload = client.calls[0][1]["json"]
    assert payload["status"] == "drafting"


# --- (e)+(f) endpoint behaviour (legacy vs opt-in) --------------------------

from fastapi.testclient import TestClient  # noqa: E402

import app as app_module  # noqa: E402

_SAMPLE_DOCX = _plain_docx()


async def _fake_generate_proposal(client, **kwargs):
    return {
        "docx_bytes": _SAMPLE_DOCX,
        "sections_meta": [{"id": "executive_summary", "title": "Executive Summary"}],
        "draft_markdown": "## Executive Summary\n\nBody.",
        "filename": "Sarvam_Proposal_Test_20260101_000000.docx",
        "included_compliance_matrix": False,
        "proposal_depth": "standard",
        "included_appendices": False,
    }


async def _fake_insert(*args, **kwargs):
    return "prop-123"


def _patch_common(monkeypatch):
    monkeypatch.setattr(app_module, "generate_proposal", _fake_generate_proposal)
    monkeypatch.setattr(app_module.supabase_client, "insert_generated_proposal", _fake_insert)


def _base_body():
    return {
        "rfp_text": "Deliver an IAM implementation.",
        "client_name": "Test Client",
        "proposal_type": "implementation",
    }


def test_legacy_response_is_docx_binary_by_default(monkeypatch):
    _patch_common(monkeypatch)
    with TestClient(app_module.app) as tc:
        resp = tc.post("/v1/generate-proposal", json=_base_body())
    assert resp.status_code == 200
    assert resp.headers["content-type"] == app_module.DOCX_MEDIA_TYPE
    assert resp.content == _SAMPLE_DOCX  # byte-for-byte unchanged
    assert "attachment; filename=" in resp.headers["content-disposition"]


def test_optin_export_returns_json_with_metadata_and_signed_urls(monkeypatch):
    _patch_common(monkeypatch)

    async def fake_upload(client, path, data, content_type):
        return path

    async def fake_sign(client, path, expires_in=3600):
        return f"https://example.supabase.co/storage/v1/object/sign/generated-drafts/{path}?token=x"

    monkeypatch.setattr(app_module.supabase_client, "upload_generated_draft", fake_upload)
    monkeypatch.setattr(app_module.supabase_client, "create_signed_url", fake_sign)
    monkeypatch.setattr(app_module.export_engine, "export_pdf", lambda b: (b"%PDF-1.4 mock", None))

    body = {**_base_body(), "lite": True, "include_pdf": True, "return_signed_urls": True}
    with TestClient(app_module.app) as tc:
        resp = tc.post("/v1/generate-proposal", json=body)
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("application/json")
    data = resp.json()
    assert data["docx"]["lite"] is True
    assert data["docx"]["size"] > 0
    assert data["pdf"]["size"] == len(b"%PDF-1.4 mock")
    assert data["signed_urls"]["docx"].startswith("https://example.supabase.co/")
    assert data["signed_urls"]["pdf"].startswith("https://example.supabase.co/")
    assert data["generated_proposal_id"] == "prop-123"


def test_optin_pdf_error_and_missing_bucket_fail_soft(monkeypatch):
    _patch_common(monkeypatch)

    async def fake_upload_fail(client, path, data, content_type):
        return None  # bucket missing

    monkeypatch.setattr(app_module.supabase_client, "upload_generated_draft", fake_upload_fail)
    monkeypatch.setattr(
        app_module.export_engine, "export_pdf",
        lambda b: (None, "LibreOffice (soffice) binary not found on PATH; PDF export unavailable"),
    )

    body = {**_base_body(), "include_pdf": True, "return_signed_urls": True}
    with TestClient(app_module.app) as tc:
        resp = tc.post("/v1/generate-proposal", json=body)
    assert resp.status_code == 200
    data = resp.json()
    assert "error" in data["pdf"]
    assert data["signed_urls"] == {}
    assert "storage_notes" in data


if __name__ == "__main__":
    import pytest

    sys.exit(pytest.main([__file__, "-v"]))
