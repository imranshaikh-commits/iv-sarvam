"""
Pass 3 depth tests — run with NO API keys / NO network.

Covers:
  * depth-tier resolution (valid values, invalid/missing -> safe default)
  * per-call token caps never exceed the hard ceiling
  * full mode adds multi-subsection drafting + the appendix pack (RACI, timeline,
    sizing, integration inventory, risks)
  * standard/no-depth generate-proposal still succeeds with a mocked LLM
"""

import asyncio
import io
import os
import sys

_HERE = os.path.dirname(os.path.abspath(__file__))
_BRAIN = os.path.dirname(_HERE)
if _BRAIN not in sys.path:
    sys.path.insert(0, _BRAIN)

import httpx
from docx import Document

import document_engine
from document_engine import assemble_docx, generate_proposal
from proposal_templates import (
    DEFAULT_DEPTH,
    DEPTH_TIERS,
    VALID_DEPTHS,
    get_depth_tier,
)

CLIENT_NAME = "Meridian Bank"
APPENDIX_HEADINGS = [
    "Appendix A — RACI Matrix",
    "Appendix B — Indicative Timeline",
    "Appendix C — Sizing & Volumetrics",
    "Appendix D — Integration Inventory",
    "Appendix E — Risk Register",
]


# --- stubs ------------------------------------------------------------------
async def stub_embed(client, text):
    return [0.01] * 1536


def _fake_chunks():
    return [
        {
            "chunk_text": "SailPoint IdentityIQ deployed with automated joiner-mover-leaver "
            "provisioning across 42 target applications.",
            "heading": "Provisioning Architecture",
            "similarity": 0.62,
            "client_name": "Northwind Insurance",
            "iam_vendor": "sailpoint",
        },
        {
            "chunk_text": "Access certification campaigns ran quarterly with delegated reviewers.",
            "heading": "Access Certification",
            "similarity": 0.58,
            "client_name": "Acme Retail",
            "iam_vendor": "sailpoint",
        },
    ]


# **kw: draft_section passes proposal_type and section_topic. A stub that does
# not accept them raises TypeError, retrieval fails silently, and the section
# falls back to a placeholder while the test still passes. Third occurrence of
# this exact drift.
async def stub_retrieve(client, embedding, query, k=8, **kw):
    return _fake_chunks()


def stub_build_system(chunks):
    return "=== EVIDENCE ===\n" + "\n".join(
        f"[{i}] {c['chunk_text']}" for i, c in enumerate(chunks, 1)
    )


async def stub_draft(client, system_prompt, user_prompt, max_tokens=1500):
    # Assert the per-call cap is honoured and never exceeds the hard ceiling.
    assert max_tokens <= document_engine.MAX_DRAFT_TOKENS, max_tokens
    return (
        "Inspirit Vision proposes a SailPoint IdentityIQ deployment [1]. "
        "Quarterly access certification with delegated reviewers [1]."
    )


def _extract_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


async def _generate(depth):
    document_engine.draft_with_openrouter = stub_draft
    async with httpx.AsyncClient() as client:
        return await generate_proposal(
            client,
            rfp_text="Deliver an enterprise IAM implementation with automated provisioning.",
            client_name=CLIENT_NAME,
            proposal_type="implementation",
            iam_vendor="SailPoint",
            embed_fn=stub_embed,
            retrieve_fn=stub_retrieve,
            build_grounded_system_fn=stub_build_system,
            sections=None,
            include_compliance_matrix=False,
            proposal_depth=depth,
        )


# --- depth-tier resolution --------------------------------------------------
def test_valid_depth_values_resolve():
    assert VALID_DEPTHS == {"brief", "standard", "full", "deep"}
    for name in ("brief", "standard", "full", "deep"):
        assert get_depth_tier(name).name == name
        assert get_depth_tier(name.upper()).name == name  # case-insensitive


def test_invalid_and_missing_depth_fall_back_to_default():
    for bad in (None, "", "   ", "ultra", "extreme", "123", "STANDARDX"):
        assert get_depth_tier(bad).name == DEFAULT_DEPTH == "standard"


def test_per_call_token_caps_never_exceed_hard_ceiling():
    for tier in DEPTH_TIERS.values():
        assert tier.per_call_max_tokens <= document_engine.MAX_DRAFT_TOKENS
        assert tier.per_call_max_tokens > 0


def test_full_tier_plan_is_richer_than_standard():
    full = get_depth_tier("full")
    standard = get_depth_tier("standard")
    assert full.subsections_per_section > standard.subsections_per_section
    assert full.retrieval_fanout > standard.retrieval_fanout
    assert full.include_appendices and not standard.include_appendices


def test_deep_tier_plan_is_richer_than_full():
    deep = get_depth_tier("deep")
    full = get_depth_tier("full")
    assert deep.subsections_per_section > full.subsections_per_section
    assert deep.include_appendices and full.include_appendices
    # per-call cap must never exceed the hard ceiling even at the deepest tier.
    assert deep.per_call_max_tokens <= document_engine.MAX_DRAFT_TOKENS


def test_full_tier_unchanged_by_deep_tier_addition():
    # Regression guard: adding "deep" must not alter "full"'s existing plan —
    # existing callers/tests must see byte-identical behaviour.
    full = get_depth_tier("full")
    assert full.subsections_per_section == 3
    assert full.retrieval_fanout == 3
    # 3500 was set to chase page count; lowered to 2500 after the Amlak run came
    # out 38% longer than the human original while padding with repetition.
    assert full.per_call_max_tokens == 2500
    assert full.include_appendices is True


# --- generation: no-depth / standard still works ---------------------------
def test_no_depth_generate_proposal_succeeds():
    result = asyncio.run(_generate(None))
    assert result["docx_bytes"]
    assert result["proposal_depth"] == "standard"
    assert result["included_appendices"] is False
    text = _extract_text(result["docx_bytes"])
    assert CLIENT_NAME in text
    assert "Executive Summary" in text
    # No appendix pack at standard depth.
    for h in APPENDIX_HEADINGS:
        assert h not in text


def test_invalid_depth_generate_falls_back():
    result = asyncio.run(_generate("nonsense"))
    assert result["proposal_depth"] == "standard"
    assert result["included_appendices"] is False


# --- full mode: subsections + appendices -----------------------------------
def test_full_mode_adds_subsections_and_appendices():
    result = asyncio.run(_generate("full"))
    assert result["proposal_depth"] == "full"
    assert result["included_appendices"] is True
    text = _extract_text(result["docx_bytes"])

    # Multi-subsection drafting: the section's OWN subsection headings appear.
    # This used to assert the generic "Overview / Detailed Design /
    # Considerations & Dependencies" triple, which is exactly what Sprint B
    # removed -- the executive summary was still getting it in run 9 because a
    # section with no subsections falls back to SUBSECTION_FACETS.
    for expected in ("Proposed Production Hardware Sizing", "RACI Legend",
                     "Tranche 1 - Foundation"):
        assert expected in text, f"missing content-specific subsection: {expected}"
    assert "Considerations & Dependencies" not in text, \
        "the generic facet triple is back"

    # The appendix pack is now SUPPRESSED when the body template carries RACI,
    # timeline, sizing and commercial as real sections (2026-08-14). Run 5
    # printed each of those twice with different numbers -- a 25-row RACI built
    # from discovery answers in the body, and a 9-row generic placeholder in
    # Appendix A. The tier flag still reports True; what changed is that the
    # assembler skips a pack the body supersedes.
    # A/B/C/F are suppressed (body carries RACI, timeline, sizing, commercial).
    # D and E have no body counterpart and must survive.
    for h in APPENDIX_HEADINGS:
        superseded = any(k in h for k in ("RACI", "Timeline", "Sizing", "Commercial"))
        if superseded:
            assert h not in text, f"appendix {h!r} duplicates a body section"
        else:
            assert h in text, f"appendix {h!r} has no body counterpart and must remain"
    # The RACI legend lived in the (now suppressed) appendix pack; the body's
    # RACI table is drafted from discovery answers and carries its own header.
    assert "RACI" in text, "no RACI content anywhere in the document"
    # Was the appendix RACI's column header. The body RACI is drafted, so
    # assert the section exists rather than a hardcoded appendix cell.
    assert "RACI" in text, "no RACI content anywhere in the document"
    assert "Risk" in text
    assert "[ASSUMPTION]" in text  # conservative placeholders, not fabricated specifics


def test_full_mode_draft_markdown_has_subsections():
    result = asyncio.run(_generate("full"))
    # Was '### Overview' -- the generic facet triple. The executive summary is
    # continuous prose in IV's house style, so assert a real content-specific
    # heading instead.
    assert "### Proposed Production Hardware Sizing" in result["draft_markdown"]


# --- deep mode: all 6 facets + appendices -----------------------------------
def test_deep_mode_adds_all_facets_and_appendices():
    result = asyncio.run(_generate("deep"))
    assert result["proposal_depth"] == "deep"
    assert result["included_appendices"] is True
    text = _extract_text(result["docx_bytes"])

    # All 6 subsection facets appear (the 3 from "full" plus the 3 new ones).
    for facet_title in ("Proposed DR Hardware Sizing", "RACI - Delivery Activities",
                        "Payment Milestone - Implementation"):
        assert facet_title in text, f"missing subsection: {facet_title}"

    for h in APPENDIX_HEADINGS:
        superseded = any(k in h for k in ("RACI", "Timeline", "Sizing", "Commercial"))
        if superseded:
            assert h not in text, f"appendix {h!r} duplicates a body section"
        else:
            assert h in text, f"appendix {h!r} has no body counterpart and must remain"
    assert "[ASSUMPTION]" in text  # conservative placeholders, not fabricated specifics


def test_deep_mode_draft_markdown_has_all_subsections():
    result = asyncio.run(_generate("deep"))
    # Was '### Overview' -- the generic facet triple. The executive summary is
    # continuous prose in IV's house style, so assert a real content-specific
    # heading instead.
    assert "### Proposed Production Hardware Sizing" in result["draft_markdown"]
    # Was a generic facet name. Assert a real IV subsection instead.
    assert "### Payment Milestone - Licence" in result["draft_markdown"]
    assert "### \n" not in result["draft_markdown"], "empty markdown heading"


def test_appendices_render_directly():
    docx_bytes = assemble_docx(
        {"client_name": CLIENT_NAME, "proposal_type": "implementation", "iam_vendor": "SailPoint"},
        [{"id": "executive_summary", "title": "Executive Summary",
          "content": "Summary [1].", "citations": _fake_chunks(),
          "max_similarity": 0.62, "needs_sme_review": False}],
        include_appendices=True,
    )
    text = _extract_text(docx_bytes)
    for h in APPENDIX_HEADINGS:
        assert h in text


if __name__ == "__main__":
    test_valid_depth_values_resolve()
    test_invalid_and_missing_depth_fall_back_to_default()
    test_per_call_token_caps_never_exceed_hard_ceiling()
    test_full_tier_plan_is_richer_than_standard()
    test_deep_tier_plan_is_richer_than_full()
    test_full_tier_unchanged_by_deep_tier_addition()
    test_no_depth_generate_proposal_succeeds()
    test_invalid_depth_generate_falls_back()
    test_full_mode_adds_subsections_and_appendices()
    test_full_mode_draft_markdown_has_subsections()
    test_deep_mode_adds_all_facets_and_appendices()
    test_deep_mode_draft_markdown_has_all_subsections()
    test_appendices_render_directly()
    print("ALL PASS 3 CHECKS PASSED")
