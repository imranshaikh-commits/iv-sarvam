#!/usr/bin/env python3
"""
Sarvam — Sprint 1: Proposal Ingestion Pipeline

Processes raw DOCX/PDF proposals into structured JSON with:
  - Text extracted per logical section (Exec Summary, Company Profile, Scope, etc.)
  - Embedded images extracted and stored separately
  - Metadata stub ready for manual tagging

Usage:
    python scripts/ingest_proposals.py --input data/raw --output data/processed

Requirements:
    pip install python-docx pypdf pillow tqdm

Output structure per proposal:
    data/processed/
        <proposal_slug>/
            proposal.json          # Structured content + metadata
            images/                # Extracted images (renumbered)
                image_001.png
                image_002.png

Author: Sarvam Project
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

from tqdm import tqdm

# ------------------------------------------------------------------
# Section detection heuristics
# ------------------------------------------------------------------
# Based on analysis of 10 sample IV proposals (SailPoint, Ping, IBM,
# Keycloak, ForgeRock across Al Qadsiah, STC, ABB, DFCC, Mannai, etc.)
# These headings recur across almost every proposal.

SECTION_PATTERNS = {
    "cover_page": [
        r"^proposal for",
        r"^response to",
        r"^prepared for",
    ],
    "executive_summary": [
        r"executive summary",
        r"proposal summary",
    ],
    "company_profile": [
        r"company profile",
        r"about (inspirit vision|iv)",
        r"who we are",
    ],
    "similar_experience": [
        r"similar experience",
        r"our clients",
        r"case studies",
        r"reference clients",
    ],
    "understanding_of_scope": [
        r"understanding of (scope|requirements?)",
        r"scope of work",
        r"project scope",
    ],
    "solution_overview": [
        r"solution overview",
        r"proposed solution",
        r"our approach",
    ],
    "why_vendor": [
        r"why sailpoint",
        r"why ping",
        r"why (ibm|okta|saviynt|keycloak|forgerock)",
        r"vendor rationale",
    ],
    "solution_architecture": [
        r"solution architecture",
        r"technical architecture",
        r"high[- ]level design",
        r"hld",
        r"reference architecture",
    ],
    "implementation_approach": [
        r"implementation (approach|methodology)",
        r"delivery approach",
        r"our methodology",
    ],
    "raci_matrix": [
        r"raci( matrix)?",
        r"roles (and|&) responsibilities",
    ],
    "project_timeline": [
        r"project timeline",
        r"project plan",
        r"milestones?",
        r"schedule",
    ],
    "hardware_sizing": [
        r"hardware sizing",
        r"infrastructure sizing",
        r"sizing requirements",
    ],
    "sla_tiers": [
        r"sla( tiers)?",
        r"service level agreements?",
        r"support tiers",
    ],
    "governance": [
        r"governance( framework)?",
        r"quality management",
    ],
    "risk_management": [
        r"risk management",
        r"risk assessment",
    ],
    "change_management": [
        r"change management",
    ],
    "commercial": [
        r"commercials?",
        r"pricing",
        r"cost breakdown",
        r"investment",
        r"financial proposal",
    ],
    "terms_conditions": [
        r"terms (and|&) conditions",
        r"assumptions (and|&) exclusions",
    ],
}


# ------------------------------------------------------------------
# Data model
# ------------------------------------------------------------------

@dataclass
class ProposalSection:
    """One logical section of a proposal."""
    section_type: str          # e.g. "executive_summary" (matches SECTION_PATTERNS keys)
    heading: str               # Raw heading text as found in doc
    order: int                 # Position in original doc
    text: str                  # Body text
    word_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.text.split())


@dataclass
class ProposalMetadata:
    """Metadata stub — to be enriched via manual tagging CSV."""
    proposal_id: str           # Auto-generated slug from filename
    source_filename: str
    file_type: str             # "docx" or "pdf"
    total_word_count: int = 0
    image_count: int = 0
    # These fields are placeholders; filled in via tagging CSV in Sprint 2
    client_name: Optional[str] = None
    industry: Optional[str] = None
    country: Optional[str] = None
    iam_vendor: Optional[str] = None
    proposal_type: Optional[str] = None  # "implementation" | "mss"
    user_count: Optional[int] = None
    app_count: Optional[int] = None
    deal_size_bucket: Optional[str] = None  # "small" | "medium" | "large" | "enterprise"
    outcome: Optional[str] = None            # "won" | "lost" | "pending" | "unknown"
    year: Optional[int] = None


@dataclass
class Proposal:
    """A fully-processed proposal ready for embedding."""
    metadata: ProposalMetadata
    sections: list[ProposalSection] = field(default_factory=list)
    unmapped_headings: list[str] = field(default_factory=list)  # For QA


# ------------------------------------------------------------------
# Section classification
# ------------------------------------------------------------------

def classify_heading(heading: str) -> Optional[str]:
    """Return section_type key if heading matches any known pattern, else None."""
    heading_lower = heading.lower().strip()
    # Strip leading numbering like "1.", "1.1", "1.1.1", "Section 1:"
    heading_lower = re.sub(r"^(section\s+)?\d+(\.\d+)*[.:)]?\s*", "", heading_lower)

    for section_type, patterns in SECTION_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, heading_lower):
                return section_type
    return None


# ------------------------------------------------------------------
# DOCX processing
# ------------------------------------------------------------------

def is_heading_paragraph(para) -> bool:
    """Detect if a python-docx paragraph is a heading."""
    if para.style and para.style.name:
        style_name = para.style.name.lower()
        if "heading" in style_name or "title" in style_name:
            return True
    # Fallback heuristic: short, all-caps or bold, non-empty
    text = para.text.strip()
    if not text or len(text) > 120:
        return False
    if text.isupper() and len(text.split()) <= 12:
        return True
    return False


def process_docx(file_path: Path, images_dir: Path) -> Proposal:
    """Extract sections and images from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(file_path))

    proposal_id = slugify(file_path.stem)
    metadata = ProposalMetadata(
        proposal_id=proposal_id,
        source_filename=file_path.name,
        file_type="docx",
    )
    proposal = Proposal(metadata=metadata)

    current_section: Optional[ProposalSection] = None
    current_buffer: list[str] = []
    order = 0

    def flush_section():
        """Save the currently-buffered section, if any."""
        nonlocal current_section
        if current_section is not None:
            current_section.text = "\n".join(current_buffer).strip()
            current_section.word_count = len(current_section.text.split())
            # Only keep sections with actual content
            if current_section.word_count > 5:
                proposal.sections.append(current_section)
        current_buffer.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if is_heading_paragraph(para):
            flush_section()
            section_type = classify_heading(text)
            if section_type is None:
                proposal.unmapped_headings.append(text)
                # Still create a section — mark as "other"
                section_type = "other"
            order += 1
            current_section = ProposalSection(
                section_type=section_type,
                heading=text,
                order=order,
                text="",
            )
        else:
            current_buffer.append(text)

    flush_section()

    # Also process tables — flatten into text under the last-known section
    for table in doc.tables:
        table_text = extract_table_text(table)
        if table_text and proposal.sections:
            proposal.sections[-1].text += f"\n\n[TABLE]\n{table_text}\n[/TABLE]"

    # Extract embedded images
    image_count = extract_docx_images(file_path, images_dir)
    metadata.image_count = image_count

    # Aggregate word count
    metadata.total_word_count = sum(s.word_count for s in proposal.sections)

    return proposal


def extract_table_text(table) -> str:
    """Flatten a docx table into pipe-delimited text."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_docx_images(docx_path: Path, output_dir: Path) -> int:
    """Extract all embedded images from a docx (which is a zip)."""
    import zipfile

    output_dir.mkdir(parents=True, exist_ok=True)
    count = 0

    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                count += 1
                ext = Path(name).suffix or ".bin"
                target = output_dir / f"image_{count:03d}{ext}"
                with z.open(name) as src, open(target, "wb") as dst:
                    dst.write(src.read())
    return count


# ------------------------------------------------------------------
# PDF processing
# ------------------------------------------------------------------

def process_pdf(file_path: Path, images_dir: Path) -> Proposal:
    """Extract sections and images from a PDF file.

    Note: PDF section detection is heuristic and less reliable than DOCX.
    We rely on font-size changes and common heading patterns.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf not installed. Run: pip install pypdf")

    reader = PdfReader(str(file_path))

    proposal_id = slugify(file_path.stem)
    metadata = ProposalMetadata(
        proposal_id=proposal_id,
        source_filename=file_path.name,
        file_type="pdf",
    )
    proposal = Proposal(metadata=metadata)

    # Simple approach: extract all text, then use regex to find section headings
    full_text = ""
    for page in reader.pages:
        full_text += page.extract_text() + "\n"

    # Split by likely heading patterns (lines that look like headings)
    lines = full_text.split("\n")
    current_section: Optional[ProposalSection] = None
    current_buffer: list[str] = []
    order = 0

    def flush_section():
        nonlocal current_section
        if current_section is not None:
            current_section.text = "\n".join(current_buffer).strip()
            current_section.word_count = len(current_section.text.split())
            if current_section.word_count > 5:
                proposal.sections.append(current_section)
        current_buffer.clear()

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # PDF heuristic: line is short + matches known section pattern
        if len(line_stripped) < 100:
            section_type = classify_heading(line_stripped)
            if section_type:
                flush_section()
                order += 1
                current_section = ProposalSection(
                    section_type=section_type,
                    heading=line_stripped,
                    order=order,
                    text="",
                )
                continue

        current_buffer.append(line_stripped)

    flush_section()

    metadata.total_word_count = sum(s.word_count for s in proposal.sections)
    # PDF image extraction is complex; skipping for MVP
    # If needed, use `pdfplumber` or `pymupdf` for image extraction

    return proposal


# ------------------------------------------------------------------
# Utilities
# ------------------------------------------------------------------

def slugify(text: str) -> str:
    """Convert filename to safe slug."""
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"-+", "-", text)
    return text.strip("-")


# ------------------------------------------------------------------
# Main
# ------------------------------------------------------------------

def process_file(file_path: Path, output_root: Path) -> Optional[Proposal]:
    """Process one file end-to-end."""
    suffix = file_path.suffix.lower()
    proposal_id = slugify(file_path.stem)
    output_dir = output_root / proposal_id
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"

    try:
        if suffix == ".docx":
            proposal = process_docx(file_path, images_dir)
        elif suffix == ".pdf":
            proposal = process_pdf(file_path, images_dir)
        else:
            logging.warning(f"Skipping unsupported file type: {file_path}")
            return None
    except Exception as e:
        logging.error(f"Failed to process {file_path.name}: {e}")
        return None

    # Write proposal.json
    output_json = output_dir / "proposal.json"
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump({
            "metadata": asdict(proposal.metadata),
            "sections": [asdict(s) for s in proposal.sections],
            "unmapped_headings": proposal.unmapped_headings,
        }, f, indent=2, ensure_ascii=False)

    return proposal


def main():
    parser = argparse.ArgumentParser(description="Sarvam proposal ingestion pipeline")
    parser.add_argument("--input", type=Path, required=True, help="Directory of raw DOCX/PDF proposals")
    parser.add_argument("--output", type=Path, required=True, help="Output directory for processed JSONs")
    parser.add_argument("--verbose", action="store_true", help="Verbose logging")
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )

    if not args.input.exists():
        logging.error(f"Input directory does not exist: {args.input}")
        sys.exit(1)

    args.output.mkdir(parents=True, exist_ok=True)

    # Collect all DOCX and PDF files (recursive)
    files = list(args.input.rglob("*.docx")) + list(args.input.rglob("*.pdf"))
    files = [f for f in files if not f.name.startswith("~$")]  # Skip Word lock files

    if not files:
        logging.error(f"No .docx or .pdf files found in {args.input}")
        sys.exit(1)

    logging.info(f"Found {len(files)} proposal(s) to process")

    stats = {"processed": 0, "failed": 0, "total_sections": 0, "unmapped_headings": 0}

    for file_path in tqdm(files, desc="Ingesting proposals"):
        result = process_file(file_path, args.output)
        if result is None:
            stats["failed"] += 1
        else:
            stats["processed"] += 1
            stats["total_sections"] += len(result.sections)
            stats["unmapped_headings"] += len(result.unmapped_headings)

    # Write summary
    summary_path = args.output / "_ingestion_summary.json"
    with open(summary_path, "w") as f:
        json.dump(stats, f, indent=2)

    print("\n" + "=" * 60)
    print("INGESTION COMPLETE")
    print("=" * 60)
    print(f"  Processed: {stats['processed']}")
    print(f"  Failed:    {stats['failed']}")
    print(f"  Total sections extracted: {stats['total_sections']}")
    print(f"  Unmapped headings (need QA): {stats['unmapped_headings']}")
    print(f"\n  Output: {args.output}")
    print(f"  Summary: {summary_path}")
    print("\n  Next steps:")
    print("    1. Review a few processed JSONs manually — spot-check section splits.")
    print("    2. Fill in metadata via data/tagging/tagging_template.csv.")
    print("    3. Run scripts/embed_and_upload.py (Sprint 2) once Supabase is live.")


if __name__ == "__main__":
    main()
