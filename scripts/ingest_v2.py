#!/usr/bin/env python3
"""
Shilpi Proposal Ingestion Pipeline v2

For each proposal (DOCX or PDF):
  1. Extract text + tables (as Markdown)
  2. Extract embedded images
  3. OCR text-heavy images with Tesseract
  4. Describe diagram-like images with Qwen 3 VL 8B (via OpenRouter)
  5. Auto-extract metadata (client, industry, vendor, etc.) with DeepSeek V3.2
  6. Chunk semantically (target ~500 tokens, 50-token overlap)
  7. Embed each chunk with openai/text-embedding-3-small (via OpenRouter)
  8. Write to Supabase (proposals + proposal_chunks tables)

Config via env vars:
  OPENROUTER_API_KEY  — required
  SUPABASE_URL        — required
  SUPABASE_KEY        — required (service_role for direct write, or anon if using RPC)
  ORG_ID              — required (uuid of Inspirit Vision org)
  INPUT_DIR           — directory of proposals (default: ./proposals)
  OUTPUT_DIR          — where to save extracted/staged JSON (default: ./out)
  LIMIT               — optional int, process only first N files

Usage:
  python ingest_v2.py --input /path/to/proposals --limit 1  # test mode
  python ingest_v2.py --input /path/to/proposals            # full batch
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import io
import json
import logging
import os
import re
import sys
import time
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

import requests
from PIL import Image

# ---------------------------------------------------------------------------
# Config & logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("ingest")

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_BASE = "https://openrouter.ai/api/v1"
SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "")
ORG_ID = os.getenv("ORG_ID", "")

# Models
EMBED_MODEL = "openai/text-embedding-3-small"
METADATA_MODEL = "deepseek/deepseek-v3.2-exp"
VISION_MODEL = "qwen/qwen3-vl-8b-instruct"  # Cheap vision model on OpenRouter

# Chunking
CHUNK_TARGET_WORDS = 350  # ~500 tokens
CHUNK_OVERLAP_WORDS = 40

# Image size classifier — small images = likely icons; medium = likely text-heavy; large = likely diagrams
MIN_IMAGE_BYTES = 5_000  # skip tiny icons
DIAGRAM_MIN_DIM = 400  # px — larger dim likely means diagram/screenshot with signal


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class Section:
    heading: str
    section_type: str  # e.g. exec_summary, scope, similar_experience, tables, ocr
    section_order: int
    text: str

@dataclass
class Proposal:
    slug: str
    source_filename: str
    file_type: str  # docx | pdf
    total_word_count: int = 0
    image_count: int = 0
    sections: list[Section] = field(default_factory=list)
    # Metadata (LLM-extracted)
    client_name: Optional[str] = None
    industry: Optional[str] = None
    country: Optional[str] = None
    iam_vendor: Optional[str] = None
    proposal_type: Optional[str] = None
    user_count: Optional[int] = None
    app_count: Optional[int] = None
    deal_size_bucket: Optional[str] = None
    year: Optional[int] = None
    notes: Optional[str] = None
    # Provenance from the reviewed manifest (see sarvam_006 migration).
    source_sha256: Optional[str] = None
    source_tier: Optional[str] = None


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_from_docx(path: Path) -> tuple[list[Section], list[bytes]]:
    """Return (sections, images_as_bytes)."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    sections: list[Section] = []
    images: list[bytes] = []

    current_heading = "Introduction"
    current_type = "intro"
    current_text_parts: list[str] = []
    section_order = 0

    def flush_section():
        nonlocal section_order
        if not current_text_parts:
            return
        text = "\n\n".join(current_text_parts).strip()
        if text and len(text.split()) > 5:  # skip tiny fragments
            sections.append(Section(
                heading=current_heading,
                section_type=classify_section(current_heading),
                section_order=section_order,
                text=text,
            ))
            section_order += 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name if para.style else "").lower()
        if "heading" in style_name and len(text) < 200:
            flush_section()
            current_heading = text
            current_type = classify_section(text)
            current_text_parts = []
        else:
            current_text_parts.append(text)

    # Extract tables as markdown
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        if not rows:
            continue
        # Skip trivial tables
        if len(rows) < 2 or all(len(r) < 2 for r in rows):
            continue
        md = "| " + " | ".join(rows[0]) + " |\n"
        md += "|" + "|".join(["---"] * len(rows[0])) + "|\n"
        for r in rows[1:]:
            # Pad row to header width
            padded = r + [""] * (len(rows[0]) - len(r))
            md += "| " + " | ".join(padded[:len(rows[0])]) + " |\n"
        sections.append(Section(
            heading=f"Table {i+1}",
            section_type="table",
            section_order=section_order,
            text=md,
        ))
        section_order += 1

    flush_section()

    # Extract images from document parts
    for rel_id, rel in doc.part.related_parts.items():
        if "image" in rel.content_type:
            try:
                images.append(rel.blob)
            except Exception as e:
                log.debug("Skipped image %s: %s", rel_id, e)

    return sections, images


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def extract_from_pdf(path: Path) -> tuple[list[Section], list[bytes]]:
    """Return (sections, images_as_bytes)."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    sections: list[Section] = []
    images: list[bytes] = []

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = text.strip()
        if text and len(text.split()) > 5:
            sections.append(Section(
                heading=f"Page {i+1}",
                section_type="page",
                section_order=i,
                text=text,
            ))
        # Extract images from page
        try:
            for img in page.images:
                images.append(img.data)
        except Exception as e:
            log.debug("Skipped images on page %d: %s", i + 1, e)

    return sections, images


# ---------------------------------------------------------------------------
# Section classification
# ---------------------------------------------------------------------------

SECTION_PATTERNS = [
    ("cover", [r"^proposal for", r"^response to", r"^prepared for", r"^cover"]),
    ("exec_summary", [r"executive summary", r"proposal summary"]),
    ("company_profile", [r"company profile", r"about (inspirit|iv)", r"who we are"]),
    ("similar_experience", [r"similar experience", r"case studies", r"reference clients", r"our clients"]),
    ("scope", [r"understanding of scope", r"scope of work", r"project scope", r"understanding of requirements"]),
    ("solution", [r"solution overview", r"proposed solution", r"our approach", r"solution architecture"]),
    ("architecture", [r"architecture", r"deployment (diagram|topology)"]),
    ("timeline", [r"timeline", r"schedule", r"project plan", r"milestones"]),
    ("pricing", [r"pricing", r"commercial", r"cost", r"investment"]),
    ("team", [r"team", r"resources", r"consultants"]),
    ("why_vendor", [r"why (sailpoint|ping|forgerock|keycloak|okta|redhat|ibm)", r"vendor advantages"]),
    ("assumptions", [r"assumptions", r"exclusions", r"risks"]),
]

def classify_section(heading: str) -> str:
    h = heading.lower()
    for section_type, patterns in SECTION_PATTERNS:
        for p in patterns:
            if re.search(p, h):
                return section_type
    return "other"


# ---------------------------------------------------------------------------
# Image OCR / vision
# ---------------------------------------------------------------------------

def is_diagram_like(img_bytes: bytes) -> bool:
    """Heuristic: diagrams are larger, have moderate aspect ratios, and few unique colors."""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        w, h = img.size
        if w < DIAGRAM_MIN_DIM and h < DIAGRAM_MIN_DIM:
            return False  # too small — likely icon/logo
        # Very wide (banner) or very tall = probably not diagram
        aspect = max(w, h) / min(w, h)
        if aspect > 4:
            return False
        return True
    except Exception:
        return False


def ocr_tesseract(img_bytes: bytes) -> str:
    """Run Tesseract on an image, return extracted text."""
    import subprocess
    try:
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode != "RGB":
            img = img.convert("RGB")
        # Resize very large images down for speed
        if max(img.size) > 2000:
            img.thumbnail((2000, 2000))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        result = subprocess.run(
            ["tesseract", "-", "-", "-l", "eng", "--psm", "6"],
            input=buf.getvalue(),
            capture_output=True,
            timeout=30,
        )
        text = result.stdout.decode("utf-8", errors="ignore").strip()
        # Filter noise: keep only if we got meaningful content
        if len(text) < 30 or len(text.split()) < 5:
            return ""
        return text
    except Exception as e:
        log.debug("Tesseract failed: %s", e)
        return ""


def describe_diagram_qwen(img_bytes: bytes) -> str:
    """Use Qwen 3 VL to describe a diagram/architecture image."""
    try:
        b64 = base64.b64encode(img_bytes).decode("ascii")
        # Detect mime type
        img = Image.open(io.BytesIO(img_bytes))
        fmt = (img.format or "PNG").lower()
        if fmt == "jpeg":
            fmt = "jpeg"
        elif fmt not in ("png", "gif", "webp"):
            fmt = "png"
            # Re-encode to PNG
            buf = io.BytesIO()
            img.convert("RGB").save(buf, format="PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        resp = requests.post(
            f"{OPENROUTER_BASE}/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": VISION_MODEL,
                "messages": [
                    {"role": "user", "content": [
                        {"type": "text", "text": (
                            "This image is from a technical IAM/cybersecurity proposal. "
                            "Describe what it shows in 2-4 sentences. Focus on: "
                            "the type of diagram (architecture, workflow, timeline, screenshot, chart), "
                            "any labeled components/products/vendors, and the relationships shown. "
                            "If it appears to be a decorative image, logo, or has no meaningful content, respond only: SKIP"
                        )},
                        {"type": "image_url", "image_url": {"url": f"data:image/{fmt};base64,{b64}"}},
                    ]},
                ],
                "max_tokens": 200,
            },
            timeout=45,
        )
        if resp.status_code != 200:
            log.warning("Vision API returned %d: %s", resp.status_code, resp.text[:200])
            return ""
        content = resp.json()["choices"][0]["message"]["content"].strip()
        if "SKIP" in content.upper()[:20]:
            return ""
        return content
    except Exception as e:
        log.debug("Vision describe failed: %s", e)
        return ""


def process_images(images: list[bytes]) -> list[Section]:
    """OCR text-heavy images, vision-describe diagram-like ones."""
    sections: list[Section] = []
    ocr_count, vision_count, skipped = 0, 0, 0
    for i, img_bytes in enumerate(images):
        if len(img_bytes) < MIN_IMAGE_BYTES:
            skipped += 1
            continue
        # Try Tesseract first — cheap and fast
        ocr_text = ocr_tesseract(img_bytes)
        if ocr_text and len(ocr_text.split()) >= 15:
            # Got meaningful text — likely a text-heavy image
            sections.append(Section(
                heading=f"Image OCR #{i+1}",
                section_type="ocr",
                section_order=1000 + i,
                text=f"[OCR from embedded image #{i+1}]\n\n{ocr_text}",
            ))
            ocr_count += 1
            continue
        # No text? Might be a diagram — use vision model
        if is_diagram_like(img_bytes):
            desc = describe_diagram_qwen(img_bytes)
            if desc:
                sections.append(Section(
                    heading=f"Diagram #{i+1}",
                    section_type="diagram",
                    section_order=2000 + i,
                    text=f"[Diagram description from embedded image #{i+1}]\n\n{desc}",
                ))
                vision_count += 1
            else:
                skipped += 1
        else:
            skipped += 1
    log.info("  Images: %d OCR'd, %d described, %d skipped", ocr_count, vision_count, skipped)
    return sections


# ---------------------------------------------------------------------------
# LLM metadata extraction
# ---------------------------------------------------------------------------

def extract_metadata(proposal: Proposal, first_pages_text: str) -> None:
    """Use DeepSeek V3.2 to extract structured metadata from the first few pages."""
    prompt = f"""Analyze this excerpt from an Inspirit Vision (IV) technical proposal document and extract structured metadata.

DOCUMENT EXCERPT:
{first_pages_text[:6000]}

Return ONLY valid JSON with these fields (use null if not found):
{{
  "client_name": "The client company name (e.g. 'STC', 'DFCC Bank')",
  "industry": "Industry sector (banking / telecom / energy / water / retail / government / sports / manufacturing / other)",
  "country": "ISO country name or region",
  "iam_vendor": "IAM technology (Ping / SailPoint / Keycloak / ForgeRock / Okta / IBM / RedHat / Microsoft / other)",
  "proposal_type": "implementation OR mss",
  "user_count": integer or null,
  "app_count": integer or null,
  "deal_size_bucket": "small OR medium OR large OR enterprise",
  "year": integer or null (proposal year),
  "notes": "One-sentence summary of the deal"
}}"""

    try:
        resp = requests.post(
            f"{OPENROUTER_BASE}/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": METADATA_MODEL,
                "messages": [
                    {"role": "system", "content": "You are a precise data extraction tool. Return ONLY valid JSON — no markdown fences, no commentary."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.1,
                "max_tokens": 500,
            },
            timeout=60,
        )
        if resp.status_code != 200:
            log.warning("Metadata extraction failed HTTP %d", resp.status_code)
            return
        content = resp.json()["choices"][0]["message"]["content"].strip()
        # Strip code fences if present
        content = re.sub(r"^```(?:json)?\n?", "", content)
        content = re.sub(r"\n?```$", "", content)
        data = json.loads(content)
        for k, v in data.items():
            if hasattr(proposal, k) and v is not None:
                setattr(proposal, k, v)
        log.info("  Metadata: client=%s industry=%s vendor=%s", 
                 proposal.client_name, proposal.industry, proposal.iam_vendor)
    except Exception as e:
        log.warning("Metadata extraction error: %s", e)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def chunk_section(section: Section, target_words: int = CHUNK_TARGET_WORDS,
                   overlap: int = CHUNK_OVERLAP_WORDS) -> list[Section]:
    """Split a section into chunks of ~target_words with overlap. Returns list of Section-like chunks."""
    words = section.text.split()
    if len(words) <= target_words:
        return [section]

    chunks: list[Section] = []
    start = 0
    part = 0
    while start < len(words):
        end = min(start + target_words, len(words))
        chunk_words = words[start:end]
        chunks.append(Section(
            heading=f"{section.heading} (part {part+1})",
            section_type=section.section_type,
            section_order=section.section_order + part,
            text=" ".join(chunk_words),
        ))
        if end == len(words):
            break
        start = end - overlap
        part += 1
    return chunks


# ---------------------------------------------------------------------------
# Embedding
# ---------------------------------------------------------------------------

def embed_texts(texts: list[str], batch_size: int = 32) -> list[list[float]]:
    """Batch-embed texts via OpenRouter."""
    all_embeddings: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        for attempt in range(3):
            try:
                resp = requests.post(
                    f"{OPENROUTER_BASE}/embeddings",
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "Content-Type": "application/json",
                    },
                    json={"model": EMBED_MODEL, "input": batch},
                    timeout=60,
                )
                if resp.status_code == 200:
                    data = resp.json()["data"]
                    all_embeddings.extend([item["embedding"] for item in data])
                    break
                else:
                    log.warning("Embed batch %d HTTP %d (attempt %d): %s",
                                i // batch_size, resp.status_code, attempt + 1, resp.text[:200])
                    time.sleep(2 * (attempt + 1))
            except Exception as e:
                log.warning("Embed batch %d error (attempt %d): %s", i // batch_size, attempt + 1, e)
                time.sleep(2 * (attempt + 1))
        else:
            raise RuntimeError(f"Failed to embed batch {i // batch_size} after 3 attempts")
    return all_embeddings


# ---------------------------------------------------------------------------
# Supabase writes
# ---------------------------------------------------------------------------

def sb_headers() -> dict:
    return {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }


def sb_insert_proposal(proposal: Proposal) -> str:
    """Insert into public.proposals and return the new UUID."""
    payload = {
        "org_id": ORG_ID,
        "proposal_slug": proposal.slug,
        "source_filename": proposal.source_filename,
        "file_type": proposal.file_type,
        "total_word_count": proposal.total_word_count,
        "image_count": proposal.image_count,
        "client_name": proposal.client_name,
        "industry": proposal.industry,
        "country": proposal.country,
        "iam_vendor": proposal.iam_vendor,
        "proposal_type": proposal.proposal_type,
        "user_count": proposal.user_count,
        "app_count": proposal.app_count,
        "deal_size_bucket": proposal.deal_size_bucket,
        "outcome": "unknown",
        "year": proposal.year,
        "notes": proposal.notes,
        "source_sha256": proposal.source_sha256,
        "source_tier": proposal.source_tier,
    }
    resp = requests.post(
        f"{SUPABASE_URL}/rest/v1/proposals",
        headers=sb_headers(),
        json=payload,
        timeout=30,
    )
    if resp.status_code not in (200, 201):
        raise RuntimeError(f"Insert proposal failed HTTP {resp.status_code}: {resp.text}")
    return resp.json()[0]["id"]


def sb_insert_chunks(proposal_id: str, chunks_with_embeddings: list[tuple[Section, list[float]]]) -> int:
    """Batch-insert chunks. Returns count inserted."""
    rows = []
    for section, emb in chunks_with_embeddings:
        rows.append({
            "proposal_id": proposal_id,
            "org_id": ORG_ID,
            "section_type": section.section_type,
            "heading": section.heading[:500],  # cap for safety
            "section_order": section.section_order,
            "text": section.text,
            "word_count": len(section.text.split()),
            "embedding": emb,
        })
    # Batch in groups of 50
    inserted = 0
    for i in range(0, len(rows), 50):
        batch = rows[i:i + 50]
        resp = requests.post(
            f"{SUPABASE_URL}/rest/v1/proposal_chunks",
            headers=sb_headers(),
            json=batch,
            timeout=60,
        )
        if resp.status_code not in (200, 201):
            raise RuntimeError(f"Insert chunks failed HTTP {resp.status_code}: {resp.text[:500]}")
        inserted += len(batch)
    return inserted


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def slugify(name: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "-", name.lower()).strip("-")
    return s[:60]


# ---------------------------------------------------------------------------
# Manifest-driven selection
#
# Ingestion used to be a directory glob. That is unsafe now: the Sales-SoWs bank
# is a sales working folder, and a content review of its 421 files moved 68 of
# 197 candidates OUT of the ingest tier -- among them the STC RFP (client-
# authored, doc properties name Saudi Telecom Company), a competitor's proposal
# authored by Smpl ID whose corporate boilerplate would have landed in IV's
# voice bank, 22 consultant CVs, 2 NDAs and four vendor marketing PDFs.
#
# A glob discards that review silently. So selection now comes from the reviewed
# manifest CSV and nothing else, and the metadata on each row -- verified by
# reading the documents -- WINS over anything the LLM infers.
# ---------------------------------------------------------------------------

MANIFEST_REQUIRED_COLUMNS = {"path", "filename", "tier", "sha256_16"}


@dataclass
class ManifestRow:
    path: str
    filename: str
    tier: str
    sha256_16: str
    iam_vendor: str = ""
    client_name: str = ""
    proposal_type: str = ""
    year: str = ""
    reason: str = ""


def load_manifest(csv_path: Path, root: Path, tier: str = "ingest",
                  vendor: Optional[str] = None) -> list[tuple[Path, ManifestRow]]:
    """Rows of the requested tier, paired with the file each refers to.

    A row whose file is missing on disk is reported and skipped rather than
    silently dropped: a manifest that has drifted from the folder is a problem
    to surface, not to route around.
    """
    import csv as _csv

    with open(csv_path, newline="", encoding="utf-8-sig") as fh:
        reader = _csv.DictReader(fh)
        missing_cols = MANIFEST_REQUIRED_COLUMNS - set(reader.fieldnames or [])
        if missing_cols:
            raise SystemExit(
                f"manifest {csv_path} is missing required columns: "
                f"{', '.join(sorted(missing_cols))}")
        raw = list(reader)

    selected: list[tuple[Path, ManifestRow]] = []
    absent: list[str] = []
    skipped_tier = 0
    for r in raw:
        if (r.get("tier") or "").strip() != tier:
            skipped_tier += 1
            continue
        row = ManifestRow(
            path=r["path"], filename=r["filename"], tier=r["tier"],
            sha256_16=r["sha256_16"],
            iam_vendor=(r.get("iam_vendor") or "").strip(),
            client_name=(r.get("client_name") or "").strip(),
            proposal_type=(r.get("proposal_type") or "").strip(),
            year=(r.get("year") or "").strip(),
            reason=(r.get("reason") or "").strip(),
        )
        if vendor and vendor.lower() not in row.iam_vendor.lower():
            continue
        full = root / row.path
        if not full.exists():
            absent.append(row.path)
            continue
        selected.append((full, row))

    log.info("manifest: %d rows total, %d not tier=%s, %d selected%s",
             len(raw), skipped_tier, tier, len(selected),
             f" (vendor filter: {vendor})" if vendor else "")
    if absent:
        log.warning("manifest references %d file(s) not found on disk; first 5: %s",
                    len(absent), ", ".join(absent[:5]))
    return selected


def already_ingested(hashes: list[str]) -> set[str]:
    """Content hashes already present in `proposals` for this org.

    Without this, re-running ingestion inserts a second copy of every proposal
    and a second full set of embeddings. Retrieval then silently favours the
    duplicated document.
    """
    if not hashes:
        return set()
    seen: set[str] = set()
    for i in range(0, len(hashes), 50):
        batch = hashes[i:i + 50]
        quoted = ",".join(f'"{h}"' for h in batch)
        resp = requests.get(
            f"{SUPABASE_URL}/rest/v1/proposals",
            headers=sb_headers(),
            params={"select": "source_sha256", "org_id": f"eq.{ORG_ID}",
                    "source_sha256": f"in.({quoted})"},
            timeout=30,
        )
        if resp.status_code != 200:
            raise RuntimeError(
                f"dedup check failed HTTP {resp.status_code}: {resp.text}. "
                "Refusing to continue -- ingesting without it risks duplicates.")
        seen.update(r["source_sha256"] for r in resp.json() if r.get("source_sha256"))
    return seen


def apply_manifest_metadata(proposal: "Proposal", row: ManifestRow) -> list[str]:
    """Overlay verified manifest metadata onto LLM-inferred values.

    The manifest WINS wherever it has a value. The LLM reads the first ~6,000
    characters, which is exactly where a partner's name sits on a cover page --
    that is how "CIAM_Mannai_IV_Technical_Proposal_V3_0.docx" reads as a Mannai
    proposal when the client is Ahlibank. The same review also corrected
    Netpolean -> ABB India and PNB -> PNB MetLife. Those corrections came from
    reading whole documents and must not be re-guessed here.

    Returns a list of human-readable overrides for the run log.
    """
    overrides: list[str] = []
    for field_name, value in (("client_name", row.client_name),
                              ("iam_vendor", row.iam_vendor),
                              ("proposal_type", row.proposal_type)):
        if not value or value.lower() in ("n/a", "none", "unknown"):
            continue
        current = getattr(proposal, field_name, None)
        if current and str(current).strip().lower() != value.strip().lower():
            overrides.append(f"{field_name}: {current!r} -> {value!r}")
        setattr(proposal, field_name, value)
    if row.year.isdigit():
        if proposal.year and int(row.year) != proposal.year:
            overrides.append(f"year: {proposal.year} -> {row.year}")
        proposal.year = int(row.year)
    return overrides


def process_file(path: Path, output_dir: Path, skip_ocr: bool = False,
                 skip_vision: bool = False, row: Optional[ManifestRow] = None,
                 dry_run: bool = False) -> Optional[Proposal]:
    log.info("═" * 70)
    log.info("Processing: %s (%.1f MB)", path.name, path.stat().st_size / 1_048_576)
    t0 = time.time()

    suffix = path.suffix.lower()
    if suffix == ".docx":
        sections, images = extract_from_docx(path)
        file_type = "docx"
    elif suffix == ".pdf":
        sections, images = extract_from_pdf(path)
        file_type = "pdf"
    else:
        log.warning("Skipping unsupported file: %s", path.name)
        return None

    log.info("  Extracted %d sections, %d images in %.1fs",
             len(sections), len(images), time.time() - t0)

    # Save raw extraction to disk first (in case OCR/upload fails later)
    proposal = Proposal(
        slug=slugify(path.stem),
        source_filename=path.name,
        file_type=file_type,
        sections=sections,
        image_count=len(images),
    )
    proposal.total_word_count = sum(len(s.text.split()) for s in sections)

    # Save intermediate raw text (for debugging)
    stage_dir = output_dir / proposal.slug
    stage_dir.mkdir(parents=True, exist_ok=True)
    with open(stage_dir / "raw_extract.json", "w", encoding="utf-8") as f:
        json.dump({
            "source_filename": proposal.source_filename,
            "file_type": proposal.file_type,
            "total_word_count": proposal.total_word_count,
            "image_count": proposal.image_count,
            "sections": [asdict(s) for s in sections],
        }, f, indent=2, ensure_ascii=False)

    # OCR + vision on images
    if not (skip_ocr and skip_vision) and images:
        log.info("  Processing %d images (Tesseract + Qwen VL hybrid)...", len(images))
        t_img = time.time()
        image_sections = process_images(images)
        proposal.sections.extend(image_sections)
        log.info("  Image processing took %.1fs, added %d sections", 
                 time.time() - t_img, len(image_sections))

    # Metadata extraction. The LLM fills in what the manifest does not carry
    # (industry, country, volumetrics, deal size); the manifest then overrides
    # the fields a human verified by reading the document.
    log.info("  Extracting metadata with DeepSeek...")
    intro_text = "\n\n".join(s.text for s in proposal.sections[:5])
    extract_metadata(proposal, intro_text)
    if row is not None:
        proposal.source_sha256 = row.sha256_16
        proposal.source_tier = row.tier
        overrides = apply_manifest_metadata(proposal, row)
        for o in overrides:
            log.info("  manifest override -> %s", o)
        if not overrides:
            log.info("  manifest metadata agreed with the model")

    # Chunking
    all_chunks: list[Section] = []
    for section in proposal.sections:
        all_chunks.extend(chunk_section(section))
    log.info("  Chunked into %d pieces", len(all_chunks))

    # Save staged JSON (before embedding, in case embedding fails)
    with open(stage_dir / "staged.json", "w", encoding="utf-8") as f:
        json.dump({
            "proposal": {k: v for k, v in asdict(proposal).items() if k != "sections"},
            "chunks": [asdict(c) for c in all_chunks],
        }, f, indent=2, ensure_ascii=False)

    # Embed. Skipped under --dry-run: embedding is a paid call and a dry run
    # exists to check SELECTION and METADATA, not to spend money proving the
    # embedding endpoint still works.
    if dry_run:
        embeddings = []
    else:
        log.info("  Embedding %d chunks via OpenRouter...", len(all_chunks))
        t_emb = time.time()
        texts = [c.text for c in all_chunks]
        embeddings = embed_texts(texts)
        log.info("  Embedded in %.1fs", time.time() - t_emb)

    if dry_run:
        log.info("  DRY RUN: %d chunks prepared, nothing written to Supabase "
                 "(client=%s vendor=%s type=%s year=%s)",
                 len(all_chunks), proposal.client_name, proposal.iam_vendor,
                 proposal.proposal_type, proposal.year)
        return proposal

    # Write to Supabase
    log.info("  Writing to Supabase...")
    proposal_uuid = sb_insert_proposal(proposal)
    log.info("  proposal_id = %s", proposal_uuid)
    inserted = sb_insert_chunks(proposal_uuid, list(zip(all_chunks, embeddings)))
    log.info("  ✅ Inserted %d chunks", inserted)

    total_time = time.time() - t0
    log.info("  Total: %.1fs (%.1f min)", total_time, total_time / 60)
    return proposal


def main():
    ap = argparse.ArgumentParser(
        description="Ingest IV proposals into the Shilpi RAG corpus. Selection "
                    "comes from a reviewed manifest CSV, never from a directory "
                    "walk.")
    ap.add_argument("--manifest", help="Reviewed manifest CSV (corpus_manifest_reviewed.csv)")
    ap.add_argument("--root", help="Folder the manifest paths are relative to (e.g. ~/Downloads/Sales-SoWs)")
    ap.add_argument("--tier", default="ingest", help="Manifest tier to ingest (default: ingest)")
    ap.add_argument("--vendor", help="Only this vendor, for batched ingestion (e.g. ForgeRock)")
    ap.add_argument("--input", help="DEPRECATED directory mode. Requires --i-know-this-skips-the-manifest.")
    ap.add_argument("--i-know-this-skips-the-manifest", action="store_true",
                    help="Required to use --input. Ingests every file found, "
                         "including client RFPs and third-party proposals.")
    ap.add_argument("--output", default="./out", help="Staging output directory")
    ap.add_argument("--limit", type=int, default=None, help="Only process first N files")
    ap.add_argument("--dry-run", action="store_true",
                    help="Extract, classify and report. No embedding calls, no writes.")
    ap.add_argument("--skip-ocr", action="store_true", help="Skip Tesseract OCR")
    ap.add_argument("--skip-vision", action="store_true", help="Skip Qwen VL")
    args = ap.parse_args()

    if not args.manifest and not args.input:
        ap.error("one of --manifest (recommended) or --input is required")
    if args.input and not args.i_know_this_skips_the_manifest:
        ap.error(
            "--input walks a directory and ingests everything it finds. A content "
            "review of the Sales-SoWs bank moved 68 of 197 candidate files out of "
            "the ingest tier, including a client-authored RFP and a competitor's "
            "proposal. Use --manifest, or pass "
            "--i-know-this-skips-the-manifest if you really mean it.")

    # Validate env. Embedding and Supabase are not needed for a dry run.
    required = ["OPENROUTER_API_KEY"] if args.dry_run else [
        "OPENROUTER_API_KEY", "SUPABASE_URL", "SUPABASE_KEY", "ORG_ID"]
    missing = [k for k in required if not os.getenv(k)]
    if missing:
        log.error("Missing env vars: %s", ", ".join(missing))
        sys.exit(1)

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    rows: list[tuple[Path, Optional[ManifestRow]]] = []
    if args.manifest:
        if not args.root:
            ap.error("--root is required with --manifest (manifest paths are relative)")
        root = Path(os.path.expanduser(args.root))
        if not root.is_dir():
            ap.error(f"--root {root} is not a directory")
        rows = [(f, r) for f, r in load_manifest(
            Path(os.path.expanduser(args.manifest)), root,
            tier=args.tier, vendor=args.vendor)]
    else:
        input_dir = Path(os.path.expanduser(args.input))
        found = sorted(list(input_dir.rglob("*.docx")) + list(input_dir.rglob("*.pdf")))
        log.warning("DIRECTORY MODE: %d files, manifest review NOT applied", len(found))
        rows = [(f, None) for f in found]

    # Skip anything already in the corpus BEFORE spending time extracting it.
    if args.manifest and not args.dry_run:
        hashes = [r.sha256_16 for _, r in rows if r]
        seen = already_ingested(hashes)
        if seen:
            before = len(rows)
            rows = [(f, r) for f, r in rows if not (r and r.sha256_16 in seen)]
            log.info("dedup: %d of %d already ingested, skipping them",
                     before - len(rows), before)

    if args.limit:
        rows = rows[:args.limit]
    log.info("Processing %d file(s)%s", len(rows), " [DRY RUN]" if args.dry_run else "")

    results, failures = [], []
    for f, row in rows:
        try:
            p = process_file(f, output_dir, args.skip_ocr, args.skip_vision,
                             row=row, dry_run=args.dry_run)
            if p:
                results.append({
                    "file": f.name, "slug": p.slug, "client": p.client_name,
                    "vendor": p.iam_vendor, "proposal_type": p.proposal_type,
                    "year": p.year, "sha256_16": p.source_sha256,
                    "chunks_estimated": len(p.sections),
                })
        except Exception as e:  # noqa: BLE001 - one bad file must not end the batch
            log.exception("FAILED on %s: %s", f.name, e)
            failures.append({"file": f.name, "error": str(e)})

    summary = {"processed": len(results), "failed": len(failures),
               "dry_run": args.dry_run, "vendor_filter": args.vendor,
               "results": results, "failures": failures}
    with open(output_dir / "run_summary.json", "w") as fh:
        json.dump(summary, fh, indent=2)
    log.info("=" * 70)
    log.info("DONE. %d processed, %d failed. Summary: %s",
             len(results), len(failures), output_dir / "run_summary.json")
    if failures:
        log.warning("Failures: %s", ", ".join(x["file"] for x in failures[:10]))


if __name__ == "__main__":
    main()
